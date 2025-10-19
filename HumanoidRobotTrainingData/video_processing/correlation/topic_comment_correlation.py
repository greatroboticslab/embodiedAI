#!/usr/bin/env python3
"""
Correlate viewer comments to *topics* (title + description),
then generate a rebuildable DOCX report with coverage and top matches.

Inputs
  1) frames_path: Either a project root containing <video_id>/raw_frames
     subfolders, or a single .../raw_frames folder.
  2) comments_docx: A DOCX file formatted like:
       <YouTube URL>
       <one or more comment paragraphs>
       <next YouTube URL>
  3) topics_root (arg): A directory holding JSON topic files, default "topics".
     Each JSON should follow a schema like (flexible, only fields used are id/title/description):
       {
         "video_id": str,
         "title": str,
         "url": str,
         "created_at": str,
         "model": str,
         "topics": [
            {"id": "T1", "title": str, "description": str,
             "start_s": float|null, "end_s": float|null}
         ]
       }

Outputs (per video)
  - JSONL cache (resume-safe):
      <video_id>_topiccorr_n{n}_p{pen}_m{min}_<embed>_k{retrieval}.results.jsonl
  - DOCX report: <video_id>_correlation_topics_n{n}_p{pen}_m{min}.docx

Notes
  - Uses Sentence-Transformers for embeddings and an LLM (via your utils.llm_utils)
    to assign correlation scores 0–100 with a short reason.
  - Retrieval prefilters candidates by cosine similarity to the topic text.
  - Rebuild mode regenerates DOCX from prior JSONL without re-running the LLM.
  - Thumbnails are optional: we attempt to show representative frames if available.

Example
  python topic_comment_correlation.py /data/project/videos_root comments.docx \
      --topics_root /data/project/topics --model llama3.1 --min_score 60 --top_k 5

  # Rebuild reports only
  python topic_comment_correlation.py /data/project/videos_root comments.docx \
      --topics_root /data/project/topics --rebuild
"""

import os
import re
import json
import argparse
from urllib.parse import urlparse, parse_qs
import sys
from collections import defaultdict

# local imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from config import Config
from utils.llm_utils import generate_response, stream_parser
from docx import Document
from docx.shared import Inches

YES = {"yes", "y", "true", "t", "1"}

# -------------------- URL parsing --------------------

def get_video_id_from_url(text: str):
    text = text.strip()
    if not text or "youtu" not in text:
        return None
    m = re.search(r"https?://[^\s)\]>]+", text)
    if not m:
        return None
    url = re.sub(r'[:.,;!?"\)\]]+$', '', m.group(0))
    try:
        u = urlparse(url)
        host = (u.netloc or "").lower()
        if "youtu.be" in host:
            vid = u.path.strip("/").split("/")[0]
            return clean_video_id(vid)
        qs = parse_qs(u.query or "")
        if "v" in qs and qs["v"]:
            return clean_video_id(qs["v"][0])
        parts = [p for p in u.path.split("/") if p]
        if "embed" in parts:
            idx = parts.index("embed")
            if idx + 1 < len(parts):
                return clean_video_id(parts[idx + 1])
    except Exception:
        pass
    return None

def clean_video_id(vid: str):
    m = re.match(r"([A-Za-z0-9_-]{6,})", vid)
    return m.group(1) if m else vid

# -------------------- DOCX parsers --------------------

def parse_comments_docx(comments_docx_path: str):
    """
    Returns dict: {video_id: [{"url": url, "comment": text}, ...]}
    Layout: URL line, followed by one or more comment paragraphs, until next URL.
    """
    doc = Document(comments_docx_path)
    by_video = defaultdict(list)
    cur_vid, cur_url, buffer = None, None, []

    def flush():
        nonlocal buffer
        if cur_vid and buffer:
            text = " ".join(buffer).strip()
            if text:
                by_video[cur_vid].append({"url": cur_url, "comment": text})
        buffer = []

    for para in doc.paragraphs:
        t = para.text.strip()
        vid = get_video_id_from_url(t)
        if vid:
            flush()
            cur_vid, cur_url = vid, t
        elif t:
            buffer.append(t)
    flush()
    return dict(by_video)

# -------------------- Topic loading --------------------

def _find_topics_json(raw_frames_dir: str, topics_root: str):
    """Look for <topics_root>/<video_id>.topics.json, else scan raw_frames for *topics.json"""
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))
    # 1) preferred path under topics_root
    if topics_root:
        p = os.path.join(topics_root, f"{video_id}.topics.json")
        if os.path.isfile(p):
            return p
    # 2) sibling/within raw_frames folder
    for f in os.listdir(raw_frames_dir):
        if f.endswith(".topics.json") and video_id in f:
            return os.path.join(raw_frames_dir, f)
    # 3) one level up under a 'topics' folder
    candidate = os.path.join(os.path.dirname(os.path.dirname(raw_frames_dir)), "topics", f"{video_id}.topics.json")
    if os.path.isfile(candidate):
        return candidate
    return None

def parse_topics(topics_json_path: str):
    """
    Returns list of topics with keys: id, title, description, start_s, end_s.
    Missing fields are tolerated.
    """
    with open(topics_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    topics = []
    for t in data.get("topics", []):
        topics.append({
            "id": t.get("id") or t.get("topic_id") or f"T{len(topics)+1}",
            "title": t.get("title", "").strip(),
            "description": t.get("description", "").strip(),
            "start_s": t.get("start_s"),
            "end_s": t.get("end_s"),
        })
    return topics

# -------------------- Embeddings + helpers --------------------

def _embed_texts(texts, model_name="all-mpnet-base-v2"):
    """Returns L2-normalized embeddings. If sentence-transformers is missing, raises ImportError."""
    try:
        from sentence_transformers import SentenceTransformer
        import numpy as np  # noqa: F401
    except Exception as e:
        raise ImportError("sentence-transformers not installed") from e

    model = SentenceTransformer(model_name)
    X = model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
    return X


def pick_representative_frames(frame_names, k=3):
    """Evenly spaced picks from the list, preserving order."""
    if not frame_names:
        return []
    if k >= len(frame_names):
        return frame_names
    idxs = [int(round(i * (len(frame_names)-1) / (k-1))) for i in range(k)]
    seen = set()
    picks = []
    for idx in idxs:
        if idx not in seen:
            picks.append(frame_names[idx])
            seen.add(idx)
    return picks


def discover_raw_frames(frames_input: str):
    frames_input = os.path.abspath(frames_input)
    targets = []
    if os.path.basename(frames_input) == "raw_frames":
        targets.append(frames_input)
        return targets
    for entry in sorted(os.listdir(frames_input)):
        d = os.path.join(frames_input, entry, "raw_frames")
        if os.path.isdir(d):
            targets.append(d)
    return targets

# -------------------- LLM scoring --------------------

def ask_model_topic(model_name: str, topic_title: str, topic_desc: str, comment: str):
    """Return {"correlated": bool, "score": int, "reason": str} for TOPIC vs COMMENT."""
    topic_text = (topic_title or "").strip()
    if topic_desc:
        topic_text += "\n" + topic_desc.strip()

    prompt = (
        "You are checking whether a TOPIC (title + description) and a VIEWER COMMENT are strongly related.\n"
        "Strong = the comment directly discusses the same steps, objects, settings, parameters, claims, or goals described by the topic; not generic praise.\n\n"
        "Give a correlation confidence 0-100, and a brief reason citing specific overlaps.\n\n"
        f"TOPIC TITLE:\n{topic_title}\n\nTOPIC DESCRIPTION:\n{topic_desc}\n\nCOMMENT:\n{comment}\n\n"
        "Respond ONLY as compact JSON: {\"correlated\": true|false, \"score\": 0-100, \"reason\": \"brief explanation\"}"
    )
    stream = generate_response(model_name, prompt)
    text = "".join(stream_parser(stream)).strip()
    m = re.search(r"\{.*?\}", text, flags=re.S)
    jtxt = m.group(0) if m else text
    try:
        obj = json.loads(jtxt)
        corr = str(obj.get("correlated", False)).lower() in YES
        score = int(obj.get("score", 0))
        score = max(0, min(100, score))
        reason = str(obj.get("reason", "")).strip()
        return {"correlated": corr, "score": score, "reason": reason}
    except Exception:
        corr = "yes" in text.lower()
        m2 = re.search(r"(\d{1,3})", text)
        score = int(m2.group(1)) if m2 else 0
        score = max(0, min(100, score))
        return {"correlated": corr, "score": score, "reason": text[:400]}

# -------------------- IO helpers (resume-safe) --------------------

def results_path(raw_frames_dir: str, video_id: str, seg_tag: str, embed_model: str, retrieval_topk: int):
    safe_model = re.sub(r"[^A-Za-z0-9_-]+", "", embed_model or "mpnet")
    return os.path.join(
        raw_frames_dir,
        f"{video_id}_topiccorr_{seg_tag}_{safe_model}_k{retrieval_topk}.results.jsonl"
    )


def load_existing_results(results_file: str):
    done = {}
    if os.path.isfile(results_file):
        with open(results_file, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    rec = json.loads(line)
                    key = (rec.get("topic_id"),)
                    done[key] = rec
                except Exception:
                    continue
    return done


def append_result(results_file: str, record: dict):
    with open(results_file, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")

# -------------------- DOCX helpers --------------------

def _add_thumbnails_row(doc, base_dir, frame_list, thumb_width_in=1.3):
    tbl = doc.add_table(rows=1, cols=max(1, len(frame_list)))
    tbl.autofit = True
    for col, relpath in enumerate(frame_list):
        cell = tbl.rows[0].cells[col]
        img_path = os.path.join(base_dir, relpath)
        if os.path.exists(img_path):
            try:
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=Inches(thumb_width_in))
            except Exception:
                cell.text = relpath
        else:
            cell.text = relpath
    return tbl


def build_docx_from_results(out_docx: str, topics: list, results_map: dict,
                            min_score: int, top_k: int, model_name: str,
                            total_comments_available: int | None = None):
    doc = Document()
    doc.add_heading("Comment–Topic Correlation Report", 0)
    doc.add_paragraph(f"Model: {model_name}")

    # Aggregate
    total_pairs = sum(rec.get("checked_pairs", 0) for rec in results_map.values())
    base_dir = os.path.dirname(out_docx)

    # Topic metadata map for quick lookup (stringify IDs for safety)
    topic_map = {str(t.get("id")): {
                    "title": t.get("title", ""),
                    "description": t.get("description", "")
                 } for t in topics}

    # Per-comment index (for coverage summary)
    comments_index = {}
    for t in topics:
        tid = t["id"]
        rec = results_map.get((tid,), {})
        for c in rec.get("candidates", []):
            key = (c.get("url", ""), c.get("comment", ""))
            entry = comments_index.setdefault(key, {
                "comment": c.get("comment", ""),
                "url": c.get("url", ""),
                "topics": set(),
                "max_score": 0,
                "per_topic": {}
            })
            s = float(c.get("score", 0))
            r = str(c.get("reason", "") or "")
            entry["per_topic"][tid] = {"score": s, "reason": r, "rep_frames": t.get("rep_frames", [])}
            if c.get("correlated") and int(s) >= min_score:
                entry["topics"].add(tid)
                entry["max_score"] = max(entry["max_score"], int(s))

    unique_corr_count = sum(1 for v in comments_index.values() if v["topics"])
    pct = None
    if total_comments_available and total_comments_available > 0:
        pct = 100.0 * unique_corr_count / total_comments_available

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Topics processed: {len(topics)}")
    if total_comments_available is not None:
        doc.add_paragraph(f"Comments available for this video: {total_comments_available}")
    doc.add_paragraph(
        f"Comments correlated to ≥ {min_score}: {unique_corr_count}" + (f" ({pct:.1f}%)" if pct is not None else "")
    )
    doc.add_paragraph(f"Total candidate pairs checked: {total_pairs}")

    # ========================= REPLACE WITH: rank topics by average score =========================
    # Compute average score from correlated comments (score >= min_score) per topic
    topic_stats = []  # list of (tid, title, avg_score, contributions)
    for t in topics:
        tid = t["id"]
        title = topic_map.get(str(tid), {}).get("title", "(untitled)")
        rec = results_map.get((tid,), {})
        scores = []
        for c in rec.get("candidates", []):
            s = int(c.get("score", 0))
            if c.get("correlated") and s >= min_score:
                scores.append(s)
        if scores:
            avg = sum(scores) / len(scores)
            contrib = len(scores)
        else:
            avg = 0.0
            contrib = 0
        topic_stats.append((tid, title, avg, contrib))

    # Sort by average score (desc), then title/id to stabilize
    ranked_desc = sorted(topic_stats, key=lambda x: (x[2], x[1], str(x[0])), reverse=True)
    ranked_asc  = sorted(topic_stats, key=lambda x: (x[2], x[1], str(x[0])))  # for bottom-3

    top3 = ranked_desc[:3]
    bottom3 = ranked_asc[:3]

    # Insert section
    doc.add_paragraph("")  # spacing
    doc.add_paragraph("Most/Least correlated topics (by average score across correlated comments):")
    doc.add_paragraph("  Top 3:")
    if top3:
        for tid, title, avg, contrib in top3:
            doc.add_paragraph(f"    • Topic {tid}: {title} — avg {avg:.2f}, contributions {contrib}")
    else:
        doc.add_paragraph("    • (none)")

    doc.add_paragraph("  Bottom 3:")
    if bottom3:
        for tid, title, avg, contrib in bottom3:
            doc.add_paragraph(f"    • Topic {tid}: {title} — avg {avg:.2f}, contributions {contrib}")
    else:
        doc.add_paragraph("    • (none)")
    # ======================= END: rank topics by average score =======================


    # Per-comment coverage
    if comments_index:
        doc.add_paragraph("Correlated comment coverage (all comments):")
        TOPK = 10
        all_list = sorted(
            comments_index.values(),
            key=lambda x: (len(x.get("topics", [])), x.get("max_score", float("-inf")), x.get("comment", "")),
            reverse=True
        )
        for e in all_list:
            coverage = len(e.get("topics", []))
            max_score = e.get("max_score")
            header = f"- matched {coverage} topic(s)"
            if max_score is not None:
                header += f", max score {max_score:.4f}"
            header += f": {e.get('comment','')}"
            doc.add_paragraph(header)
            if e.get("url"):
                doc.add_paragraph(f"  Source: {e['url']}")

            items = list(e.get("per_topic", {}).items())  # (tid, info)

            # Top-10
            top_t = sorted(items, key=lambda kv: (kv[1].get("score", 0.0), str(kv[0])), reverse=True)[:TOPK]
            doc.add_paragraph("  Top correlated topics:")
            if top_t:
                for tid, info in top_t:
                    _add_thumbnails_row(doc, base_dir, info.get("rep_frames", []), thumb_width_in=1.2)
                    title = topic_map.get(str(tid), {}).get("title", "(untitled)")
                    desc  = topic_map.get(str(tid), {}).get("description", "")
                    p = doc.add_paragraph()
                    p.add_run(f"Topic {tid}: {title} — score {float(info.get('score', 0.0)):.4f}")
                    if desc:
                        doc.add_paragraph(f"    Description: {desc}")
                    if info.get("reason"):
                        doc.add_paragraph(f"    Reason: {info['reason']}")
            else:
                doc.add_paragraph("    • (none)")

            # Bottom-10
            bot_t = sorted(items, key=lambda kv: (kv[1].get("score", 0.0), str(kv[0])))[:TOPK]
            doc.add_paragraph("  Least correlated topics:")
            if bot_t:
                for tid, info in bot_t:
                    _add_thumbnails_row(doc, base_dir, info.get("rep_frames", []), thumb_width_in=1.2)
                    title = topic_map.get(str(tid), {}).get("title", "(untitled)")
                    desc  = topic_map.get(str(tid), {}).get("description", "")
                    p = doc.add_paragraph()
                    p.add_run(f"Topic {tid}: {title} — score {float(info.get('score', 0.0)):.4f}")
                    if desc:
                        doc.add_paragraph(f"    Description: {desc}")
                    if info.get("reason"):
                        doc.add_paragraph(f"    Reason: {info['reason']}")
            else:
                doc.add_paragraph("    • (none)")

    # Per-topic details
    for t in topics:
        tid = t["id"]
        title = t.get("title", "").strip()
        desc = t.get("description", "").strip()
        rep = t.get("rep_frames", [])

        doc.add_heading(f"Topic {tid}: {title}", level=1)
        if rep:
            _add_thumbnails_row(doc, base_dir, rep, thumb_width_in=1.4)
        if desc:
            doc.add_paragraph(f"Description: {desc}")

        rec = results_map.get((tid,), {})
        candidates = rec.get("candidates", [])
        kept = [c for c in candidates if c.get("score", 0) >= min_score and c.get("correlated", False)]
        kept.sort(key=lambda x: x.get("score", 0), reverse=True)
        kept = kept[:top_k]

        if kept:
            for r in kept:
                doc.add_paragraph(f"Comment (score {r['score']}): {r['comment']}")
                if r.get("reason"): doc.add_paragraph(f"Reason: {r['reason']}")
                if r.get("url"): doc.add_paragraph(f"Source: {r['url']}")
        else:
            doc.add_paragraph("No strongly correlated comments found for this topic.")

    doc.save(out_docx)


# -------------------- Frame discovery for thumbnails --------------------

def list_all_frames(raw_frames_dir: str):
    files = [f for f in os.listdir(raw_frames_dir) if f.lower().endswith((".jpg", ".jpeg", ".png"))]
    files.sort()
    return files

# -------------------- Core processing --------------------

def process_one_video(raw_frames_dir: str, comments_by_video: dict, topics_root: str, model: str,
                      min_score: int, top_k_scan: int,
                      seg_n_sections: int | None, seg_penalty: int | None,
                      seg_min_size: int, embed_model: str,
                      retrieval_topk: int, rep_k: int):
    """Segment knobs (n/penalty/min) are kept for tag consistency only; topics are pre-defined."""
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))

    topics_json = _find_topics_json(raw_frames_dir, topics_root)
    if not topics_json:
        print(f"[SKIP] {video_id}: topics JSON not found (looked under {topics_root} and nearby).")
        return

    topics = parse_topics(topics_json)
    if not topics:
        print(f"[SKIP] {video_id}: no topics parsed.")
        return

    # Representative frames per topic (best effort): evenly from whole folder
    # If you later add per-topic frame ranges, swap to filtered subset here.
    all_frames = list_all_frames(raw_frames_dir)
    for t in topics:
        t["rep_frames"] = pick_representative_frames(all_frames, k=rep_k)

    comments_for_video = comments_by_video.get(video_id, [])
    if not comments_for_video:
        comments_for_video = [c for v in comments_by_video.values() for c in v]
        print(f"[WARN] {video_id}: no comments matched by video id; falling back to ALL comments ({len(comments_for_video)}).")
    else:
        print(f"[INFO] {video_id}: {len(comments_for_video)} comments found for this video.")

    # Pre-embed comments for retrieval (if possible)
    comment_texts = [c["comment"] for c in comments_for_video]
    C = None
    retrieval_ok = False
    try:
        C = _embed_texts(comment_texts, model_name=embed_model)
        retrieval_ok = True
    except ImportError:
        print(f"[WARN] {video_id}: sentence-transformers missing; skipping fast retrieval (LLM will score all comments).")

    # Resume state
    seg_tag = f"n{seg_n_sections or 0}_p{seg_penalty or 0}_m{seg_min_size}"
    res_file = results_path(raw_frames_dir, video_id, seg_tag, embed_model, retrieval_topk)
    existing = load_existing_results(res_file)
    print(f"[RESUME] {video_id}: {len(existing)}/{len(topics)} topics already processed.")

    import numpy as np

    # Process topics
    for t in topics:
        tid = t["id"]
        key = (tid,)
        if key in existing:
            continue

        title = t.get("title", "")
        desc = t.get("description", "")
        topic_text = (title or "").strip()
        if desc:
            topic_text += "\n" + desc.strip()

        checked_pairs = 0
        candidates = []

        # choose candidate comments via retrieval
        candidate_indices = range(len(comment_texts))
        if retrieval_ok and retrieval_topk > 0 and topic_text:
            t_vec = _embed_texts([topic_text], model_name=embed_model)[0]
            sims = np.dot(C, t_vec)
            idx = np.argsort(-sims)[:min(retrieval_topk, len(comment_texts))]
            candidate_indices = idx

        for i in candidate_indices:
            item = comments_for_video[i]
            res = ask_model_topic(model, title, desc, item["comment"])
            checked_pairs += 1
            candidates.append({
                "correlated": bool(res["correlated"]),
                "score": int(res["score"]),
                "reason": res.get("reason", ""),
                "comment": item["comment"],
                "url": item["url"]
            })

        record = {
            "video_id": video_id,
            "topic_id": tid,
            "title": title,
            "checked_pairs": checked_pairs,
            "candidates": candidates,
        }
        append_result(res_file, record)
        existing[key] = record
        print(f"[OK] {video_id}: topic {tid} processed ({checked_pairs} pairs).")

    # Build DOCX
    out_docx = os.path.join(raw_frames_dir, f"{video_id}_correlation_topics_{seg_tag}.docx")
    build_docx_from_results(
        out_docx, topics, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[DONE] {video_id}: report saved -> {out_docx}")


def rebuild_docx_only(raw_frames_dir: str, comments_by_video: dict, topics_root: str, model: str,
                      min_score: int, top_k_scan: int,
                      seg_n_sections: int | None, seg_penalty: int | None,
                      seg_min_size: int, embed_model: str,
                      retrieval_topk: int, rep_k: int):
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))

    topics_json = _find_topics_json(raw_frames_dir, topics_root)
    if not topics_json:
        print(f"[SKIP] {video_id}: topics JSON not found for rebuild.")
        return
    topics = parse_topics(topics_json)
    if not topics:
        print(f"[SKIP] {video_id}: no topics parsed for rebuild.")
        return

    # thumbnails best-effort
    all_frames = list_all_frames(raw_frames_dir)
    for t in topics:
        t["rep_frames"] = pick_representative_frames(all_frames, k=rep_k)

    seg_tag = f"n{seg_n_sections or 0}_p{seg_penalty or 0}_m{seg_min_size}"
    res_file = results_path(raw_frames_dir, video_id, seg_tag, embed_model, retrieval_topk)
    if not os.path.isfile(res_file):
        print(f"[SKIP] {video_id}: no results file found to rebuild.")
        return

    existing = load_existing_results(res_file)
    comments_for_video = comments_by_video.get(video_id, [])

    out_docx = os.path.join(raw_frames_dir, f"{video_id}_correlation_topics_{seg_tag}.docx")
    build_docx_from_results(
        out_docx, topics, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[REBUILT] {video_id}: report saved -> {out_docx}")

# -------------------- CLI --------------------

def main():
    ap = argparse.ArgumentParser(description="Correlate viewer comments to topics (resume-safe) and build a DOCX report.")
    ap.add_argument("--frames_path", default="../frames", help="Path to frames root (containing <video_id>/raw_frames) OR a single raw_frames folder.")
    ap.add_argument("--comments_docx", default="../../../Data For Classes-Analysis/final_merged_data/final_merged_class_data.docx", help="DOCX with video URLs and comments.")

    ap.add_argument("--topics_root", default="topics", help="Directory containing <video_id>.topics.json files.")
    ap.add_argument("--model", default=None, help="LLM name for scoring (defaults to Config.ollama_models[0]).")
    ap.add_argument("--min_score", type=int, default=60, help="Minimum score to include a matched comment in the final report.")
    ap.add_argument("--top_k", type=int, default=5, help="Top-K comments per topic shown in the final report.")

    # Tags only (for naming parity with your section pipeline)
    ap.add_argument("--seg_n_sections", type=int, default=None, help="Unused for topics; kept for tag naming consistency.")
    ap.add_argument("--seg_penalty", type=int, default=15, help="Unused for topics; kept for tag naming consistency.")
    ap.add_argument("--seg_min_size", type=int, default=6, help="Unused for topics; kept for tag naming consistency.")

    ap.add_argument("--embed_model", default="all-mpnet-base-v2", help="SentenceTransformer model for embeddings.")
    ap.add_argument("--retrieval_topk", type=int, default=50, help="How many candidate comments to retrieve per topic before LLM re-score (0 = score all).")
    ap.add_argument("--rep_k", type=int, default=3, help="Representative thumbnails per topic in the report.")

    ap.add_argument("--rebuild", action="store_true", help="Rebuild DOCX from existing results without running the model.")

    args = ap.parse_args()
    model = args.model or Config.ollama_models[0]

    comments_by_video = parse_comments_docx(args.comments_docx)
    if not comments_by_video:
        print("[FATAL] No comments parsed from comments DOCX.")
        return

    targets = discover_raw_frames(args.frames_path)
    if not targets:
        print("[FATAL] No raw_frames folders found.")
        return

    print(f"[INFO] Found {len(targets)} video(s) to process.")

    for raw_frames in targets:
        print(f"\n=== Processing: {os.path.basename(os.path.dirname(raw_frames))} ===")
        try:
            if args.rebuild:
                rebuild_docx_only(
                    raw_frames, comments_by_video, args.topics_root, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    seg_n_sections=args.seg_n_sections, seg_penalty=args.seg_penalty,
                    seg_min_size=args.seg_min_size, embed_model=args.embed_model,
                    retrieval_topk=args.retrieval_topk, rep_k=args.rep_k
                )
            else:
                process_one_video(
                    raw_frames, comments_by_video, args.topics_root, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    seg_n_sections=args.seg_n_sections, seg_penalty=args.seg_penalty,
                    seg_min_size=args.seg_min_size, embed_model=args.embed_model,
                    retrieval_topk=args.retrieval_topk, rep_k=args.rep_k
                )
        except KeyboardInterrupt:
            print("\n[INTERRUPTED] Stopping cleanly. You can rerun to resume.")
            break
        except Exception as e:
            print(f"[ERROR] {raw_frames}: {e}")

if __name__ == "__main__":
    main()
