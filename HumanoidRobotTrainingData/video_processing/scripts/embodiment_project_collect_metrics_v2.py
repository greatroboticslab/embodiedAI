#!/usr/bin/env python3
"""
Collect per-video metrics for an Embodied vs. Conventional video study, save a JSON summary,
and generate a DOCX table report. Supports merging an external per-video data file
(e.g., an Excel/CSV/JSON sheet with a row per video_id) and appending selected columns
from that file to the DOCX table.

Inputs (paths you pass via CLI):
  --videos_root       Path to the folder holding raw videos (optional; used for duration if ffprobe available)
  --frames_root       Path to the folder holding extracted frames. Expected structure: <frames_root>/<video_id>/raw_frames/
                      NOTE: The correlation result files are expected near frames and can be .results.jsonl or .json.
  --transcripts_root  Path to the folder with transcripts. Accepts either <video_id>.txt OR <video_id>.json (with "segments").
  --topics_root       Path to the folder with topic JSON files (default name: <video_id>.topics.json).
  --out_dir           Output directory to write results JSON and DOCX.
  --extra_file        (Optional) Path to an external per-video data file (XLSX/CSV/JSON) with a 'video_id' column/key.
  --extra_key         (Optional) Key/column name to join on (default: video_id).
  --extra_cols        (Optional) Comma-separated columns from --extra_file to append to the DOCX; if omitted, include all non-key columns.

Run the same command twice—once for Conventional, once for Embodied.

What this script computes per video_id (best-effort, robust to missing pieces):
  - frames_count: count of image files in raw_frames
  - transcript_word_count: words in <video_id>.txt or sum over JSON segments
  - topics_count: len of topics array in <video_id>.topics.json
  - comments_count: number of UNIQUE comments found in correlation results (JSONL/JSON) near frames
  - corr_scores: list of correlation scores (if present) and summary stats (avg, median)
  - num_correlated_comments_thresholded: number of scores >= --corr_threshold
  - duration_seconds: if video exists and ffprobe is available, duration in seconds
  - extra: any extra fields merged from --extra_file

Outputs:
  - <out_dir>/<label>_metrics.json
  - <out_dir>/<label>_metrics.docx
"""

import argparse
import json
import os
import re
import subprocess
import glob
from dataclasses import dataclass, asdict, field
from statistics import mean, median
from typing import Dict, List, Optional, Tuple, Any

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    raise SystemExit("python-docx is required. Install with: pip install python-docx")

IMG_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}

# ---------------------- Utilities ----------------------

def list_video_ids_from_frames(frames_root: str) -> List[str]:
    """Use frames_root/<video_id>/raw_frames as canonical list of videos."""
    if not os.path.isdir(frames_root):
        return []
    vids = []
    for name in os.listdir(frames_root):
        vid_dir = os.path.join(frames_root, name)
        if not os.path.isdir(vid_dir):
            continue
        raw_frames = os.path.join(vid_dir, "raw_frames")
        if os.path.isdir(raw_frames):
            vids.append(name)
    return sorted(vids)


def count_frames(raw_frames_dir: str) -> int:
    if not os.path.isdir(raw_frames_dir):
        return 0
    return sum(1 for fn in os.listdir(raw_frames_dir) if os.path.splitext(fn)[1].lower() in IMG_EXTS)


def read_transcript_word_count(transcripts_root: str, vid: str) -> int:
    """Try <vid>.txt then <vid>.json with segments[].text."""
    txt_path = os.path.join(transcripts_root, f"{vid}.txt")
    if os.path.isfile(txt_path):
        try:
            with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return len(re.findall(r"\b\w+\b", text))
        except Exception:
            return 0
    json_path = os.path.join(transcripts_root, f"{vid}.json")
    if os.path.isfile(json_path):
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            segs = data.get("segments") or []
            text = " ".join(str(s.get("text", "")) for s in segs)
            return len(re.findall(r"\b\w+\b", text))
        except Exception:
            return 0
    return 0


def read_topics_count(topics_root: str, vid: str) -> int:
    path = os.path.join(topics_root, f"{vid}.topics.json")
    if not os.path.isfile(path):
        if not os.path.isdir(topics_root):
            return 0
        # fallback: any json starting with vid and containing 'topics'
        for fn in os.listdir(topics_root):
            if not (fn.startswith(vid) and fn.endswith('.json')):
                continue
            try:
                with open(os.path.join(topics_root, fn), 'r', encoding='utf-8') as f:
                    d = json.load(f)
                if isinstance(d, dict) and isinstance(d.get('topics'), list):
                    return len(d['topics'])
            except Exception:
                continue
        return 0
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        topics = data.get('topics') or []
        return len(topics)
    except Exception:
        return 0


def list_correlation_result_paths(frames_video_dir: str) -> List[str]:
    """Return correlation result files (.results.jsonl preferred) near frames."""
    out: List[str] = []
    for folder in [frames_video_dir, os.path.join(frames_video_dir, 'raw_frames')]:
        if not os.path.isdir(folder):
            continue
        out.extend(glob.glob(os.path.join(folder, "*_topiccorr_*.results.jsonl")))
        out.extend(glob.glob(os.path.join(folder, "*.results.jsonl")))
        out.extend(glob.glob(os.path.join(folder, "*.jsonl")))
        out.extend(glob.glob(os.path.join(folder, "*.json")))
    # Remove obvious non-correlation JSONs by name if needed later
    return sorted(set(out))


def parse_correlation_results(paths: List[str]) -> Tuple[int, List[float]]:
    """Aggregate (unique comments_count, scores[]) across correlation result files.
    Supports JSONL with per-topic objects containing 'candidates'.
    """
    comments_seen = set()
    scores: List[float] = []

    def _get_comment_text(rec: Dict[str, Any]) -> Optional[str]:
        for k in ("comment", "comment_text", "text", "body"):
            v = rec.get(k)
            if isinstance(v, str) and v.strip():
                return v.strip()
        return None

    for path in paths:
        if not os.path.isfile(path):
            continue
        # JSONL
        if path.lower().endswith('.jsonl'):
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                    except Exception:
                        continue
                    cands = obj.get('candidates') or []
                    if isinstance(cands, list):
                        for c in cands:
                            if not isinstance(c, dict):
                                continue
                            ct = _get_comment_text(c)
                            if ct:
                                comments_seen.add(ct)
                            sc = c.get('score')
                            if isinstance(sc, (int, float)):
                                scores.append(float(sc))
            continue
        # JSON
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception:
            continue
        if isinstance(data, list):
            for obj in data:
                if not isinstance(obj, dict):
                    continue
                for c in obj.get('candidates', []) or []:
                    if not isinstance(c, dict):
                        continue
                    ct = _get_comment_text(c)
                    if ct:
                        comments_seen.add(ct)
                    sc = c.get('score')
                    if isinstance(sc, (int, float)):
                        scores.append(float(sc))
        elif isinstance(data, dict):
            for c in data.get('candidates', []) or []:
                if not isinstance(c, dict):
                    continue
                ct = _get_comment_text(c)
                if ct:
                    comments_seen.add(ct)
                sc = c.get('score')
                if isinstance(sc, (int, float)):
                    scores.append(float(sc))

    return len(comments_seen), scores


def get_duration_seconds(videos_root: Optional[str], vid: str) -> Optional[float]:
    if not videos_root or not os.path.isdir(videos_root):
        return None
    cand = None
    for fn in os.listdir(videos_root):
        if not fn.startswith(vid):
            continue
        if os.path.splitext(fn)[1].lower() in {'.mp4', '.mkv', '.mov', '.avi', '.webm'}:
            cand = os.path.join(videos_root, fn)
            break
    if not cand or not os.path.isfile(cand):
        return None
    try:
        proc = subprocess.run([
            'ffprobe', '-v', 'error', '-select_streams', 'v:0',
            '-show_entries', 'format=duration', '-of', 'default=nw=1:nk=1', cand
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=10)
        out = proc.stdout.strip()
        if out:
            return float(out)
    except Exception:
        return None
    return None

# ---------------------- External data merge ----------------------

def load_extra_file(path: Optional[str], key: str) -> Tuple[Dict[str, Dict[str, Any]], List[str]]:
    """Load an optional external per-video table and index by 'key'.
    Supports .xlsx/.xls (requires pandas+openpyxl), .csv, .json.
    Returns (mapping video_id->row_without_key, columns_in_file).
    """
    if not path:
        return {}, []
    if not os.path.isfile(path):
        print(f"[WARN] Extra file not found: {path}")
        return {}, []
    ext = os.path.splitext(path)[1].lower()
    data_rows: List[Dict[str, Any]] = []
    try:
        if ext in ('.xlsx', '.xls'):
            try:
                import pandas as pd  # type: ignore
            except Exception:
                raise SystemExit("Reading Excel requires pandas and openpyxl. Install with: pip install pandas openpyxl")
            df = pd.read_excel(path)
            data_rows = df.to_dict(orient='records')
        elif ext == '.csv':
            import csv
            with open(path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                data_rows = list(reader)
        elif ext == '.json':
            with open(path, 'r', encoding='utf-8') as f:
                obj = json.load(f)
            if isinstance(obj, list):
                data_rows = obj
            elif isinstance(obj, dict) and isinstance(obj.get('rows'), list):
                data_rows = obj['rows']
            else:
                # attempt dict-of-dicts
                data_rows = [v for v in obj.values() if isinstance(v, dict)]
        else:
            print(f"[WARN] Unsupported extra file type: {ext}")
            return {}, []
    except Exception as e:
        print(f"[WARN] Failed to read extra file {path}: {e}")
        return {}, []

    emap: Dict[str, Dict[str, Any]] = {}
    cols: List[str] = []
    for row in data_rows:
        if not isinstance(row, dict):
            continue
        if not cols:
            cols = [c for c in row.keys()]
        vid_val = row.get(key)
        if vid_val is None:
            continue
        # store all other fields; normalize video_id to str
        emap[str(vid_val)] = {k: v for k, v in row.items() if k != key}
    return emap, cols

# ---------------------- Data Classes ----------------------

@dataclass
class VideoMetrics:
    video_id: str
    frames_count: int
    transcript_word_count: int
    topics_count: int
    comments_count: int
    corr_avg: Optional[float]
    corr_median: Optional[float]
    corr_n: int
    corr_n_ge_threshold: int
    duration_seconds: Optional[float]
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Aggregate:
    label: str
    num_videos: int
    avg_frames: Optional[float]
    avg_transcript_words: Optional[float]
    avg_topics: Optional[float]
    avg_comments: Optional[float]
    avg_corr_avg: Optional[float]
    median_corr_avg: Optional[float]

# ---------------------- Core ----------------------

def compute_metrics_for_set(
    videos_root: Optional[str],
    frames_root: str,
    transcripts_root: str,
    topics_root: str,
    label: str,
    corr_threshold: float = 0.0,
    extra_map: Optional[Dict[str, Dict[str, Any]]] = None,
) -> Tuple[List[VideoMetrics], Aggregate]:

    video_ids = list_video_ids_from_frames(frames_root)
    metrics: List[VideoMetrics] = []
    corr_avgs: List[float] = []

    for vid in video_ids:
        vid_dir = os.path.join(frames_root, vid)
        raw_frames_dir = os.path.join(vid_dir, 'raw_frames')
        frames_count = count_frames(raw_frames_dir)
        transcript_words = read_transcript_word_count(transcripts_root, vid)
        topics_count = read_topics_count(topics_root, vid)

        corr_paths = list_correlation_result_paths(vid_dir)
        comments_count, scores = parse_correlation_results(corr_paths)

        corr_avg = mean(scores) if scores else None
        corr_median = median(scores) if scores else None
        corr_n = len(scores)
        corr_n_ge = sum(1 for s in scores if s >= corr_threshold)
        dur_s = get_duration_seconds(videos_root, vid)

        extra_row = (extra_map or {}).get(vid, {})

        vm = VideoMetrics(
            video_id=vid,
            frames_count=frames_count,
            transcript_word_count=transcript_words,
            topics_count=topics_count,
            comments_count=comments_count,
            corr_avg=corr_avg,
            corr_median=corr_median,
            corr_n=corr_n,
            corr_n_ge_threshold=corr_n_ge,
            duration_seconds=dur_s,
            extra=extra_row,
        )
        metrics.append(vm)
        if corr_avg is not None:
            corr_avgs.append(corr_avg)

    def _avg(vals: List[int]) -> Optional[float]:
        return mean(vals) if vals else None

    agg = Aggregate(
        label=label,
        num_videos=len(metrics),
        avg_frames=_avg([m.frames_count for m in metrics]),
        avg_transcript_words=_avg([m.transcript_word_count for m in metrics]),
        avg_topics=_avg([m.topics_count for m in metrics]),
        avg_comments=_avg([m.comments_count for m in metrics]),
        avg_corr_avg=mean(corr_avgs) if corr_avgs else None,
        median_corr_avg=median(corr_avgs) if corr_avgs else None,
    )

    return metrics, agg

# ---------------------- DOCX ----------------------

def build_docx_table(out_docx_path: str, metrics: List[VideoMetrics], agg: Aggregate, corr_threshold: float,
                     extra_cols: Optional[List[str]] = None):
    doc = Document()

    # Title
    title = doc.add_paragraph(f"{agg.label} Set Metrics")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(16)
    run.bold = True

    # Summary paragraph
    p = doc.add_paragraph()
    p.add_run("Summary\n").bold = True
    summary_lines = [
        f"Videos: {agg.num_videos}",
        f"Avg frames: {round(agg.avg_frames, 2) if agg.avg_frames is not None else '—'}",
        f"Avg transcript words: {round(agg.avg_transcript_words, 2) if agg.avg_transcript_words is not None else '—'}",
        f"Avg topics: {round(agg.avg_topics, 2) if agg.avg_topics is not None else '—'}",
        f"Avg comments: {round(agg.avg_comments, 2) if agg.avg_comments is not None else '—'}",
        f"Avg of per-video correlation averages: {round(agg.avg_corr_avg, 3) if agg.avg_corr_avg is not None else '—'}",
        f"Median of per-video correlation averages: {round(agg.median_corr_avg, 3) if agg.median_corr_avg is not None else '—'}",
        f"Threshold for counted correlations: >= {corr_threshold}",
    ]
    p.add_run("\n".join(summary_lines))

    # Table headers
    base_headers = [
        "video_id", "frames", "topics", "comments",
        "corr_avg", "corr_median", "corr_n", f"corr_n≥{corr_threshold}",
        "transcript_words", "duration_s"
    ]
    headers = list(base_headers)

    # Append extra columns (if any)
    extra_cols = [c for c in (extra_cols or []) if c and c not in headers]
    headers.extend(extra_cols)

    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for m in metrics:
        row = table.add_row().cells
        row[0].text = m.video_id
        row[1].text = str(m.frames_count)
        row[2].text = str(m.topics_count)
        row[3].text = str(m.comments_count)
        row[4].text = "—" if m.corr_avg is None else f"{m.corr_avg:.3f}"
        row[5].text = "—" if m.corr_median is None else f"{m.corr_median:.3f}"
        row[6].text = str(m.corr_n)
        row[7].text = str(m.corr_n_ge_threshold)
        row[8].text = str(m.transcript_word_count)
        row[9].text = "—" if m.duration_seconds is None else f"{m.duration_seconds:.2f}"
        # populate extra cols
        for idx, col in enumerate(extra_cols, start=len(base_headers)):
            val = m.extra.get(col, "—")
            row[idx].text = str(val)

    doc.save(out_docx_path)

# ---------------------- Main ----------------------

def main():
    ap = argparse.ArgumentParser(description="Collect metrics for embodied vs conventional video sets and export JSON + DOCX table")
    ap.add_argument('--videos_root', type=str, default="../rawvideos/downloaded_videos", help='Raw videos folder (optional; used for duration via ffprobe)')
    ap.add_argument('--frames_root', type=str, default="../frames", help='Frames root: <frames_root>/<video_id>/raw_frames')
    ap.add_argument('--transcripts_root', type=str, default="../transcription/transcripts", help='Transcripts root containing <video_id>.txt or .json')
    ap.add_argument('--topics_root', type=str, default="../correlation/topics", help='Topics root containing <video_id>.topics.json')
    ap.add_argument('--out_dir', type=str, default="../output/embodied", help='Output directory for JSON + DOCX')
    ap.add_argument('--label', type=str, default="Embodied", help='Label to tag this run (e.g., Conventional or Embodied)')
    ap.add_argument('--corr_threshold', type=float, default=60, help='Threshold for counting a correlation as present')

    # New: external data file
    ap.add_argument('--extra_file', type=str, default="../../../Data For Classes-Analysis/student_video_analysis/sa_embodied_videos.xlsx", help='Optional XLSX/CSV/JSON with additional per-video fields (must include a video_id column/key).')
    ap.add_argument('--extra_key', type=str, default='video_id', help='Column/key name in --extra_file to join on (default: video_id).')
    ap.add_argument('--extra_cols', type=str, default="Total Likes", help='Comma-separated list of columns from --extra_file to append to the DOCX table. If omitted, include all non-key columns.')

    args = ap.parse_args()

    os.makedirs(args.out_dir, exist_ok=True)

    extra_map, extra_all_cols = load_extra_file(args.extra_file, args.extra_key)
    if args.extra_cols:
        extra_cols = [c.strip() for c in args.extra_cols.split(',') if c.strip()]
    else:
        extra_cols = [c for c in extra_all_cols if c != args.extra_key] if extra_all_cols else []

    metrics, agg = compute_metrics_for_set(
        videos_root=args.videos_root,
        frames_root=args.frames_root,
        transcripts_root=args.transcripts_root,
        topics_root=args.topics_root,
        label=args.label,
        corr_threshold=args.corr_threshold,
        extra_map=extra_map,
    )

    # Save JSON
    out_json = os.path.join(args.out_dir, f"{args.label}_metrics.json")
    blob = {
        'label': agg.label,
        'aggregate': asdict(agg),
        'corr_threshold': args.corr_threshold,
        'videos': [asdict(m) for m in metrics],
        'extra_cols_in_docx': extra_cols,
        'extra_file_used': args.extra_file or None,
    }
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(blob, f, ensure_ascii=False, indent=2)

    # Save DOCX
    out_docx = os.path.join(args.out_dir, f"{args.label}_metrics.docx")
    build_docx_table(out_docx, metrics, agg, args.corr_threshold, extra_cols=extra_cols)

    print(f"Wrote: {out_json}")
    print(f"Wrote: {out_docx}")


if __name__ == '__main__':
    main()
