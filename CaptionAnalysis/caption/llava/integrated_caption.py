import os
import re
import json
import signal
import argparse
from contextlib import contextmanager
from docx import Document
from docx.shared import Inches
from config import Config
from utils.llm_utils import stream_parser, generate_response


class TimeoutException(Exception): pass


# Cross-platform timeout
from contextlib import contextmanager
import signal


class TimeoutException(Exception): pass


@contextmanager
def timeout(time):
    if hasattr(signal, 'SIGALRM'):
        def raise_timeout(signum, frame):
            raise TimeoutException

        signal.signal(signal.SIGALRM, raise_timeout)
        signal.alarm(time)
        try:
            yield
        finally:
            signal.alarm(0)
    else:
        # Timeout not supported on this platform (e.g., Windows)
        yield


llama_model = Config.ollama_models[0]


def load_captions_from_json(json_path):
    """Load captions from a JSON file into a dictionary {frame_id: caption}."""
    captions = {}
    if os.path.exists(json_path):
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                for frame in data.get('frames', []):
                    captions[frame['frame_id'].lower()] = frame['caption']
        except Exception as e:
            print(f"⚠️ Error reading JSON {json_path}: {e}")
    return captions


def generate_integrated_caption(llava_caption, minicpm_caption, mock=False):
    if mock:
        return f"[MOCK INTEGRATION] {llava_caption[:50]}... + {minicpm_caption[:50]}..."

    llama_prompt = f"""
    Carefully integrate these two captions into a single highly detailed caption. Your integrated caption must:
    - Retain all accurate and relevant details explicitly mentioned in either original caption.
    - Avoid redundancy and contradictions.
    - Strictly avoid any speculation, assumptions, or introduction of details not explicitly stated.

    Original Captions:
    1. LLaVA: {llava_caption}
    2. MiniCPM: {minicpm_caption}

    Provide your response strictly in this format:
    Final Caption: <Integrated caption>
    """
    try:
        with timeout(600):
            response_stream = generate_response(llama_model, llama_prompt)
            response_text = ''.join(stream_parser(response_stream)).strip()
    except TimeoutException:
        print("⚠️ Timeout during caption integration, retrying...")
        return generate_integrated_caption(llava_caption, minicpm_caption)

    match = re.search(r'Final Caption:\s*(.+)', response_text, re.DOTALL)
    return match.group(1).strip() if match else llava_caption


def integrate_captions_for_video(video_id, llava_json_path, minicpm_json_path, out_dir, video_frames_dir, mock=False):
    """
    Integrate captions for a single video and save to DOCX and JSON.
    """
    video_out_dir = os.path.join(out_dir, video_id)
    os.makedirs(video_out_dir, exist_ok=True)

    out_docx_path = os.path.join(video_out_dir, f"{video_id}_captions_integrated.docx")
    out_json_path = os.path.join(video_out_dir, f"{video_id}_captions_integrated.json")

    # Load existing progress if any
    existing_captions = load_captions_from_json(out_json_path)

    llava_captions = load_captions_from_json(llava_json_path)
    minicpm_captions = load_captions_from_json(minicpm_json_path)

    if not llava_captions:
        print(f"⚠️ No LLaVA captions found for {video_id}")
        return 0

    if not minicpm_captions:
        print(f"⚠️ No MiniCPM captions found for {video_id}")
        return 0

    # Prepare DOCX
    if os.path.exists(out_docx_path):
        try:
            doc = Document(out_docx_path)
        except:
            doc = Document()
            doc.add_heading(f"Integrated Captions: {video_id}", 0)
    else:
        doc = Document()
        doc.add_heading(f"Integrated Captions: {video_id}", 0)

    # Prepare JSON data
    json_payload = {
        "video_id": video_id,
        "frames": []
    }

    if os.path.exists(out_json_path):
        try:
            with open(out_json_path, 'r', encoding='utf-8') as f:
                existing_data = json.load(f)
                json_payload["frames"] = existing_data.get("frames", [])
        except:
            pass

    processed_count = 0

    # Get all unique frame IDs from both sources
    all_frame_ids = sorted(set(llava_captions.keys()) | set(minicpm_captions.keys()))

    for frame_id in all_frame_ids:
        # Check if already processed
        if frame_id in existing_captions:
            continue

        llava_text = llava_captions.get(frame_id, "")
        minicpm_text = minicpm_captions.get(frame_id, "")

        if not llava_text or not minicpm_text:
            final_caption = llava_text if llava_text else minicpm_text
            if not final_caption:
                continue
        else:
            print(f"🔄 Integrating frame: {frame_id} for video: {video_id}")
            final_caption = generate_integrated_caption(llava_text, minicpm_text, mock=mock)

        # Add to DOCX
        doc.add_heading(f"Frame: {frame_id}", level=1)

        # Insert Image if available
        if video_frames_dir:
            image_inserted = False

            # Normalize frame_id (remove extension if present)
            base_frame_id = os.path.splitext(frame_id)[0]

            # Potential filenames to try
            candidates = [
                f"{base_frame_id}.jpg", f"{base_frame_id}.jpeg", f"{base_frame_id}.png",
                f"{base_frame_id}.JPG", f"{base_frame_id}.JPEG", f"{base_frame_id}.PNG",
                frame_id  # In case it already has extension and matches exactly
            ]

            for candidate in candidates:
                image_path = os.path.join(video_frames_dir, candidate)
                if os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(5.0))
                        image_inserted = True
                        # print(f"   🖼️ Inserted image: {candidate}") # Optional: too verbose?
                        break
                    except Exception as e:
                        print(f"   ⚠️ Failed to insert image {image_path}: {e}")

            if not image_inserted:
                # Fallback: Case-insensitive search in directory
                try:
                    files_in_dir = os.listdir(video_frames_dir)
                    for f in files_in_dir:
                        if os.path.splitext(f)[0].lower() == base_frame_id.lower() and f.lower().endswith(
                                ('.jpg', '.jpeg', '.png')):
                            image_path = os.path.join(video_frames_dir, f)
                            try:
                                doc.add_picture(image_path, width=Inches(5.0))
                                image_inserted = True
                                break
                            except Exception as e:
                                print(f"   ⚠️ Failed to insert image (fallback) {image_path}: {e}")
                except Exception as e:
                    print(f"   ⚠️ Error listing frames dir {video_frames_dir}: {e}")

            if not image_inserted:
                print(f"   ❌ Image not found for {frame_id} in {video_frames_dir}")
                doc.add_paragraph(f"[Image not found: {frame_id}]")

        doc.add_paragraph(final_caption)

        # Add to JSON
        json_payload["frames"].append({
            "frame_id": frame_id,
            "caption": final_caption
        })

        # Save incrementally
        try:
            doc.save(out_docx_path)
            with open(out_json_path, "w", encoding="utf-8") as f:
                json.dump(json_payload, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"⚠️ Failed to save after {frame_id}: {e}")

        processed_count += 1

    if processed_count > 0:
        print(f"✅ Processed {processed_count} frames for video {video_id}")
    else:
        print(f"⏭️ No new frames to process for video {video_id}")

    return processed_count


def integrate_all_folders(llava_dir, minicpm_dir, out_dir, frames_root, mock=False):
    llava_dir = os.path.abspath(llava_dir)
    minicpm_dir = os.path.abspath(minicpm_dir)
    out_dir = os.path.abspath(out_dir)
    frames_root = os.path.abspath(frames_root)

    print(f"📂 LLaVA Input: {llava_dir}")
    print(f"📂 MiniCPM Input: {minicpm_dir}")
    print(f"📂 Frames Root: {frames_root}")
    print(f"📂 Output: {out_dir}")
    if mock:
        print("⚠️ RUNNING IN MOCK MODE - No LLM calls will be made.")

    total_videos = 0

    for root, dirs, files in os.walk(llava_dir):
        for file in files:
            if file.endswith("_captions_llava.json"):
                llava_json_path = os.path.join(root, file)

                # Infer video_id from filename
                video_id = file.replace("_captions_llava.json", "")

                # Try to find corresponding MiniCPM file
                minicpm_json_path = None
                found_minicpm = False
                for mc_root, mc_dirs, mc_files in os.walk(minicpm_dir):
                    if f"{video_id}_captions_MiniCPM.json" in mc_files:
                        minicpm_json_path = os.path.join(mc_root, f"{video_id}_captions_MiniCPM.json")
                        found_minicpm = True
                        break

                if not found_minicpm:
                    print(f"⚠️ Could not find MiniCPM captions for video: {video_id}")
                    continue

                # Determine output directory structure and frames directory
                # root is .../frames_conventional_captions_llava/<video_id>
                parent_dir = os.path.dirname(root)
                dataset_folder_name = os.path.basename(parent_dir)  # e.g. frames_conventional_captions_llava

                dataset_name = dataset_folder_name.replace("_captions_llava", "_captions_integrated")

                # Infer frames directory
                # If dataset_folder_name is "frames_embodied_captions_llava", frames should be in "frames_embodied"
                frames_folder_name = dataset_folder_name.replace("_captions_llava", "")
                video_frames_dir = os.path.join(frames_root, frames_folder_name, video_id)

                # Check for nested structure (e.g. raw_frames)
                if os.path.exists(video_frames_dir):
                    for subdir in ["raw_frames", "frames", "images"]:
                        if os.path.exists(os.path.join(video_frames_dir, subdir)):
                            video_frames_dir = os.path.join(video_frames_dir, subdir)
                            break

                if not os.path.exists(video_frames_dir):
                    print(f"⚠️ Frames directory not found: {video_frames_dir}")
                    # Fallback: try searching for video_id in frames_root
                    found_frames = False
                    for fr_root, fr_dirs, fr_files in os.walk(frames_root):
                        if os.path.basename(fr_root) == video_id:
                            # Found the video folder, now check inside for subdirs
                            potential_dir = fr_root
                            for subdir in ["raw_frames", "frames", "images"]:
                                if os.path.exists(os.path.join(potential_dir, subdir)):
                                    potential_dir = os.path.join(potential_dir, subdir)
                                    break

                            video_frames_dir = potential_dir
                            found_frames = True
                            break
                    if not found_frames:
                        print(f"   ❌ Could not locate frames for {video_id}")
                        video_frames_dir = None

                # Construct final output dir
                final_out_dir = os.path.join(out_dir, dataset_name)

                if video_frames_dir:
                    print(f"   📂 Resolved frames dir: {video_frames_dir}")

                integrate_captions_for_video(video_id, llava_json_path, minicpm_json_path, final_out_dir,
                                             video_frames_dir, mock=mock)
                total_videos += 1

    print(f"\n✅ All done. Processed {total_videos} videos.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate integrated captions from LLaVA and MiniCPM JSON outputs.")
    parser.add_argument("--llava_dir", default="../../data/llava_caption", help="Root directory for LLaVA captions")
    parser.add_argument("--minicpm_dir", default="../../data/minicpm_caption",
                        help="Root directory for MiniCPM captions")
    parser.add_argument("--out_dir", default="../../data/integrated_caption",
                        help="Output directory for integrated captions")
    parser.add_argument("--frames_dir", default="../../../VideoAnalysis/data/frames",
                        help="Root directory containing frame subfolders (e.g. frames_embodied)")
    parser.add_argument("--mock", action="store_true", help="Run in mock mode (no LLM calls)")
    args = parser.parse_args()

    integrate_all_folders(args.llava_dir, args.minicpm_dir, args.out_dir, args.frames_dir, mock=args.mock)
