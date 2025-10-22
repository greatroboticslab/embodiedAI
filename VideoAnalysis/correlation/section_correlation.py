import os
import re
import json
import argparse
from urllib.parse import urlparse, parse_qs
import sys
from collections import defaultdict

# local imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from config import Config
from utils.llm_utils import generate_response, stream_parser
from docx import Document
from docx.shared import Inches

YES = {"yes", "y", "true", "t", "1"}

# -------------------- URL parsing --------------------

def get_video_id_from_url(text: str):
    text = text.strip()
    if not text or "youtu" not in text:
        return None
    m = re.search(r"https?://[^\s)\]>]+", text)
    if not m:
        return None
    url = re.sub(r'[:.,;!?")\]]+$', '', m.group(0))
    try:
        u = urlparse(url)
        host = (u.netloc or "").lower()
        if "youtu.be" in host:
            vid = u.path.strip("/").split("/")[0]
            return clean_video_id(vid)
        qs = parse_qs(u.query or "")
        if "v" in qs and qs["v"]:
            return clean_video_id(qs["v"][0])
        parts = [p for p in u.path.split("/") if p]
        if "embed" in parts:
            idx = parts.index("embed")
            if idx + 1 < len(parts):
                return clean_video_id(parts[idx + 1])
    except Exception:
        pass
    return None

def clean_video_id(vid: str):
    m = re.match(r"([A-Za-z0-9_-]{6,})", vid)
    return m.group(1) if m else vid

# -------------------- DOCX parsers --------------------

def parse_integrated_docx(docx_path: str):
    """
    Expect sections like:
      Heading/paragraph 'Frame: frame_0001.jpg'
      then lines including 'Caption:' etc. until next 'Frame:'.
    Returns list[{'frame': 'frame_0001.jpg', 'caption': '...'}] in order.
    """
    doc = Document(docx_path)
    frames = []
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i].text.strip()
        if p.lower().startswith("frame:"):
            frame_name = p.split(":", 1)[1].strip()
            i += 1
            chunks = []
            while i < len(doc.paragraphs):
                t = doc.paragraphs[i].text.strip()
                if t.lower().startswith("frame:"):
                    i -= 1
                    break
                if t:
                    # prefer lines after "Caption:" but keep all text blocks
                    chunks.append(t)
                i += 1
            caption = " ".join(chunks).strip()
            frames.append({"frame": frame_name, "caption": caption})
        i += 1
    return frames

def parse_comments_docx(comments_docx_path: str):
    """
    Returns dict: {video_id: [{"url": url, "comment": text}, ...]}
    Layout: URL line, followed by one or more comment paragraphs, until next URL.
    """
    doc = Document(comments_docx_path)
    by_video = defaultdict(list)
    cur_vid, cur_url, buffer = None, None, []

    def flush():
        nonlocal buffer
        if cur_vid and buffer:
            text = " ".join(buffer).strip()
            if text:
                by_video[cur_vid].append({"url": cur_url, "comment": text})
        buffer = []

    for para in doc.paragraphs:
        t = para.text.strip()
        vid = get_video_id_from_url(t)
        if vid:
            flush()
            cur_vid, cur_url = vid, t
        elif t:
            buffer.append(t)
    flush()
    return dict(by_video)

# -------------------- Embeddings + segmentation --------------------

def _embed_texts(texts, model_name="all-mpnet-base-v2"):
    """
    Returns L2-normalized embeddings. If sentence-transformers is missing, raises ImportError.
    """
    try:
        from sentence_transformers import SentenceTransformer
        import numpy as np
    except Exception as e:
        raise ImportError("sentence-transformers not installed") from e

    model = SentenceTransformer(model_name)
    X = model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
    return X

def segment_by_caption_embeddings(frames, n_sections=None, pen=None, min_size=6, embed_model="all-mpnet-base-v2"):
    """
    Use change-point detection over caption embeddings to split into contiguous sections.
    Returns list of dicts:
      {
        "section_id": "sXXX_startIdx_endIdx",
        "start": start_index,   # inclusive frame index
        "end": end_index,       # exclusive frame index
        "frames": [frame dicts...],  # original frame dicts
        "text": "... concatenated captions ..."
      }
    """
    try:
        import ruptures as rpt
        import numpy as np
    except Exception as e:
        raise ImportError("ruptures not installed") from e

    texts = [(f.get("caption") or "").strip() for f in frames]
    if not any(t for t in texts):
        # fallback: one big section (no captions)
        return [{
            "section_id": f"s000_0_{len(frames)}",
            "start": 0, "end": len(frames),
            "frames": frames[:],
            "text": ""
        }]

    X = _embed_texts(texts, model_name=embed_model)  # (N, d) normalized
    N = len(frames)

    if N < max(2*min_size, 8):
        return [{
            "section_id": f"s000_0_{N}",
            "start": 0, "end": N,
            "frames": frames[:],
            "text": " ".join(texts)
        }]

    # Prefer PELT with RBF; choose either target count or penalty
    algo = rpt.Pelt(model="rbf", min_size=min_size).fit(X)
    if n_sections:
        # Binseg is deterministic for a fixed # of breakpoints; n_sections -> n_bkps = n_sections-1
        bkps = rpt.Binseg(model="rbf").fit(X).predict(n_bkps=max(0, n_sections-1))
    else:
        # penalty controls how many segments; tune 10..40 depending on density
        bkps = algo.predict(pen=pen or 15)

    if bkps[-1] != N:
        bkps[-1] = N

    sections = []
    s = 0
    sec_idx = 0
    for b in bkps:
        e = b
        if e - s <= 0:
            s = e
            continue
        chunk = frames[s:e]
        section_text = " ".join((t or "") for t in texts[s:e]).strip()
        sections.append({
            "section_id": f"s{sec_idx:03d}_{s}_{e}",
            "start": s, "end": e,
            "frames": chunk,
            "text": section_text
        })
        sec_idx += 1
        s = e
    return sections

def summarize_section_text(section_text, max_chars=600):
    """
    Simple, fast summarizer: keep the first ~max_chars while trimming.
    (You can swap in an LLM summarizer later.)
    """
    if not section_text:
        return ""
    s = re.sub(r"\s+", " ", section_text).strip()
    if len(s) > max_chars:
        s = s[:max_chars].rsplit(" ", 1)[0] + " ..."
    return s

def pick_representative_frames(frame_names, k=3):
    """
    Evenly spaced picks from the list, preserving order.
    """
    if not frame_names:
        return []
    if k >= len(frame_names):
        return frame_names
    idxs = [int(round(i * (len(frame_names)-1) / (k-1))) for i in range(k)]
    seen = set()
    picks = []
    for idx in idxs:
        if idx not in seen:
            picks.append(frame_names[idx])
            seen.add(idx)
    return picks

# -------------------- LLM scoring --------------------

def ask_model(model_name: str, section_summary: str, comment: str):
    """
    Return {"correlated": bool, "score": int, "reason": str}
    """
    prompt = (
        "You are checking whether a VIDEO SECTION SUMMARY and a VIEWER COMMENT are strongly related.\n"
        "Strong = the comment directly discusses the same steps, objects, settings, parameters, or goals described in the section (not generic praise).\n\n"
        "Give a correlation confidence 0-100, and a brief reason citing specific overlaps.\n\n"
        f"SECTION SUMMARY:\n{section_summary}\n\nCOMMENT:\n{comment}\n\n"
        "Respond ONLY as compact JSON: "
        '{"correlated": true|false, "score": 0-100, "reason": "brief explanation"}'
    )
    stream = generate_response(model_name, prompt)
    text = "".join(stream_parser(stream)).strip()
    m = re.search(r"\{.*?\}", text, flags=re.S)
    jtxt = m.group(0) if m else text
    try:
        obj = json.loads(jtxt)
        corr = str(obj.get("correlated", False)).lower() in YES
        score = int(obj.get("score", 0))
        score = max(0, min(100, score))
        reason = str(obj.get("reason", "")).strip()
        return {"correlated": corr, "score": score, "reason": reason}
    except Exception:
        corr = "yes" in text.lower()
        m2 = re.search(r"(\d{1,3})", text)
        score = int(m2.group(1)) if m2 else 0
        score = max(0, min(100, score))
        return {"correlated": corr, "score": score, "reason": text[:400]}

# -------------------- IO helpers (resume-safe) --------------------

def results_path(raw_frames_dir: str, video_id: str,
                 seg_tag: str, embed_model: str, retrieval_topk: int):
    """
    Separate from your single-frame & group files.
    """
    safe_model = re.sub(r"[^A-Za-z0-9_-]+", "", embed_model or "mpnet")
    return os.path.join(
        raw_frames_dir,
        f"{video_id}_sectioncorr_{seg_tag}_{safe_model}_k{retrieval_topk}.results.jsonl"
    )

def load_existing_results(results_file: str):
    done = {}
    if os.path.isfile(results_file):
        with open(results_file, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    rec = json.loads(line)
                    done[rec["section_id"]] = rec
                except Exception:
                    continue
    return done

def append_result(results_file: str, record: dict):
    with open(results_file, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")

# -------------------- DOCX helpers --------------------

def _add_thumbnails_row(doc, base_dir, frame_list, thumb_width_in=1.3):
    tbl = doc.add_table(rows=1, cols=max(1, len(frame_list)))
    tbl.autofit = True
    for col, relpath in enumerate(frame_list):
        cell = tbl.rows[0].cells[col]
        img_path = os.path.join(base_dir, relpath)
        if os.path.exists(img_path):
            try:
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=Inches(thumb_width_in))
            except Exception:
                cell.text = relpath
        else:
            cell.text = relpath
    return tbl

def build_docx_from_results(out_docx: str, sections: list, results_map: dict,
                            min_score: int, top_k: int, model_name: str,
                            total_comments_available: int | None = None):
    """
    sections: list of dicts with keys: section_id, frames (list of {'frame','caption'}), summary, rep_frames
    results_map: {section_id: {"candidates":[{...}], "checked_pairs":int, ...}}
    """
    doc = Document()
    doc.add_heading("Comment–Section Correlation Report", 0)
    doc.add_paragraph(f"Model: {model_name}")

    # Aggregate
    total_pairs = sum(rec.get("checked_pairs", 0) for rec in results_map.values())
    base_dir = os.path.dirname(out_docx)

    # Per-comment index (for Top/Least sections)
    comments_index = {}
    for sec in sections:
        sid = sec["section_id"]
        rec = results_map.get(sid, {})
        for c in rec.get("candidates", []):
            key = (c.get("url",""), c.get("comment",""))
            entry = comments_index.setdefault(key, {
                "comment": c.get("comment",""),
                "url": c.get("url",""),
                "sections": set(),
                "max_score": 0,
                "per_section": {}
            })
            s = float(c.get("score", 0))
            r = str(c.get("reason", "") or "")
            entry["per_section"][sid] = {"score": s, "reason": r, "rep_frames": sec.get("rep_frames", [])}
            if c.get("correlated") and int(s) >= min_score:
                entry["sections"].add(sid)
                entry["max_score"] = max(entry["max_score"], int(s))

    unique_corr_count = sum(1 for v in comments_index.values() if v["sections"])
    pct = None
    if total_comments_available and total_comments_available > 0:
        pct = 100.0 * unique_corr_count / total_comments_available

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Sections processed: {len(sections)}")
    if total_comments_available is not None:
        doc.add_paragraph(f"Comments available for this video: {total_comments_available}")
    doc.add_paragraph(
        f"Comments correlated to ≥ {min_score}: {unique_corr_count}"
        + (f" ({pct:.1f}%)" if pct is not None else "")
    )
    doc.add_paragraph(f"Total candidate pairs checked: {total_pairs}")

    # Per-comment coverage (Top & Least sections with thumbnails + reasons)
    if comments_index:
        doc.add_paragraph("Correlated comment coverage (all comments):")
        TOPK = 10
        all_list = sorted(
            comments_index.values(),
            key=lambda x: (len(x.get("sections", [])), x.get("max_score", float("-inf")), x.get("comment","")),
            reverse=True
        )
        for e in all_list:
            coverage = len(e.get("sections", []))
            max_score = e.get("max_score")
            header = f"- matched {coverage} section(s)"
            if max_score is not None:
                header += f", max score {max_score:.4f}"
            header += f": {e.get('comment','')}"
            doc.add_paragraph(header)
            if e.get("url"):
                doc.add_paragraph(f"  Source: {e['url']}")

            items = list(e["per_section"].items())  # (sid, info)
            # Top-10
            top_secs = sorted(items, key=lambda kv: (kv[1]["score"], kv[0]), reverse=True)[:TOPK]
            doc.add_paragraph("  Top correlated sections:")
            if top_secs:
                for sid, info in top_secs:
                    _add_thumbnails_row(doc, base_dir, info.get("rep_frames", []), thumb_width_in=1.2)
                    p = doc.add_paragraph()
                    p.add_run(f"score {info['score']:.4f}")
                    if info.get("reason"):
                        p.add_run(f" — reason: {info['reason']}")
            else:
                doc.add_paragraph("    • (none)")

            # Bottom-10
            bot_secs = sorted(items, key=lambda kv: (kv[1]["score"], kv[0]))[:TOPK]
            doc.add_paragraph("  Least correlated sections:")
            if bot_secs:
                for sid, info in bot_secs:
                    _add_thumbnails_row(doc, base_dir, info.get("rep_frames", []), thumb_width_in=1.2)
                    p = doc.add_paragraph()
                    p.add_run(f"score {info['score']:.4f}")
                    if info.get("reason"):
                        p.add_run(f" — reason: {info['reason']}")
            else:
                doc.add_paragraph("    • (none)")

    # Per-section details
    for sec in sections:
        sid = sec["section_id"]
        frames = sec.get("frames", [])
        rep = sec.get("rep_frames", [])
        summary = sec.get("summary", "")

        doc.add_heading(f"Section {sid}  (frames {sec['start']}–{sec['end']})", level=1)
        if rep:
            _add_thumbnails_row(doc, base_dir, rep, thumb_width_in=1.4)
        else:
            # fallback: show first 3 frames if rep missing
            _add_thumbnails_row(doc, base_dir, [f["frame"] for f in frames[:3]], thumb_width_in=1.4)

        if summary:
            doc.add_paragraph(f"Summary: {summary}")

        rec = results_map.get(sid, {})
        candidates = rec.get("candidates", [])

        kept = [c for c in candidates if c.get("score", 0) >= min_score and c.get("correlated", False)]
        kept.sort(key=lambda x: x.get("score", 0), reverse=True)
        kept = kept[:top_k]

        if kept:
            for r in kept:
                doc.add_paragraph(f"Comment (score {r['score']}): {r['comment']}")
                if r.get("reason"): doc.add_paragraph(f"Reason: {r['reason']}")
                if r.get("url"): doc.add_paragraph(f"Source: {r['url']}")
        else:
            doc.add_paragraph("No strongly correlated comments found for this section.")

    doc.save(out_docx)

# -------------------- Core processing --------------------

def discover_targets(frames_input: str):
    frames_input = os.path.abspath(frames_input)
    targets = []
    if os.path.basename(frames_input) == "raw_frames":
        targets.append(frames_input)
        return targets
    for entry in sorted(os.listdir(frames_input)):
        d = os.path.join(frames_input, entry, "raw_frames")
        if os.path.isdir(d):
            targets.append(d)
    return targets

def _find_integrated_docx(raw_frames_dir: str):
    for f in os.listdir(raw_frames_dir):
        if f.endswith("_raw_frames_captions_integrated.docx") or f.endswith("_raw_frames_integrated.docx"):
            return os.path.join(raw_frames_dir, f)
    return None

def process_one_video(raw_frames_dir: str, comments_by_video: dict, model: str,
                      min_score: int, top_k_scan: int,
                      seg_n_sections: int | None, seg_penalty: int | None,
                      seg_min_size: int, embed_model: str,
                      retrieval_topk: int, rep_k: int):
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))
    integrated_docx = _find_integrated_docx(raw_frames_dir)
    if not integrated_docx:
        print(f"[SKIP] {video_id}: integrated docx not found in {raw_frames_dir}")
        return

    frames = parse_integrated_docx(integrated_docx)
    if not frames:
        print(f"[SKIP] {video_id}: no frames parsed.")
        return

    # Segment into sections
    seg_tag = f"n{seg_n_sections or 0}_p{seg_penalty or 0}_m{seg_min_size}"
    try:
        sections = segment_by_caption_embeddings(
            frames, n_sections=seg_n_sections, pen=seg_penalty,
            min_size=seg_min_size, embed_model=embed_model
        )
    except ImportError as e:
        print(f"[FATAL] {video_id}: {e}. Install the missing package(s).")
        return

    # Summaries + representative frames
    for sec in sections:
        fnames = [f["frame"] for f in sec["frames"]]
        sec["rep_frames"] = pick_representative_frames(fnames, k=rep_k)
        sec["summary"] = summarize_section_text(sec.get("text",""), max_chars=600)

    comments_for_video = comments_by_video.get(video_id, [])
    if not comments_for_video:
        comments_for_video = [c for v in comments_by_video.values() for c in v]
        print(f"[WARN] {video_id}: no comments matched; falling back to ALL comments ({len(comments_for_video)}).")
    else:
        print(f"[INFO] {video_id}: {len(comments_for_video)} comments found for this video.")

    # Pre-embed comments for retrieval (if possible)
    comment_texts = [c["comment"] for c in comments_for_video]
    C = None
    retrieval_ok = False
    try:
        C = _embed_texts(comment_texts, model_name=embed_model)  # normalized
        retrieval_ok = True
    except ImportError:
        print(f"[WARN] {video_id}: sentence-transformers missing; skipping fast retrieval (LLM will score all comments).")

    # resume state
    res_file = results_path(raw_frames_dir, video_id, seg_tag, embed_model, retrieval_topk)
    existing = load_existing_results(res_file)
    already = set(existing.keys())
    print(f"[RESUME] {video_id}: {len(already)}/{len(sections)} sections already processed.")

    # process sections
    import numpy as np
    for sec in sections:
        sid = sec["section_id"]
        if sid in already:
            continue

        summary = sec.get("summary","")
        checked_pairs = 0
        candidates = []

        # choose candidate comments
        candidate_indices = range(len(comment_texts))
        if retrieval_ok and retrieval_topk > 0 and summary:
            s_vec = _embed_texts([summary], model_name=embed_model)[0]  # normalized
            sims = np.dot(C, s_vec)  # cosine because both normalized
            idx = np.argsort(-sims)[:min(retrieval_topk, len(comment_texts))]
            candidate_indices = idx

        for i in candidate_indices:
            item = comments_for_video[i]
            res = ask_model(model, summary, item["comment"])
            checked_pairs += 1
            candidates.append({
                "correlated": bool(res["correlated"]),
                "score": int(res["score"]),
                "reason": res.get("reason",""),
                "comment": item["comment"],
                "url": item["url"]
            })

        record = {
            "video_id": video_id,
            "section_id": sid,
            "start": sec["start"],
            "end": sec["end"],
            "rep_frames": sec.get("rep_frames", []),
            "checked_pairs": checked_pairs,
            "candidates": candidates
        }
        append_result(res_file, record)
        existing[sid] = record
        print(f"[OK] {video_id}: section {sid} processed ({checked_pairs} pairs).")

    # build docx
    out_docx = os.path.join(raw_frames_dir, f"{video_id}_correlation_sections_{seg_tag}.docx")
    # enrich sections from saved map (rep_frames already included)
    build_docx_from_results(
        out_docx, sections, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[DONE] {video_id}: report saved -> {out_docx}")

def rebuild_docx_only(raw_frames_dir: str, comments_by_video: dict, model: str,
                      min_score: int, top_k_scan: int,
                      seg_n_sections: int | None, seg_penalty: int | None,
                      seg_min_size: int, embed_model: str,
                      retrieval_topk: int, rep_k: int):
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))
    integrated_docx = _find_integrated_docx(raw_frames_dir)
    if not integrated_docx:
        print(f"[SKIP] {video_id}: integrated docx not found for rebuild.")
        return
    frames = parse_integrated_docx(integrated_docx)
    if not frames:
        print(f"[SKIP] {video_id}: no frames parsed for rebuild.")
        return

    seg_tag = f"n{seg_n_sections or 0}_p{seg_penalty or 0}_m{seg_min_size}"
    try:
        sections = segment_by_caption_embeddings(
            frames, n_sections=seg_n_sections, pen=seg_penalty,
            min_size=seg_min_size, embed_model=embed_model
        )
    except ImportError as e:
        print(f"[FATAL] {video_id}: {e}. Install the missing package(s).")
        return
    for sec in sections:
        fnames = [f["frame"] for f in sec["frames"]]
        sec["rep_frames"] = pick_representative_frames(fnames, k=rep_k)
        sec["summary"] = summarize_section_text(sec.get("text",""), max_chars=600)

    res_file = results_path(raw_frames_dir, video_id, seg_tag, embed_model, retrieval_topk)
    if not os.path.isfile(res_file):
        print(f"[SKIP] {video_id}: no results file found to rebuild.")
        return
    existing = load_existing_results(res_file)

    comments_for_video = comments_by_video.get(video_id, [])
    out_docx = os.path.join(raw_frames_dir, f"{video_id}_correlation_sections_{seg_tag}.docx")

    build_docx_from_results(
        out_docx, sections, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[REBUILT] {video_id}: report saved -> {out_docx}")

# -------------------- CLI --------------------

def main():
    ap = argparse.ArgumentParser(description="Segment video into sections from captions, then correlate comments to sections (resume-safe).")
    ap.add_argument("frames_path", help="Path to frames root (containing <video_id>/raw_frames) OR a single raw_frames folder.")
    ap.add_argument("comments_docx", help="DOCX with video URLs and comments.")
    ap.add_argument("--model", default=None, help="LLM name for scoring (defaults to Config.ollama_models[0]).")
    ap.add_argument("--min_score", type=int, default=60, help="Minimum score to include a matched comment in the final report.")
    ap.add_argument("--top_k", type=int, default=5, help="Top-K comments per section shown in the final report.")

    # Segmentation controls
    ap.add_argument("--seg_n_sections", type=int, default=None, help="Target number of sections (overrides penalty).")
    ap.add_argument("--seg_penalty", type=int, default=15, help="Penalty for PELT (higher = fewer sections) if seg_n_sections not set.")
    ap.add_argument("--seg_min_size", type=int, default=6, help="Minimum frames per section to avoid over-segmentation.")
    ap.add_argument("--embed_model", default="all-mpnet-base-v2", help="SentenceTransformer model for embeddings.")

    # Retrieval + visuals
    ap.add_argument("--retrieval_topk", type=int, default=50, help="How many candidate comments to retrieve per section before LLM re-score (0 = score all).")
    ap.add_argument("--rep_k", type=int, default=3, help="Representative thumbnails per section in the report.")

    ap.add_argument("--rebuild", action="store_true", help="Rebuild DOCX from existing results without running the model.")

    args = ap.parse_args()
    model = args.model or Config.ollama_models[0]

    comments_by_video = parse_comments_docx(args.comments_docx)
    if not comments_by_video:
        print("[FATAL] No comments parsed from comments DOCX.")
        return

    targets = discover_targets(args.frames_path)
    if not targets:
        print("[FATAL] No raw_frames folders found.")
        return

    print(f"[INFO] Found {len(targets)} video(s) to process.")

    for raw_frames in targets:
        print(f"\n=== Processing: {os.path.basename(os.path.dirname(raw_frames))} ===")
        try:
            if args.rebuild:
                rebuild_docx_only(
                    raw_frames, comments_by_video, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    seg_n_sections=args.seg_n_sections, seg_penalty=args.seg_penalty,
                    seg_min_size=args.seg_min_size, embed_model=args.embed_model,
                    retrieval_topk=args.retrieval_topk, rep_k=args.rep_k
                )
            else:
                process_one_video(
                    raw_frames, comments_by_video, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    seg_n_sections=args.seg_n_sections, seg_penalty=args.seg_penalty,
                    seg_min_size=args.seg_min_size, embed_model=args.embed_model,
                    retrieval_topk=args.retrieval_topk, rep_k=args.rep_k
                )
        except KeyboardInterrupt:
            print("\n[INTERRUPTED] Stopping cleanly. You can rerun to resume.")
            break
        except Exception as e:
            print(f"[ERROR] {raw_frames}: {e}")

if __name__ == "__main__":
    main()
