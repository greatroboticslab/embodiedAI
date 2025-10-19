import os
import torch
from PIL import Image
from transformers import AutoModel, AutoTokenizer
from docx import Document
import traceback
import argparse


def generate_caption(image_path, model, tokenizer, device):
    try:
        image = Image.open(image_path).convert('RGB')
    except Exception as e:
        return f"[Error loading image: {e}]"

    question = 'our caption is too short, please show detailed and overview for the image.'
    msgs = [{'role': 'user', 'content': question}]
    context = "This is an image captioning task."

    try:
        res = model.chat(
            image=image,
            context=context,
            msgs=msgs,
            tokenizer=tokenizer,
            sampling=True,
            temperature=0.7
        )
        return res[0]
    except Exception as e:
        traceback.print_exc()
        return f"[Error generating caption: {e}]"


def caption_folder(folder_path, model, tokenizer, device, save_path=None):
    image_files = sorted([
        f for f in os.listdir(folder_path)
        if f.lower().endswith(('.jpg', '.jpeg', '.png'))
    ])
    if not image_files:
        return None, 0

    doc = Document()
    folder_name = os.path.basename(folder_path.rstrip('/'))
    doc.add_heading(f"Captions for {folder_name}", level=0)

    count = 0
    for filename in image_files:
        image_path = os.path.join(folder_path, filename)
        print(f"  Processing {filename}")
        caption = generate_caption(image_path, model, tokenizer, device)
        doc.add_heading(filename, level=1)
        doc.add_paragraph(caption)
        count += 1

    if save_path:
        doc.save(save_path)
        print(f"  Saved individual .docx: {save_path}")

    return doc, count


def caption_all_folders(root_folder, combined_filename):
    print("Loading model and tokenizer...")
    model_id = 'openbmb/MiniCPM-V'
    model = AutoModel.from_pretrained(model_id, trust_remote_code=True, torch_dtype=torch.float32)
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    model = model.to(device)
    model.eval()
    tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)
    print("Model and tokenizer loaded.")

    combined_doc = Document()
    combined_doc.add_heading('Combined Captions from All Folders', level=0)
    total_images = 0

    for subdir, _, files in os.walk(root_folder):
        image_files = [f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
        if not image_files:
            continue  # Skip folders without images

        rel_path = os.path.relpath(subdir, root_folder)
        output_docx_name = f"{rel_path.replace(os.sep, '_')}_captions_MiniCPM.docx"
        output_docx_path = os.path.join(subdir, output_docx_name)

        print(f"\nProcessing folder: {rel_path}")
        sub_doc, count = caption_folder(subdir, model, tokenizer, device, save_path=output_docx_path)
        total_images += count

        if sub_doc:
            combined_doc.add_page_break()
            combined_doc.add_heading(f"Folder: {rel_path}", level=1)
            for element in sub_doc.element.body:
                combined_doc.element.body.append(element)

    combined_path = os.path.join(root_folder, f"{combined_filename}.docx")
    if total_images > 0:
        combined_doc.save(combined_path)
        print(f"\nCombined .docx saved at: {combined_path}")
    else:
        print("\nNo images were processed. Combined document not saved.")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Batch caption all folders using MiniCPM-V.")
    parser.add_argument('-i', required=True, help='Root folder containing subfolders of images')
    parser.add_argument('-o', required=True, help='Filename (without extension) for combined .docx output')
    args = parser.parse_args()

    caption_all_folders(args.i, args.o)
