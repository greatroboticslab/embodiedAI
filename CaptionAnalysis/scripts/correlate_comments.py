#!/usr/bin/env python3
"""
Correlate viewer comments to video frames (using integrated JSON captions).

Inputs:
  1) frames_path: Root directory containing video frames (can be nested).
     The script discovers all `raw_frames` subdirectories.
  2) comments_docx: DOCX file with video URLs and comments.
  3) integrated_captions_root: Root directory for JSON captions.
     The script searches for `<video_id>_captions_integrated.json`.

Outputs:
  Results are saved to `results_root`, mirroring the relative structure of `frames_path` if applicable.

Example:
  python correlate_comments.py ../data/integrated_caption/frames_conventional_captions_integrated \
      comments.docx --results_root ../results --integrated_captions_root ../data/integrated_caption
"""

import os
import re
import json
import argparse
from urllib.parse import urlparse, parse_qs
import sys
from collections import defaultdict

# local imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from config import Config
from utils.llm_utils import generate_response, stream_parser
from docx import Document
from docx.shared import Inches


YES = {"yes", "y", "true", "t", "1"}

# -------------------- Parsers --------------------

def get_video_id_from_url(text: str):
    """
    Extract YouTube video ID from a paragraph that may contain a URL plus extra text.
    Handles youtube.com, youtu.be, and trims any trailing punctuation/characters after the ID.
    """
    text = text.strip()
    if not text or "youtu" not in text:
        return None

    # Grab first URL-looking token
    m = re.search(r"https?://[^\s)\]>]+", text)
    if not m:
        return None
    url = m.group(0)

    # Remove common trailing punctuation/suffix that might stick to the URL
    url = re.sub(r'[:.,;!?")\]]+$', '', url)

    try:
        u = urlparse(url)
        host = (u.netloc or "").lower()

        if "youtu.be" in host:
            vid = u.path.strip("/").split("/")[0]
            return clean_video_id(vid)

        qs = parse_qs(u.query or "")
        if "v" in qs and qs["v"]:
            return clean_video_id(qs["v"][0])

        parts = [p for p in u.path.split("/") if p]
        if "embed" in parts:
            idx = parts.index("embed")
            if idx + 1 < len(parts):
                return clean_video_id(parts[idx + 1])
    except Exception:
        pass
    return None

def clean_video_id(vid: str):
    """
    Trim any trailing non-YouTube-ID characters (YouTube IDs are 11 chars of [A-Za-z0-9-_]).
    """
    m = re.match(r"([A-Za-z0-9_-]{6,})", vid)
    return m.group(1) if m else vid


def parse_integrated_json(json_path: str):
    """
    Return list[{'frame': 'frame_0000.jpg', 'caption': '...'}] from JSON.
    Assumes JSON structure: {"frames": [{"frame_id": "frame_0000", "caption": "..."}, ...]}
    The raw frame filename is assumed to be <frame_id>.jpg (or similar).
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    frames = []
    for item in data.get("frames", []):
        fid = item.get("frame_id")
        # Ensure we have a filename-like string for compatibility
        if fid and not fid.lower().endswith(('.jpg', '.png', '.jpeg')):
            fid += ".jpg"
            
        frames.append({
            "frame": fid,
            "caption": item.get("caption", "").strip()
        })
    return frames


def parse_comments_docx(comments_docx_path: str):
    """
    Returns dict: {video_id: [{"url": url, "comment": text}, ...]}
    Layout: URL line, followed by one or more comment paragraphs, until next URL.
    """
    doc = Document(comments_docx_path)
    by_video = defaultdict(list)
    cur_vid, cur_url, buffer = None, None, []

    def flush():
        nonlocal buffer
        if cur_vid and buffer:
            text = " ".join(buffer).strip()
            if text:
                by_video[cur_vid].append({"url": cur_url, "comment": text})
        buffer = []

    for para in doc.paragraphs:
        t = para.text.strip()
        vid = get_video_id_from_url(t)
        if vid:
            flush()
            cur_vid, cur_url = vid, t
        elif t:
            buffer.append(t)
    flush()
    return dict(by_video)

# -------------------- Model --------------------

def ask_model(model_name: str, paragraph: str, comment: str):
    """
    Return {"correlated": bool, "score": int, "reason": str}
    """
    prompt = (
        "You are checking whether a video FRAME CAPTION and a VIEWER COMMENT are strongly related.\n"
        "Strong = the comment directly discusses the same technical steps, objects, settings, "
        "or goals described in the caption (not generic praise or unrelated info).\n\n"
        "Include a correlation confidence score from 0 to 100, where 100 indicates perfect correlation based on the connection between the comment and paragraph.\n\n"
        f"FRAME CAPTION:\n{paragraph}\n\nCOMMENT:\n{comment}\n\n"
        "Respond ONLY as compact JSON: "
        '{"correlated": true|false, "score": 0-100, "reason": "provide explanation"}'
    )
    stream = generate_response(model_name, prompt)
    text = "".join(stream_parser(stream)).strip()
    m = re.search(r"\{.*?\}", text, flags=re.S)
    jtxt = m.group(0) if m else text
    try:
        obj = json.loads(jtxt)
        corr = str(obj.get("correlated", False)).lower() in YES
        score = int(obj.get("score", 0))
        score = max(0, min(100, score))
        reason = str(obj.get("reason", "")).strip()
        return {"correlated": corr, "score": score, "reason": reason}
    except Exception:
        corr = "yes" in text.lower()
        m2 = re.search(r"(\d{1,3})", text)
        score = int(m2.group(1)) if m2 else 0
        score = max(0, min(100, score))
        return {"correlated": corr, "score": score, "reason": text[:400]}

# -------------------- IO helpers (resume-safe) --------------------

def results_path_jsonl(output_dir: str, video_id: str):
    return os.path.join(output_dir, f"{video_id}_correlation.results.jsonl")

def load_existing_results(results_file: str):
    done = {}
    if os.path.isfile(results_file):
        with open(results_file, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    rec = json.loads(line)
                    done[rec["frame"]] = rec
                except Exception:
                    continue
    return done

def append_result(results_file: str, record: dict):
    with open(results_file, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")

def build_docx_from_results(out_docx: str, integrated_frames: list, results_map: dict,
                            min_score: int, top_k: int, model_name: str,
                            raw_frames_dir: str,
                            total_comments_available: int | None = None):
    doc = Document()
    doc.add_heading("Comment–Caption Correlation Report", 0)
    doc.add_paragraph(f"Model: {model_name}")
    #doc.add_paragraph(f"Min score: {min_score}, Top-K per frame: {top_k}")

    # --- Aggregate across ALL comments using scores present in results_map ---
    # Count total pairs
    total_pairs = sum(rec.get("checked_pairs", 0) for rec in results_map.values())

    # Build complete per-comment index
    comments_index = {}  # {(url, comment): {...}}
    for frame_name, rec in results_map.items():
        for c in rec.get("candidates", []):
            key = (c.get("url", ""), c.get("comment", ""))
            entry = comments_index.setdefault(key, {
                "comment": c.get("comment", ""),
                "url": c.get("url", ""),
                "frames": set(),
                "max_score": 0,
                "scores": {}
            })
            s = float(c.get("score", 0))
            entry["scores"][frame_name] = s
            if c.get("correlated") and int(s) >= min_score:
                entry["frames"].add(frame_name)
                entry["max_score"] = max(entry["max_score"], int(s))

    # Final counts for the summary
    unique_corr_count = sum(1 for v in comments_index.values() if v["frames"])
    pct = None
    if total_comments_available and total_comments_available > 0:
        pct = 100.0 * unique_corr_count / total_comments_available

    # --- Summary at the beginning ---
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Frames processed: {len(integrated_frames)}")
    if total_comments_available is not None:
        doc.add_paragraph(f"Comments available for this video: {total_comments_available}")
    doc.add_paragraph(
        f"Comments correlated to ≥ {min_score}: {unique_corr_count}"
        + (f" ({pct:.1f}%)" if pct is not None else "")
    )
    doc.add_paragraph(f"Total candidate pairs checked: {total_pairs}")


    # --- Per-comment coverage: ALL comments, with top/bottom frames by score ---
    if comments_index:
        doc.add_paragraph("Correlated comment coverage (all comments):")

        # Sort by: number of matched frames (desc), then max_score (desc), then comment text for stability
        all_list = sorted(
            comments_index.values(),
            key=lambda x: (len(x.get("frames", [])), x.get("max_score", float("-inf")), x.get("comment","")),
            reverse=True
        )

        TOPK = 10

        for e in all_list:
            coverage = len(e.get("frames", []))
            max_score = e.get("max_score")
            header = f"- matched {coverage} frame(s)"
            if max_score is not None:
                header += f", max score {max_score:.4f}"
            header += f": {e.get('comment', '')}"
            doc.add_paragraph(header)
            if e.get("url"):
                doc.add_paragraph(f"  Source: {e['url']}")

            # We have a score for every (comment, frame) pair, so ranking is straightforward
            items = list(e["scores"].items())  # [(frame_id, score), ...]

            # Top-10 highest scores (most correlated)
            top_correlated = sorted(items, key=lambda kv: (kv[1], kv[0]), reverse=True)[:TOPK]
            # Bottom-10 lowest scores (least correlated)
            bottom_uncorrelated = sorted(items, key=lambda kv: (kv[1], kv[0]))[:TOPK]

            doc.add_paragraph("  Top correlated frames:")
            if top_correlated:
                for fid, s in top_correlated:
                    frame_path = os.path.join(raw_frames_dir, fid)
                    if os.path.exists(frame_path):
                        try:
                            doc.add_picture(frame_path, width=Inches(1.5))  # small thumbnail
                            doc.paragraphs[-1].add_run(f" (score {s:.4f})")  # add score text after image
                        except Exception:
                            doc.add_paragraph(f"{fid} (score {s:.4f})")
                    else:
                        doc.add_paragraph(f"{fid} (score {s:.4f})")
            else:
                doc.add_paragraph("    • (none)")

            doc.add_paragraph("  Least correlated frames:")
            if bottom_uncorrelated:
                for fid, s in bottom_uncorrelated:
                    frame_path = os.path.join(raw_frames_dir, fid)
                    if os.path.exists(frame_path):
                        try:
                            doc.add_picture(frame_path, width=Inches(1.5))  # small thumbnail
                            doc.paragraphs[-1].add_run(f" (score {s:.4f})")  # add score text after image
                        except Exception:
                            doc.add_paragraph(f"{fid} (score {s:.4f})")
                    else:
                        doc.add_paragraph(f"{fid} (score {s:.4f})")
            else:
                doc.add_paragraph("    • (none)")


    # --- Per-frame details (as before) ---
    total_kept = 0
    for f in integrated_frames:
        frame = f["frame"]; caption = f["caption"]
        doc.add_heading(f"Frame: {frame}", level=1)
        if caption: doc.add_paragraph(f"Caption: {caption}")

        rec = results_map.get(frame, {})
        candidates = rec.get("candidates", [])

        kept = [c for c in candidates if c.get("score", 0) >= min_score and c.get("correlated", False)]
        kept.sort(key=lambda x: x.get("score", 0), reverse=True)
        kept = kept[:top_k]
        total_kept += len(kept)

        if kept:
            for r in kept:
                doc.add_paragraph(f"Comment (score {r['score']}): {r['comment']}")
                if r.get("reason"): doc.add_paragraph(f"Reason: {r['reason']}")
                if r.get("url"): doc.add_paragraph(f"Source: {r['url']}")
        else:
            doc.add_paragraph("No strongly correlated comments found for this frame.")

    doc.save(out_docx)


# -------------------- Core processing --------------------

def find_caption_json(video_id: str, search_root: str):
    """
    Search for <video_id>_captions_integrated.json in search_root (recursive).
    """
    fname = f"{video_id}_captions_integrated.json"
    for root, dirs, files in os.walk(search_root):
        if fname in files:
            return os.path.join(root, fname)
    return None

def process_one_video(raw_frames_dir: str, frames_input_base: str, output_base: str,
                      comments_by_video: dict, model: str,
                      min_score: int, top_k_scan: int, integrated_captions_root: str):
    """
    raw_frames_dir: absolute path to e.g. .../<video_id>/raw_frames
    frames_input_base: origin path user pointed to (to mirror structure).
    output_base: root where all results should go.
    """
    # video_id is parent of raw_frames
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))

    # Calculate mirror output path
    # e.g. input:  /A/B/video_id/raw_frames
    #      base:   /A
    #      rel:    B/video_id/raw_frames
    #      actual output dir: /Out/B/video_id (we drop raw_frames from the end usually, or keep it?)
    # Let's clean it up: we want results in /Out/.../video_id/
    
    # 1. Get relative path from input base to the video folder
    # video_folder = .../video_id
    video_folder_abs = os.path.dirname(raw_frames_dir)
    try:
        rel = os.path.relpath(video_folder_abs, frames_input_base)
    except ValueError:
        # separate drives on Windows or similar issue
        rel = video_id

    # 2. Results output dir
    video_out_dir = os.path.join(output_base, rel)
    os.makedirs(video_out_dir, exist_ok=True)

    # find integrated json
    integrated_json = find_caption_json(video_id, integrated_captions_root)
    if not integrated_json:
        print(f"[SKIP] {video_id}: integrated JSON not found in {integrated_captions_root}")
        return

    frames = parse_integrated_json(integrated_json)
    if not frames:
        print(f"[SKIP] {video_id}: no frames parsed from JSON.")
        return

    comments_for_video = comments_by_video.get(video_id, [])
    if not comments_for_video:
        # fallback: use all comments (still resume-safe)
        comments_for_video = [c for v in comments_by_video.values() for c in v]
        print(f"[WARN] {video_id}: no comments matched; falling back to ALL comments ({len(comments_for_video)}).")
    else:
        print(f"[INFO] {video_id}: {len(comments_for_video)} comments found for this video.")

    # resume state
    res_file = results_path_jsonl(video_out_dir, video_id)
    existing = load_existing_results(res_file)
    already = set(existing.keys())
    print(f"[RESUME] {video_id}: {len(already)}/{len(frames)} frames already processed.")

    # process frames
    for f in frames:
        frame_name = f["frame"]
        if frame_name in already:
            continue

        caption = f["caption"]
        candidates = []
        checked_pairs = 0

        for item in comments_for_video:
            res = ask_model(model, caption, item["comment"])
            checked_pairs += 1
            # keep all raw candidates; filtering happens when building docx
            candidates.append({
                "correlated": bool(res["correlated"]),
                "score": int(res["score"]),
                "reason": res.get("reason", ""),
                "comment": item["comment"],
                "url": item["url"]
            })

        record = {
            "video_id": video_id,
            "frame": frame_name,
            "checked_pairs": checked_pairs,
            "candidates": candidates
        }
        append_result(res_file, record)
        existing[frame_name] = record
        print(f"[OK] {video_id}: frame {frame_name} processed ({checked_pairs} pairs).")

    # always (re)build docx from accumulated results so resume is seamless
    out_docx = os.path.join(video_out_dir, f"{video_id}_correlation.docx")
    
    build_docx_from_results(
        out_docx, frames, existing,
        min_score=min_score,
        top_k=top_k_scan,
        model_name=model,
        raw_frames_dir=raw_frames_dir, # need this for thumbnails
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[DONE] {video_id}: report saved -> {out_docx}")

# -------------------- Batch driver --------------------

def discover_targets(frames_input: str):
    """
    Returns list of raw_frames directories to process.
    Accepts either:
      - path ending with /raw_frames  (single video)
      - frames root containing many <video_id>/raw_frames (recursive)
    """
    frames_input = os.path.abspath(frames_input)
    targets = []
    
    # Direct target?
    if os.path.basename(frames_input) == "raw_frames":
        targets.append(frames_input)
        return targets

    # Recursive scan
    for root, dirs, files in os.walk(frames_input):
        if "raw_frames" in dirs:
            targets.append(os.path.join(root, "raw_frames"))
            
    return sorted(targets)

def rebuild_docx_only(raw_frames_dir: str, frames_input_base: str, output_base: str,
                      comments_by_video: dict,
                      min_score: int, top_k_scan: int, model: str, integrated_captions_root: str):
    
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))
    
    # Calculate output path
    video_folder_abs = os.path.dirname(raw_frames_dir)
    try:
        rel = os.path.relpath(video_folder_abs, frames_input_base)
    except ValueError:
        rel = video_id
    video_out_dir = os.path.join(output_base, rel)
    
    # Check if results exist
    res_file = results_path_jsonl(video_out_dir, video_id)
    if not os.path.isfile(res_file):
        print(f"[SKIP] {video_id}: no results file found in {video_out_dir} to rebuild.")
        return

    # find integrated json
    integrated_json = find_caption_json(video_id, integrated_captions_root)
    if not integrated_json:
        print(f"[SKIP] {video_id}: integrated JSON not found for rebuild.")
        return
    frames = parse_integrated_json(integrated_json)
    if not frames:
        return

    existing = load_existing_results(res_file)
    comments_for_video = comments_by_video.get(video_id, [])
    
    out_docx = os.path.join(video_out_dir, f"{video_id}_correlation.docx")

    build_docx_from_results(
        out_docx, frames, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        raw_frames_dir=raw_frames_dir,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[REBUILT] {video_id}: report saved -> {out_docx}")


def main():
    ap = argparse.ArgumentParser(description="Batch correlate integrated frame captions with YouTube comments (resume-safe).")
    ap.add_argument("--frames_path", default="../../VideoAnalysis/data/frames", help="Path to frames root (containing <video_id>/raw_frames) OR a single raw_frames folder.")
    ap.add_argument("--comments_docx", default="../../Data For Classes-Analysis/final_merged_data/final_merged_class_data.docx", help="DOCX with video URLs and comments.")
    ap.add_argument("--results_root", default="../results", help="Directory where results/reports will be saved (mirrors input structure).")
    ap.add_argument("--integrated_captions_root", default="../data/integrated_caption", help="Root directory containing <video_id>_captions_integrated.json files.")
    
    ap.add_argument("--model", default=None, help="Model name (defaults to Config.ollama_models[0]).")
    ap.add_argument("--min_score", type=int, default=60, help="Minimum score to include a match in the final report.")
    ap.add_argument("--top_k", type=int, default=5, help="Top-K comments per frame in the final report.")
    ap.add_argument("--rebuild", action="store_true",
                    help="Rebuild DOCX from existing results without running the model.")
    args = ap.parse_args()

    model = args.model or Config.ollama_models[0]
    results_root = os.path.abspath(args.results_root)
    frames_input_base = os.path.abspath(args.frames_path)

    # load all comments once (used for % in summary)
    comments_by_video = parse_comments_docx(args.comments_docx)
    if not comments_by_video:
        print("[FATAL] No comments parsed from comments DOCX.")
        return

    targets = discover_targets(args.frames_path)
    if not targets:
        print("[FATAL] No raw_frames folders found.")
        return

    print(f"[INFO] Found {len(targets)} video(s) to process. Results -> {results_root}")

    for raw_frames in targets:
        print(f"\n=== Processing: {os.path.basename(os.path.dirname(raw_frames))} ===")
        try:
            if args.rebuild:
                rebuild_docx_only(
                    raw_frames, frames_input_base, results_root,
                    comments_by_video,
                    min_score=args.min_score, top_k_scan=args.top_k, model=model,
                    integrated_captions_root=args.integrated_captions_root
                )
            else:
                process_one_video(
                    raw_frames, frames_input_base, results_root,
                    comments_by_video, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    integrated_captions_root=args.integrated_captions_root
                )
        except KeyboardInterrupt:
            print("\n[INTERRUPTED] Stopping cleanly. You can rerun to resume.")
            break
        except Exception as e:
            print(f"[ERROR] {raw_frames}: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    main()
