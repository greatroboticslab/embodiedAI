import os
import re
import signal
from contextlib import contextmanager
from docx import Document
from config import Config
from utils.llm_utils import stream_parser, generate_response

class TimeoutException(Exception): pass

@contextmanager
def timeout(time):
    def raise_timeout(signum, frame):
        raise TimeoutException
    signal.signal(signal.SIGALRM, raise_timeout)
    signal.alarm(time)
    try:
        yield
    finally:
        signal.alarm(0)

llama_model = Config.ollama_models[0]

def extract_processed_frames(docx_path):
    processed_frames = set()
    if os.path.exists(docx_path):
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading 1"):
                match = re.match(r'(?:Frame: )?(frame_\d{4}\.jpg)', para.text.strip(), re.I)
                if match:
                    processed_frames.add(match.group(1).lower())
    return processed_frames

def extract_captions(docx_file):
    doc = Document(docx_file)
    captions = {}
    current_frame = None
    for para in doc.paragraphs:
        text = para.text.strip()
        match = re.match(r'(?:Frame: )?(frame_\d{4}\.jpg)', text)
        if match:
            current_frame = match.group(1)
            captions[current_frame] = ""
        elif current_frame and text:
            captions[current_frame] += (" " + text)
    return captions

def generate_integrated_caption(llava_caption, minicpm_caption):
    llama_prompt = f"""
    Carefully integrate these two captions into a single highly detailed caption. Your integrated caption must:
    - Retain all accurate and relevant details explicitly mentioned in either original caption.
    - Avoid redundancy and contradictions.
    - Strictly avoid any speculation, assumptions, or introduction of details not explicitly stated.

    Original Captions:
    1. LLaVA: {llava_caption}
    2. MiniCPM: {minicpm_caption}

    Provide your response strictly in this format:
    Final Caption: <Integrated caption>
    """
    try:
        with timeout(600):
            response_stream = generate_response(llama_model, llama_prompt)
            response_text = ''.join(stream_parser(response_stream)).strip()
    except TimeoutException:
        return generate_integrated_caption(llava_caption, minicpm_caption)

    match = re.search(r'Final Caption:\s*(.+)', response_text, re.DOTALL)
    return match.group(1).strip() if match else llava_caption

def integrate_captions(root_folder, combined_filename):
    combined_doc = Document()
    combined_doc.add_heading('Combined Integrated Captions', 0)

    for subdir, _, files in os.walk(root_folder):
        images = [f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
        if not images:
            continue

        folder_name = os.path.relpath(subdir, root_folder)
        parent_dir = os.path.basename(os.path.dirname(subdir))
        current_dir = os.path.basename(subdir)
        llava_docx = os.path.join(subdir, f"{parent_dir}_{current_dir}_captions_llava.docx")
        minicpm_docx = os.path.join(subdir, f"{parent_dir}_{current_dir}_captions_MiniCPM.docx")
        integrated_docx = os.path.join(subdir, f"{parent_dir}_{current_dir}_captions_integrated.docx")

        processed_frames = extract_processed_frames(integrated_docx)
        llava_captions = extract_captions(llava_docx)
        minicpm_captions = extract_captions(minicpm_docx)

        integrated_doc = Document() if not os.path.exists(integrated_docx) else Document(integrated_docx)
        if not os.path.exists(integrated_docx):
            integrated_doc.add_heading(f'Integrated Captions: {folder_name}', 0)

        new_captions = 0
        for img in sorted(images):
            if img.lower() in processed_frames:
                continue
            if img in llava_captions and img in minicpm_captions:
                final_caption = generate_integrated_caption(llava_captions[img], minicpm_captions[img])
                integrated_doc.add_heading(f'Frame: {img}', level=1)
                integrated_doc.add_paragraph(final_caption)
                integrated_doc.save(integrated_docx)
                new_captions += 1
                print(f"Processed: {img}")

        if new_captions > 0:
            combined_doc.add_page_break()
            combined_doc.add_heading(f"Folder: {folder_name}", level=1)
            for element in integrated_doc.element.body:
                combined_doc.element.body.append(element)

    combined_doc_path = os.path.join(root_folder, f"{combined_filename}.docx")
    combined_doc.save(combined_doc_path)
    print(f"✅ All folders processed. Combined doc saved: {combined_doc_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate integrated captions from two model outputs.")
    parser.add_argument("root_folder", help="Root folder containing subfolders of frames")
    parser.add_argument("combined_docx_filename", help="Filename for the combined .docx (no extension)")
    args = parser.parse_args()

    integrate_captions(args.root_folder, args.combined_docx_filename)
