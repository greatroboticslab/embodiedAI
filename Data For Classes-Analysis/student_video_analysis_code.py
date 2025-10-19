import os
import re
from collections import defaultdict, Counter
from docx import Document
from config import Config
from utils.llm_utils import stream_parser, generate_response

# Load model
llm_model = Config.ollama_models[1]

# Prompts
LIKE_DISLIKE_PROMPT = (
    "Based on the student's paragraph below, determine if they liked the video, "
    "disliked parts of it, or were neutral. Give a clear label: Like, Dislike, or Neutral."
)

REASON_PROMPT = (
    "What is the main reason the student liked or disliked the video below? Respond with a concise phrase."
)

EMBODIED_LEARNING_PROMPT = (
    "Does the following paragraph show signs of embodied learning (i.e., learning through physical interaction, sensory/motor engagement, hands-on experience)? "
    "Answer Yes or No. If Yes, briefly explain why."
)

OTHER_LEARNING_PROMPT = (
    "If the student enjoyed the video but it doesn't involve embodied learning, identify which other learning paradigm are involved. The examples of alternative learning paradigms include:"
    "powerpoint slides based learning, observing learning, simulation based learning, project based learning, peer learning and study groups, learning through research papers, online courses, competition among others."
)

LOGICAL_REASONING_PROMPT = (
    "Based on the student's paragraph below, determine if students have learnt something significant from the video, "
    "The determination is based on their logical reasoning process of the materials learnt through the process."
)

LOGICAL_EMBODIED_PROMPT = (
    "Based on the student's paragraph below, can you find any embodied learning components based on their logical reasoning of the learning."
)

# Extract opinions grouped by video URL
def extract_opinions_by_video(doc_path):
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    video_opinions = defaultdict(list)
    current_link = None

    for para in paragraphs:
        links = re.findall(r"https://www\.youtube\.com/watch\?v=[\w-]+", para)
        if links:
            current_link = links[0]
        elif current_link:
            video_opinions[current_link].append(para)

    return video_opinions

# Analyze all opinions for one video
def analyze_opinions(video_link, opinions):
    analysis_data = {
        "link": video_link,
        "opinions": opinions,
        "like": 0,
        "dislike": 0,
        "embodied_like": 0,
        "embodied_dislike": 0,
        "like_reasons": Counter(),
        "dislike_reasons": Counter(),
        "embodied_reasons": Counter(),
        "other_learning": Counter(),
        "detailed_feedback": []
    }

    for para in opinions:
        sentiment_stream = generate_response(llm_model, f"{LIKE_DISLIKE_PROMPT}\n\nParagraph: {para}")
        sentiment = ''.join(stream_parser(sentiment_stream)).lower()

        reason_stream = generate_response(llm_model, f"{REASON_PROMPT}\n\nParagraph: {para}")
        reason = ''.join(stream_parser(reason_stream)).strip().lower()

        emb_stream = generate_response(llm_model, f"{EMBODIED_LEARNING_PROMPT}\n\nParagraph: {para}")
        emb_response = ''.join(stream_parser(emb_stream)).strip().lower()
        emb_match = re.match(r"yes.*?(\-|:)?\s*(.*)", emb_response)
        is_embodied = 'yes' in emb_response
        emb_reason = emb_match.group(2) if emb_match else None

        logic_stream = generate_response(llm_model, f"{LOGICAL_REASONING_PROMPT}\n\nParagraph: {para}")
        logic_reasoning = ''.join(stream_parser(logic_stream)).strip()

        logic_embodied_stream = generate_response(llm_model, f"{LOGICAL_EMBODIED_PROMPT}\n\nParagraph: {para}")
        logic_embodied_reason = ''.join(stream_parser(logic_embodied_stream)).strip()

        if "like" in sentiment:
            analysis_data["like"] += 1
            if reason:
                analysis_data["like_reasons"][reason] += 1
            if is_embodied:
                analysis_data["embodied_like"] += 1
                if emb_reason:
                    analysis_data["embodied_reasons"][emb_reason] += 1
        elif "dislike" in sentiment:
            analysis_data["dislike"] += 1
            if reason:
                analysis_data["dislike_reasons"][reason] += 1
            if is_embodied:
                analysis_data["embodied_dislike"] += 1
                if emb_reason:
                    analysis_data["embodied_reasons"][emb_reason] += 1

        if "like" in sentiment and not is_embodied:
            other_stream = generate_response(llm_model, f"{OTHER_LEARNING_PROMPT}\n\nParagraph: {para}")
            other = ''.join(stream_parser(other_stream)).strip()
            if other and other.lower() != 'embodied':
                analysis_data["other_learning"][other] += 1

        # Save per-student data
        analysis_data["detailed_feedback"].append({
            "opinion": para,
            "logical": logic_reasoning,
            "embodied": emb_response,
            "logical_embodied": logic_embodied_reason
        })

    return analysis_data

# Generate per-video output docx
def generate_individual_output_docs(video_analysis, output_dir):
    os.makedirs(output_dir, exist_ok=True)

    for video_link, data in video_analysis.items():
        video_id_match = re.search(r"v=([\w-]+)", video_link)
        video_id = video_id_match.group(1) if video_id_match else "unknown"
        filename = os.path.join(output_dir, f"video_feedback_{video_id}.docx")

        doc = Document()
        doc.add_heading(f"Video: {video_link}", level=1)

        doc.add_heading("Analysis Summary:", level=2)
        total_students = len(data["opinions"])

        # Final logic based on actual content of logical and logical_embodied fields
        total_logical = 0
        total_logical_embodied = 0

        for entry in data["detailed_feedback"]:
            logic_text = entry["logical"].strip().lower()
            logic_embodied_text = entry["logical_embodied"].strip().lower()

            if any(keyword in logic_text for keyword in ["yes", "student shows", "logical reasoning", "they reasoned", "demonstrates understanding"]):
                total_logical += 1
                if any(key in logic_embodied_text for key in ["yes", "embodied", "hands-on", "physical", "sensor", "motor"]):
                    total_logical_embodied += 1

        doc.add_paragraph(f"Total Student Opinions: {total_students}")
        doc.add_paragraph(f"Total Likes: {data['like']}")
        for reason, count in data['like_reasons'].most_common():
            doc.add_paragraph(f"  - {reason} ({count})")

        doc.add_paragraph(f"Total Dislikes: {data['dislike']}")
        for reason, count in data['dislike_reasons'].most_common():
            doc.add_paragraph(f"  - {reason} ({count})")

        doc.add_paragraph(f"Embodied Learning in Likes: {data['embodied_like']}")
        doc.add_paragraph(f"Embodied Learning in Dislikes: {data['embodied_dislike']}")
        if data['embodied_reasons']:
            doc.add_paragraph("Embodied Learning Evidence:")
            for reason, count in data['embodied_reasons'].most_common():
                doc.add_paragraph(f"  - {reason} ({count})")

        if data['other_learning']:
            doc.add_paragraph("Other Learning Paradigms:")
            for paradigm, count in data['other_learning'].most_common():
                doc.add_paragraph(f"  - {paradigm} ({count})")

        doc.add_paragraph(f"Total Logical Reasoning Occurrences: {total_logical}")
        doc.add_paragraph(f"Logical Reasoning with Embodied Learning: {total_logical_embodied}")

        doc.add_page_break()
        doc.add_heading("Detailed Student Feedback", level=2)
        for i, entry in enumerate(data['detailed_feedback'], 1):
            doc.add_paragraph(f"Student Opinion {i}:")
            doc.add_paragraph(entry['opinion'], style='Intense Quote')
            doc.add_paragraph(f"Logical Reasoning: {entry['logical'] or 'None'}")
            doc.add_paragraph(f"Embodied Learning: {entry['embodied'] or 'None'}")
            doc.add_paragraph(f"Embodied Learning in Logical Reasoning: {entry['logical_embodied'] or 'None'}")
            doc.add_paragraph("---")

        doc.save(filename)
        print(f"Saved: {filename}")

# Main execution
if __name__ == '__main__':
    input_path = "/home/mtsu/workspace/reshma/student_video_analysis/final_merged_class_data.docx"
    output_dir = "/home/mtsu/workspace/reshma/student_video_analysis/video_outputs6"

    print("Extracting opinions...")
    grouped_opinions = extract_opinions_by_video(input_path)

    print(f"Found {len(grouped_opinions)} unique video links.")

    all_analysis = {}
    for link, paras in grouped_opinions.items():
        print(f"Analyzing video: {link} with {len(paras)} opinions...")
        result = analyze_opinions(link, paras)
        all_analysis[link] = result

    print("Generating individual output documents...")
    generate_individual_output_docs(all_analysis, output_dir)
    print(f"All analysis saved to directory: {output_dir}")
