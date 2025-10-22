#!/usr/bin/env python3
"""
Collect per-video metrics for an Embodied vs. Conventional video study, save a JSON summary,
and generate a DOCX table report.

Inputs (paths you pass via CLI):
  --videos_root       Path to the folder holding raw videos (optional; used for duration if ffprobe/ffprobe-python available)
  --frames_root       Path to the folder holding extracted frames. Expected structure: <frames_root>/<video_id>/raw_frames/
                      NOTE: The correlation result JSON is expected to live either in <video_id>/ or in <video_id>/raw_frames/.
  --transcripts_root  Path to the folder with transcripts. Accepts either <video_id>.txt OR <video_id>.json (with "segments").
  --topics_root       Path to the folder with topic JSON files (default name: <video_id>.topics.json).
  --out_dir           Output directory to write results JSON and DOCX.

Run the same command twice—once for your Conventional set, once for your Embodied set.

Example:
  python embodiment_project_collect_metrics.py \
      --videos_root ../../video_processing/rawvideos \
      --frames_root ../../video_processing/frames \
      --transcripts_root ../../video_processing/transcripts \
      --topics_root ../../video_processing/topics \
      --out_dir ./outputs/conventional \
      --label Conventional

Then repeat with a different --out_dir and --label for Embodied.

What this script computes per video_id (best-effort, robust to missing pieces):
  - frames_count: count of image files in raw_frames
  - transcript_word_count: words in <video_id>.txt or sum over JSON segments
  - topics_count: len of topics array in <video_id>.topics.json
  - comments_count: number of comments found in a correlation JSON near frames
  - corr_scores: list of correlation scores (if present) and summary stats (avg, median)
  - num_correlated_comments_thresholded: number of comments with score >= --corr_threshold
  - duration_seconds: if a video exists and ffprobe is available, duration in seconds

It also outputs aggregated summary across all videos.

Assumptions & Heuristics for correlation JSON:
  - We search for a JSON file inside <frames_root>/<video_id>/ and .../<video_id>/raw_frames/.
  - We try common shapes, e.g. {'correlations': [...]}, list of dicts with keys like 'comment', 'score', etc.
  - If multiple candidates are found, we pick the one with the most entries containing numeric 'score' fields.

Dependencies:
  - python-docx (pip install python-docx)
  - (optional) ffprobe available in PATH to compute duration (or install ffmpeg). If unavailable, duration stays None.

Outputs:
  - <out_dir>/<label>_metrics.json
  - <out_dir>/<label>_metrics.docx
"""

import csv
import argparse
import json
import os
import re
import subprocess
from collections import defaultdict
from dataclasses import dataclass, asdict
from statistics import mean, median
from typing import Dict, List, Optional, Tuple, Any

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    raise SystemExit("python-docx is required. Install with: pip install python-docx")

IMG_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}

# ---------------------- Utilities ----------------------

def safe_stem(path: str) -> str:
    """Return filename stem without extension."""
    base = os.path.basename(path)
    stem, _ = os.path.splitext(base)
    return stem

def _none(x):
    return "" if x is None else x




def list_video_ids_from_frames(frames_root: str) -> List[str]:
    """Assume frames_root/<video_id>/raw_frames exists; use those subdirs as canonical IDs."""
    if not os.path.isdir(frames_root):
        return []
    vids = []
    for name in os.listdir(frames_root):
        vid_dir = os.path.join(frames_root, name)
        if not os.path.isdir(vid_dir):
            continue
        raw_frames = os.path.join(vid_dir, "raw_frames")
        if os.path.isdir(raw_frames):
            vids.append(name)
    return sorted(vids)

import glob

def list_correlation_result_paths(frames_video_dir: str) -> list[str]:
    """
    Return all correlation result JSONL/JSON files near frames.
    Looks in <vid_dir>/ and <vid_dir>/raw_frames/.
    """
    out = []
    for folder in [frames_video_dir, os.path.join(frames_video_dir, 'raw_frames')]:
        if not os.path.isdir(folder):
            continue
        # Your generator uses *.results.jsonl
        out.extend(glob.glob(os.path.join(folder, "*_topiccorr_*.results.jsonl")))
        # Fallbacks (if you ever export .json or .jsonl without that stem)
        out.extend(glob.glob(os.path.join(folder, "*.results.jsonl")))
        out.extend(glob.glob(os.path.join(folder, "*.jsonl")))
        out.extend(glob.glob(os.path.join(folder, "*.json")))
    return sorted(set(out))



def count_frames(raw_frames_dir: str) -> int:
    cnt = 0
    if not os.path.isdir(raw_frames_dir):
        return 0
    for fn in os.listdir(raw_frames_dir):
        if os.path.splitext(fn)[1].lower() in IMG_EXTS:
            cnt += 1
    return cnt


def read_transcript_word_count(transcripts_root: str, vid: str) -> int:
    """Try <vid>.txt then <vid>.json with segments[].text."""
    txt_path = os.path.join(transcripts_root, f"{vid}.txt")
    if os.path.isfile(txt_path):
        try:
            with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return len(re.findall(r"\b\w+\b", text))
        except Exception:
            return 0
    json_path = os.path.join(transcripts_root, f"{vid}.json")
    if os.path.isfile(json_path):
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            segs = data.get("segments") or []
            text = " ".join([str(s.get("text", "")) for s in segs])
            return len(re.findall(r"\b\w+\b", text))
        except Exception:
            return 0
    return 0


def read_topics_count(topics_root: str, vid: str) -> int:
    # default filename
    path = os.path.join(topics_root, f"{vid}.topics.json")
    if not os.path.isfile(path):
        # fallback: any json starting with vid and containing 'topics'
        candidates = [
            os.path.join(topics_root, fn) for fn in os.listdir(topics_root)
            if fn.startswith(vid) and fn.endswith('.json')
        ] if os.path.isdir(topics_root) else []
        for c in candidates:
            try:
                with open(c, 'r', encoding='utf-8') as f:
                    d = json.load(f)
                if isinstance(d, dict) and isinstance(d.get('topics'), list):
                    return len(d['topics'])
            except Exception:
                continue
        return 0
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        topics = data.get('topics') or []
        return len(topics)
    except Exception:
        return 0


def find_best_correlation_json(frames_video_dir: str) -> Optional[str]:
    if not os.path.isdir(frames_video_dir):
        return None

    candidates = []
    for folder in [frames_video_dir, os.path.join(frames_video_dir, "raw_frames")]:
        if not os.path.isdir(folder):
            continue
        for fn in os.listdir(folder):
            if fn.lower().endswith(".json"):
                candidates.append(os.path.join(folder, fn))
    if not candidates:
        return None

    def score_json(path: str) -> Tuple[int, int]:
        """Return (#records, #records_with_numeric_score), drilling into 'candidates'."""
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            return (0, 0)

        total = 0
        with_score = 0

        def _acc(obj: Any):
            nonlocal total, with_score
            if isinstance(obj, dict):
                # If topic-level object has a candidates list, inspect those
                if isinstance(obj.get("candidates"), list):
                    for c in obj["candidates"]:
                        if isinstance(c, dict):
                            total += 1
                            if isinstance(c.get("score"), (int, float)):
                                with_score += 1
                # Also check other common list keys
                for k in ["correlations", "results", "items", "data"]:
                    if isinstance(obj.get(k), list):
                        for it in obj[k]:
                            _acc(it)
            elif isinstance(obj, list):
                for it in obj:
                    _acc(it)
            elif isinstance(obj, (str, int, float, type(None))):
                return

        _acc(data)
        return (total, with_score)

    best = None
    best_tuple = (-1, -1)
    for p in candidates:
        t = score_json(p)
        if t > best_tuple:
            best_tuple = t
            best = p
    return best



def parse_correlation_json(path: str) -> Tuple[int, List[float]]:
    """
    Return (comments_count, scores[]) for a correlation JSON.

    Handles shapes like:
      [
        { "video_id": "...", "topic_id": "T1", "title": "...",
          "checked_pairs": 7,
          "candidates": [
             {"correlated": false, "score": 20, "comment": "...", "url": "..."},
             ...
          ]
        },
        ...
      ]

    Also still supports older shapes where items themselves have 'comment' and 'score'.
    Counts UNIQUE comment texts across the whole file as 'comments_count'.
    """
    if not path or not os.path.isfile(path):
        return 0, []

    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return 0, []

    # Collect all leaf-level candidate dicts in one list
    leaf_items: List[Dict[str, Any]] = []

    def _maybe_extend_from(obj: Any):
        """
        If obj is a dict with a list of correlation records under a known key,
        extend leaf_items with those. Otherwise if obj itself looks like a record,
        append it.
        """
        if isinstance(obj, dict):
            # common container keys in various pipelines
            for key in ["candidates", "correlations", "results", "items", "data"]:
                if isinstance(obj.get(key), list):
                    for it in obj[key]:
                        if isinstance(it, dict):
                            leaf_items.append(it)
                    return
            # If it's a topic-level dict that has 'candidates', handle above.
            # Otherwise, it might itself be a leaf correlation record.
            leaf_items.append(obj)
        elif isinstance(obj, list):
            for it in obj:
                _maybe_extend_from(it)

    _maybe_extend_from(data)

    comments_seen = set()
    scores: List[float] = []

    # Helper: extract a comment text from a record using several possible keys
    def _get_comment_text(rec: Dict[str, Any]) -> Optional[str]:
        for k in ("comment", "comment_text", "text", "body"):
            v = rec.get(k)
            if isinstance(v, str) and v.strip():
                return v.strip()
        return None

    for rec in leaf_items:
        if not isinstance(rec, dict):
            continue

        # If this record itself has nested 'candidates', drill down
        if isinstance(rec.get("candidates"), list):
            for c in rec["candidates"]:
                if not isinstance(c, dict):
                    continue
                ct = _get_comment_text(c)
                if ct:
                    comments_seen.add(ct)
                sc = c.get("score")
                if isinstance(sc, (int, float)):
                    scores.append(float(sc))
            continue

        # Flat record path
        ct = _get_comment_text(rec)
        if ct:
            comments_seen.add(ct)
        sc = rec.get("score")
        if isinstance(sc, (int, float)):
            scores.append(float(sc))

        # Some pipelines store nested matches
        if isinstance(rec.get("matches"), list):
            for m in rec["matches"]:
                if isinstance(m, dict):
                    ct2 = _get_comment_text(m)
                    if ct2:
                        comments_seen.add(ct2)
                    sc2 = m.get("score")
                    if isinstance(sc2, (int, float)):
                        scores.append(float(sc2))

    comments_count = len(comments_seen)
    return comments_count, scores

from typing import Any, Dict, List, Optional, Tuple

def parse_correlation_results(paths: List[str]) -> Tuple[int, List[float]]:
    """
    Aggregate (unique comments_count, scores[]) across one or more
    correlation result files. Supports your per-topic JSONL schema:

      {"video_id": "...", "topic_id": "T1", "title": "...",
       "checked_pairs": 7,
       "candidates": [
          {"correlated": false, "score": 20, "reason": "...",
           "comment": "...", "url": "..."},
          ...
       ]}

    Returns:
      comments_count: # of UNIQUE comment texts across all topics/files
      scores: all numeric scores found in candidates (floats)
    """
    comments_seen = set()
    scores: List[float] = []

    def _get_comment_text(rec: Dict[str, Any]) -> Optional[str]:
        for k in ("comment", "comment_text", "text", "body"):
            v = rec.get(k)
            if isinstance(v, str) and v.strip():
                return v.strip()
        return None

    for path in paths:
        if not os.path.isfile(path):
            continue

        # JSONL: iterate lines
        if path.lower().endswith(".jsonl"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                    except Exception:
                        continue

                    cands = obj.get("candidates") or []
                    if isinstance(cands, list):
                        for c in cands:
                            if not isinstance(c, dict):
                                continue
                            ct = _get_comment_text(c)
                            if ct:
                                comments_seen.add(ct)
                            sc = c.get("score")
                            if isinstance(sc, (int, float)):
                                scores.append(float(sc))
            continue

        # Plain JSON fallback (single object or list)
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            continue

        # If it's a list of topic records
        if isinstance(data, list):
            for obj in data:
                if not isinstance(obj, dict):
                    continue
                for c in obj.get("candidates", []) or []:
                    if not isinstance(c, dict):
                        continue
                    ct = _get_comment_text(c)
                    if ct:
                        comments_seen.add(ct)
                    sc = c.get("score")
                    if isinstance(sc, (int, float)):
                        scores.append(float(sc))
        elif isinstance(data, dict):
            # Single topic record in JSON (rare)
            for c in data.get("candidates", []) or []:
                if not isinstance(c, dict):
                    continue
                ct = _get_comment_text(c)
                if ct:
                    comments_seen.add(ct)
                sc = c.get("score")
                if isinstance(sc, (int, float)):
                    scores.append(float(sc))

    return len(comments_seen), scores




def get_duration_seconds(videos_root: Optional[str], vid: str) -> Optional[float]:
    if not videos_root:
        return None
    # Find a video file that starts with vid.*
    if not os.path.isdir(videos_root):
        return None
    cand = None
    for fn in os.listdir(videos_root):
        if not fn.startswith(vid + ""):
            continue
        if os.path.splitext(fn)[1].lower() in {'.mp4', '.mkv', '.mov', '.avi', '.webm'}:
            cand = os.path.join(videos_root, fn)
            break
    if not cand or not os.path.isfile(cand):
        return None
    try:
        # Use ffprobe to get duration in seconds
        proc = subprocess.run([
            'ffprobe', '-v', 'error', '-select_streams', 'v:0',
            '-show_entries', 'format=duration', '-of', 'default=nw=1:nk=1', cand
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=10)
        out = proc.stdout.strip()
        if out:
            return float(out)
    except Exception:
        return None
    return None

# ---------------------- Data Classes ----------------------

@dataclass
class VideoMetrics:
    video_id: str
    frames_count: int
    transcript_word_count: int
    topics_count: int
    comments_count: int
    corr_avg: Optional[float]
    corr_median: Optional[float]
    corr_n: int
    corr_n_ge_threshold: int
    duration_seconds: Optional[float]


@dataclass
class Aggregate:
    label: str
    num_videos: int
    avg_frames: Optional[float]
    avg_transcript_words: Optional[float]
    avg_topics: Optional[float]
    avg_comments: Optional[float]
    avg_corr_avg: Optional[float]
    median_corr_avg: Optional[float]

# ---------------------- Core ----------------------

def build_csv_files(out_dir: str, label: str, metrics: List[VideoMetrics], agg: Aggregate, corr_threshold: float):
    """
    Write two CSVs:
      1) <label>_metrics.csv    — per-video rows (good for plotting in Overleaf)
      2) <label>_aggregate.csv  — single-row summary (handy for quick references)
    """
    # --- Per-video CSV ---
    headers = [
        "video_id", "frames", "topics", "comments",
        "corr_avg", "corr_median", "corr_n", f"corr_n_ge_{corr_threshold}",
        "transcript_words", "duration_s"
    ]
    per_video_csv = os.path.join(out_dir, f"{label}_metrics.csv")
    with open(per_video_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(headers)
        for m in metrics:
            w.writerow([
                m.video_id,
                m.frames_count,
                m.topics_count,
                m.comments_count,
                _none(m.corr_avg),
                _none(m.corr_median),
                m.corr_n,
                m.corr_n_ge_threshold,
                m.transcript_word_count,
                _none(m.duration_seconds),
            ])

    # --- Aggregate CSV (single row) ---
    agg_headers = [
        "label", "num_videos", "avg_frames", "avg_transcript_words",
        "avg_topics", "avg_comments", "avg_corr_avg", "median_corr_avg",
        "corr_threshold"
    ]
    agg_csv = os.path.join(out_dir, f"{label}_aggregate.csv")
    with open(agg_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(agg_headers)
        w.writerow([
            agg.label,
            agg.num_videos,
            _none(agg.avg_frames),
            _none(agg.avg_transcript_words),
            _none(agg.avg_topics),
            _none(agg.avg_comments),
            _none(agg.avg_corr_avg),
            _none(agg.median_corr_avg),
            corr_threshold,
        ])

    print(f"Wrote: {per_video_csv}")
    print(f"Wrote: {agg_csv}")



def compute_metrics_for_set(
    videos_root: Optional[str],
    frames_root: str,
    transcripts_root: str,
    topics_root: str,
    label: str,
    corr_threshold: float = 0.0,
) -> Tuple[List[VideoMetrics], Aggregate]:

    video_ids = list_video_ids_from_frames(frames_root)

    metrics: List[VideoMetrics] = []
    corr_avgs: List[float] = []

    for vid in video_ids:
        vid_dir = os.path.join(frames_root, vid)
        raw_frames_dir = os.path.join(vid_dir, 'raw_frames')
        frames_count = count_frames(raw_frames_dir)
        transcript_words = read_transcript_word_count(transcripts_root, vid)
        topics_count = read_topics_count(topics_root, vid)
        corr_paths = list_correlation_result_paths(vid_dir)
        comments_count, scores = parse_correlation_results(corr_paths)

        corr_avg = mean(scores) if scores else None
        corr_median = median(scores) if scores else None
        corr_n = len(scores)
        corr_n_ge = sum(1 for s in scores if s >= corr_threshold)
        dur_s = get_duration_seconds(videos_root, vid)

        vm = VideoMetrics(
            video_id=vid,
            frames_count=frames_count,
            transcript_word_count=transcript_words,
            topics_count=topics_count,
            comments_count=comments_count,
            corr_avg=corr_avg,
            corr_median=corr_median,
            corr_n=corr_n,
            corr_n_ge_threshold=corr_n_ge,
            duration_seconds=dur_s,
        )
        metrics.append(vm)
        if corr_avg is not None:
            corr_avgs.append(corr_avg)

    def _avg(vals: List[int]) -> Optional[float]:
        return mean(vals) if vals else None

    agg = Aggregate(
        label=label,
        num_videos=len(metrics),
        avg_frames=_avg([m.frames_count for m in metrics]),
        avg_transcript_words=_avg([m.transcript_word_count for m in metrics]),
        avg_topics=_avg([m.topics_count for m in metrics]),
        avg_comments=_avg([m.comments_count for m in metrics]),
        avg_corr_avg=mean(corr_avgs) if corr_avgs else None,
        median_corr_avg=median(corr_avgs) if corr_avgs else None,
    )

    return metrics, agg

# ---------------------- DOCX ----------------------

def build_docx_table(out_docx_path: str, metrics: List[VideoMetrics], agg: Aggregate, corr_threshold: float):
    doc = Document()

    # Title
    title = doc.add_paragraph(f"{agg.label} Set Metrics")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(16)
    run.bold = True

    # Summary paragraph
    p = doc.add_paragraph()
    p.add_run("Summary\n").bold = True
    summary_lines = [
        f"Videos: {agg.num_videos}",
        f"Avg frames: {round(agg.avg_frames, 2) if agg.avg_frames is not None else '—'}",
        f"Avg transcript words: {round(agg.avg_transcript_words, 2) if agg.avg_transcript_words is not None else '—'}",
        f"Avg topics: {round(agg.avg_topics, 2) if agg.avg_topics is not None else '—'}",
        f"Avg comments: {round(agg.avg_comments, 2) if agg.avg_comments is not None else '—'}",
        f"Avg of per-video correlation averages: {round(agg.avg_corr_avg, 3) if agg.avg_corr_avg is not None else '—'}",
        f"Median of per-video correlation averages: {round(agg.median_corr_avg, 3) if agg.median_corr_avg is not None else '—'}",
        f"Threshold for counted correlations: >= {corr_threshold}",
    ]
    p.add_run("\n".join(summary_lines))

    # Table
    headers = [
        "video_id", "frames", "topics", "comments",
        "corr_avg", "corr_median", "corr_n", f"corr_n≥{corr_threshold}",
        "transcript_words", "duration_s"
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for m in metrics:
        row = table.add_row().cells
        row[0].text = m.video_id
        row[1].text = str(m.frames_count)
        row[2].text = str(m.topics_count)
        row[3].text = str(m.comments_count)
        row[4].text = "—" if m.corr_avg is None else f"{m.corr_avg:.3f}"
        row[5].text = "—" if m.corr_median is None else f"{m.corr_median:.3f}"
        row[6].text = str(m.corr_n)
        row[7].text = str(m.corr_n_ge_threshold)
        row[8].text = str(m.transcript_word_count)
        row[9].text = "—" if m.duration_seconds is None else f"{m.duration_seconds:.2f}"

    doc.save(out_docx_path)

# ---------------------- Main ----------------------

def main():
    ap = argparse.ArgumentParser(description="Collect metrics for embodied vs conventional video sets and export JSON + DOCX table")
    ap.add_argument('--videos_root', type=str, default="../rawvideos/downloaded_videos", help='Raw videos folder (optional; used for duration via ffprobe)')
    ap.add_argument('--frames_root', type=str, default="../frames", help='Frames root: <frames_root>/<video_id>/raw_frames')
    ap.add_argument('--transcripts_root', type=str, default="../transcription/transcripts", help='Transcripts root containing <video_id>.txt or .json')
    ap.add_argument('--topics_root', type=str, default="../correlation/topics", help='Topics root containing <video_id>.topics.json')
    ap.add_argument('--out_dir', type=str, default="../output/embodied", help='Output directory for JSON + DOCX')
    ap.add_argument('--label', type=str, default="Embodied", help='Label to tag this run (e.g., Conventional or Embodied)')
    ap.add_argument('--corr_threshold', type=float, default=60, help='Threshold for counting a correlation as present')

    args = ap.parse_args()

    os.makedirs(args.out_dir, exist_ok=True)

    metrics, agg = compute_metrics_for_set(
        videos_root=args.videos_root,
        frames_root=args.frames_root,
        transcripts_root=args.transcripts_root,
        topics_root=args.topics_root,
        label=args.label,
        corr_threshold=args.corr_threshold,
    )

    # Save JSON
    out_json = os.path.join(args.out_dir, f"{args.label}_metrics.json")
    blob = {
        'label': agg.label,
        'aggregate': asdict(agg),
        'corr_threshold': args.corr_threshold,
        'videos': [asdict(m) for m in metrics],
    }
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(blob, f, ensure_ascii=False, indent=2)

    # Save DOCX
    out_docx = os.path.join(args.out_dir, f"{args.label}_metrics.docx")
    build_docx_table(out_docx, metrics, agg, args.corr_threshold)

    print(f"Wrote: {out_json}")
    print(f"Wrote: {out_docx}")

    # Save CSVs
    build_csv_files(args.out_dir, args.label, metrics, agg, args.corr_threshold)


if __name__ == '__main__':
    main()
