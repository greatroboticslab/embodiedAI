#!/usr/bin/env python3
"""
generate_sections.py

Part 1 of Fusion Analysis.
1.  Aligns Transcript and Caption data.
2.  Uses LLM to segment video into "Engagement Sections".
3.  Outputs:
    - JSON: Sections data for correlation script.
    - DOCX: Human-readable report with section details, full transcripts, and representative frames.

Usage:
    python generate_sections.py \
        --transcripts_root ../data/transcripts_conventional \
        --captions_root ../../CaptionAnalysis/data/integrated_caption/frames_conventional_captions_integrated \
        --frames_root ../../VideoAnalysis/data/frames/frames_conventional \
        --metrics_csv ../../CaptionAnalysis/output/Conventional_metrics.csv \
        --output_dir ../results/fusion_analysis
"""

import os
import sys
import json
import csv
import re
import argparse
import math
from dataclasses import dataclass, asdict
from typing import List, Dict, Any, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("[ERR] python-docx not installed. Install with 'pip install python-docx' to generate reports.")
    sys.exit(1)

# Add local scripts dir to path for utils import
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
try:
    from config import Config
    from utils.llm_utils import generate_response, stream_parser
except ImportError:
    print("[WARN] Could not import utils.llm_utils. Ensure you are running from FusionAnalysis/scripts/")
    Config = None
    def generate_response(*args): raise NotImplementedError("LLM utils not found")
    def stream_parser(x): return x


# -------------------- Data Structures --------------------

@dataclass
class FusionEvent:
    timestamp: float
    end_timestamp: float
    source_type: str  # "AUDIO" or "VISUAL"
    content: str

@dataclass
class EngagementSection:
    id: int
    start_time: float
    end_time: float
    title: str
    summary: str
    visual_cues: str
    verbal_cues: str
    # We don't store comments here anymore, that's step 2
    
    def to_dict(self):
        return asdict(self)

# -------------------- Data Loading --------------------

def load_metrics_for_fps(metrics_csv_paths: List[str]) -> Dict[str, float]:
    """Load duration and frame counts to calculate FPS."""
    vid_meta = {} 
    
    for path in metrics_csv_paths:
        if not os.path.exists(path): continue
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                vid = row.get('video_id') or row.get('VideoID')
                if not vid: continue
                frames = row.get('frames_count') or row.get('frames') or row.get('total_frames')
                dur = row.get('duration_seconds') or row.get('duration') or row.get('DurationSec') or row.get('duration_s')
                
                if frames and dur:
                    try:
                        f_val = float(frames)
                        d_val = float(dur)
                        if vid not in vid_meta:
                            vid_meta[vid] = {'frames': f_val, 'duration': d_val}
                        else:
                            if f_val > 0: vid_meta[vid]['frames'] = f_val
                            if d_val > 0: vid_meta[vid]['duration'] = d_val
                    except ValueError: continue

    fps_map = {}
    for vid, data in vid_meta.items():
        if data['duration'] > 0 and data['frames'] > 0:
            fps_map[vid] = data['frames'] / data['duration']
        else:
            fps_map[vid] = 1.0 
    return fps_map

def load_transcript(path: str) -> List[Tuple[float, float, str]]:
    """Parses _timestamped.txt or .json. Returns [(start, end, text), ...]"""
    segments = []
    if path.endswith('.json'):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                segs = data.get('segments', [])
                for s in segs:
                    segments.append((float(s.get('start', 0)), float(s.get('end', 0)), s.get('text', '')))
        except Exception as e:
            print(f"[ERR] Failed to load transcript JSON {path}: {e}")
            
    elif path.endswith('.txt'):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                for line in f:
                    # Regex for [HH:MM:SS.mmm → HH:MM:SS.mmm]
                    m = re.search(r'\[(\d{2}):(\d{2}):(\d{2}\.\d+)\s+(?:→|->)\s+(\d{2}):(\d{2}):(\d{2}\.\d+)\]\s+(.*)', line)
                    if m:
                        h1, m1, s1 = int(m.group(1)), int(m.group(2)), float(m.group(3))
                        start_sec = h1*3600 + m1*60 + s1
                        h2, m2, s2 = int(m.group(4)), int(m.group(5)), float(m.group(6))
                        end_sec = h2*3600 + m2*60 + s2
                        content = m.group(7).strip()
                        segments.append((start_sec, end_sec, content))
        except Exception as e:
            print(f"[WARN] Text transcript parse error {path}: {e}")
    return segments

def load_captions(path: str, fps: float) -> List[Tuple[float, float, str]]:
    """Parses _captions_integrated.json."""
    captions = []
    if not os.path.exists(path): return captions
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            frames = data.get('frames', [])
            for item in frames:
                fid_str = item.get('frame_id', '')
                cols = re.findall(r'\d+', fid_str)
                if not cols: continue
                frame_idx = int(cols[-1])
                timestamp = frame_idx / fps
                captions.append((timestamp, timestamp + 2.0, item.get('caption', '')))
    except Exception as e:
        print(f"[ERR] Failed to load captions {path}: {e}")
    captions.sort(key=lambda x: x[0])
    return captions

# -------------------- Core Logic --------------------

SAFE_CHAR_LIMIT = 500000

def create_fusion_stream(transcripts: List[Tuple], captions: List[Tuple], caption_step: int = 1) -> List[FusionEvent]:
    events = []
    for s, e, txt in transcripts:
        events.append(FusionEvent(s, e, "AUDIO", txt))
    
    # Filter captions based on step
    for i, (s, e, txt) in enumerate(captions):
        if i % caption_step == 0:
            events.append(FusionEvent(s, e, "VISUAL", txt))
            
    events.sort(key=lambda x: x.timestamp)
    return events

def format_stream_for_llm(events: List[FusionEvent]) -> str:
    lines = []
    for ev in events:
        time_str = f"{int(ev.timestamp // 60):02d}:{int(ev.timestamp % 60):02d}"
        lines.append(f"[{time_str}] [{ev.source_type}] {ev.content}")
    return "\n".join(lines)

def segment_video(context_stream_text: str, duration_sec: float, model_name: str) -> List[EngagementSection]:
    """Segment video into sections, chunking if needed."""
    
    if len(context_stream_text) <= SAFE_CHAR_LIMIT:
        return _segment_chunk(context_stream_text, duration_sec, model_name)
    else:
        print(f"    [INFO] Stream length {len(context_stream_text)} chars exceeds limit. Chunking...")
        lines = context_stream_text.split('\n')
        chunks = []
        current_chunk = []
        current_len = 0
        
        for line in lines:
            if current_len + len(line) > SAFE_CHAR_LIMIT:
                chunks.append("\n".join(current_chunk))
                current_chunk = []
                current_len = 0
            current_chunk.append(line)
            current_len += len(line)
        if current_chunk:
            chunks.append("\n".join(current_chunk))
            
        all_sections = []
        offset_id = 1
        for i, chunk in enumerate(chunks):
            print(f"    > Processing chunk {i+1}/{len(chunks)}...")
            chunk_sections = _segment_chunk(chunk, duration_sec, model_name)
            for sec in chunk_sections:
                sec.id = offset_id
                offset_id += 1
                all_sections.append(sec)
        return all_sections

def _segment_chunk(text: str, duration_sec: float, model_name: str) -> List[EngagementSection]:
    prompt = f"""
    You are an expert AI Video Analyst. 
    Below is a chronological stream of a video's content (AUDIO + VISUAL).
    
    YOUR TASK:
    Analyze this stream and divide it into distinct "Engagement Sections" based on topic shifts, activity changes, or narrative progression.
    
    GUIDELINES:
    1. **Coherence**: Each section should represent a cohesive segment of the video (e.g., an intro, a specific demo, a Q&A session, a conclusion).
    2. **Granularity**: Avoid making sections too short (<30s) unless strictly necessary. Aim for standard video chapters.
    3. **Titles**: Create descriptive, engaging titles (e.g., "Robot Demonstration: Picking up the Cup" instead of just "Demo").
    4. **Summary**: Synthesize both what is seen (VISUAL) and what is said (AUDIO).
    5. **Cues**: explicitely list the key visuals and words that define the section.
    
    INPUT STREAM:
    {text}

    OUTPUT FORMAT (Strict JSON list):
    [
      {{
        "start_time": 0.0,
        "end_time": 45.5,
        "title": "Descriptive Title",
        "summary": "Comprehensive summary of the section's content.",
        "visual_cues": "Key actions or objects seen.",
        "verbal_cues": "Key phrases or topics discussed."
      }}
    ]
    Response must be ONLY valid JSON.
    """
    try:
        stream = generate_response(model_name, prompt)
        full_text = "".join(stream_parser(stream)).strip()
        
        # Look for the first JSON list [...] 
        # We use non-greedy matching .*? inside but wait, we need to match nested brackets potentially?
        # A simple non-greedy match might fail on nested structures, but for this strict format 
        # it usually works. Let's try to find the first '[' and the last ']' that makes valid JSON.
        
        # Robust strategy: Find start '[', then Try to parse from there.
        start_idx = full_text.find('[')
        if start_idx != -1:
            potential_json = full_text[start_idx:]
            # Try to find the matching closing bracket by parsing?
            # Or just use the original regex which was strict: r'\[.*\]' with DOTALL matches EVERYTHING from first [ to LAST ]
            # If the LLM adds text AFTER the last ], the original regex `r'\[.*\]'` matches it if inside.
            # But if it output: [ ... ] \n Hope that helps!
            # The regex `\[.*\]` (greedy) will match up to the LAST ] in "Hope that helps!", which doesn't exist. 
            # So it matches up to the actual last ]. 
            
            # Wait, `Extra data` error usually means we passed "[...] some garbage" to json.loads.
            # json.loads only accepts ONE valid JSON object. 
            # If `json_match.group(0)` captured "[...] \n text", then json.loads fails.
            
            # Fix: Use a regex that stops at the closing bracket of the list.
            # Since we can't easily balance brackets with regex, let's just clean up the `full_text` 
            # to extract the largest substring starting with [ and ending with ] that parses.
            
            # Simplified approach for this specific error "Extra data":
            # The regex `\[.*\]` is greedy. If there is text like "Here is the list: [...] Some notes",
            # it matches "[...]". Wait, regex `\[.*\]` will match from first [ to LAST ]. 
            # IF the text is "[...]\nNotes", the last ] is the end of the list. 
            # So regex SHOULD extract "[...]". 
            
            # However, if the regex is `\[.*\]` and the text is `[...]`, json.loads works.
            # If text is `[...] \n Text`, regex `\[.*\]` matches `[...]`. 
            # BUT if text is `[...] ... [other thing]`, it matches the whole span.
            
            # The error "Extra data: line X col Y" specifically comes from `json.loads` when the string HAS valid json but continues.
            # This suggests my regex might be capturing too much or `json_match.group(0)` is NOT doing what I think.
            # Let's verify standard `re.search(r'\[.*\]', text, flags=re.DOTALL)`.
            # If text = "[A]\nB", match is "[A]". json.loads("[A]") -> OK.
            
            # Ah, maybe the LLM output something like:  
            # ```json
            # [...]
            # ```
            # My current code uses regex. 
            # Let's try to be safer: parse only the valid JSON part.
            
            # Fallback: Just assume the list ends at the last ']' found.
            end_idx = potential_json.rfind(']')
            if end_idx != -1:
                json_str = potential_json[:end_idx+1]
                data = json.loads(json_str)
                
                sections = []
                for item in data:
                    sections.append(EngagementSection(
                        id=0,
                        start_time=float(item.get('start_time', 0)),
                        end_time=float(item.get('end_time', 0)),
                        title=item.get('title', 'Unknown'),
                        summary=item.get('summary', ''),
                        visual_cues=item.get('visual_cues', ''),
                        verbal_cues=item.get('verbal_cues', '')
                    ))
                return sections
    except Exception as e:
        print(f"    [ERR] Chunk processing failed: {e}")
        # Debug: print the bad text to see what happened
        # print(f"DEBUG TEXT: {full_text}") 
    return []

# -------------------- DOCX Report Generation --------------------

def find_frame_path(video_frames_dir: str, timestamp: float, fps: float) -> Optional[str]:
    """Find the frame image file closest to the given timestamp."""
    target_frame_idx = int(timestamp * fps)
    
    # Try exact match first
    # Filename format: frame_001.jpg or frame_1.jpg?
    # Usually frame_{number}.jpg
    # Check what files exist
    if not os.path.exists(video_frames_dir): return None
    
    # We scan directory files once per report? Or just construct path?
    # Constructing path is faster if naming is consistent.
    # Assuming "frame_X.jpg"
    cand = os.path.join(video_frames_dir, f"frame_{target_frame_idx}.jpg")
    if os.path.exists(cand): return cand
    
    # Check leading zeros?
    cand = os.path.join(video_frames_dir, f"frame_{target_frame_idx:04d}.jpg") # just in case
    if os.path.exists(cand): return cand

    # Fallback: list dir and find closest number
    # This is expensive, maybe do it if exact fails
    files = [f for f in os.listdir(video_frames_dir) if f.endswith('.jpg')]
    best_diff = float('inf')
    best_file = None
    
    for f in files:
        nums = re.findall(r'\d+', f)
        if nums:
            f_idx = int(nums[-1])
            diff = abs(f_idx - target_frame_idx)
            if diff < best_diff:
                best_diff = diff
                best_file = f
                
    if best_file and best_diff < (5 * fps): # Within 5 seconds?
        return os.path.join(video_frames_dir, best_file)
        
    return None

def generate_docx_report(video_id: str, sections: List[EngagementSection], transcripts: List[Tuple], 
                         video_frames_dir: str, fps: float, out_path: str):
    doc = Document()
    doc.add_heading(f"Engagement Analysis: {video_id}", 0)
    
    for sec in sections:
        doc.add_heading(f"Section {sec.id}: {sec.title}", level=1)
        doc.add_paragraph(f"Time: {int(sec.start_time // 60)}:{int(sec.start_time % 60):02d} - {int(sec.end_time // 60)}:{int(sec.end_time % 60):02d}")
        
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(sec.summary)
        
        # Representative Frames
        doc.add_heading("Representative Frames", level=2)
        # Select 5 timestamps evenly spaced
        step = (sec.end_time - sec.start_time) / 6
        timestamps = [sec.start_time + step * (i+1) for i in range(5)]
        
        # Create a table for images
        table = doc.add_table(rows=1, cols=5)
        table.autofit = True
        
        for i, ts in enumerate(timestamps):
            cell = table.rows[0].cells[i]
            # Add text timestamp
            p = cell.paragraphs[0]
            p.add_run(f"{int(ts // 60)}:{int(ts % 60):02d}").bold = True
            
            # Find and add image
            img_path = find_frame_path(video_frames_dir, ts, fps)
            if img_path:
                try:
                    run = p.add_run()
                    run.add_break()
                    run.add_picture(img_path, width=Inches(1.2)) 
                except Exception as e:
                    p.add_run("\n[Image Error]")
            else:
                p.add_run("\n[No Frame]")
                
        # Full Transcript
        doc.add_heading("Transcript", level=2)
        # Filter transcript segments
        sec_text = []
        for s, e, txt in transcripts:
            # Check overlap: (StartA <= EndB) and (EndA >= StartB)
            # Strictly inside? Or overlapping? Let's say midpoint is inside
            mid = (s + e) / 2
            if sec.start_time <= mid <= sec.end_time:
                sec_text.append(f"[{int(s // 60)}:{int(s % 60):02d}] {txt}")
                
        if sec_text:
            doc.add_paragraph("\n".join(sec_text))
        else:
            doc.add_paragraph("(No speech detected in this section)")
            
        doc.add_page_break()
        
    doc.save(out_path)

# -------------------- Main --------------------

def main():
    parser = argparse.ArgumentParser(description="Fusion Analysis - Step 1: Segmentation")
    parser.add_argument("--transcripts_root", default="../../TranscriptAnalysis/data/transcripts_embodied", help="Directory containing <video_id>.transcripts.json files.")
    parser.add_argument("--captions_root", default="../../CaptionAnalysis/data/integrated_caption/frames_embodied_captions_integrated")
    parser.add_argument("--frames_root", default="../../VideoAnalysis/data/frames/frames_embodied", help="Root dir containing <video_id>/raw_frames or similar")
    parser.add_argument("--metrics_csv", default=["../../CaptionAnalysis/output/Embodied_metrics.csv"], nargs='+', help="One or more CSVs to load FPS from")
    parser.add_argument("--output_dir", default="../results/fusion_analysis")
    parser.add_argument("--model", default="llama3.1", help="LLM model name")
    
    args = parser.parse_args()
    
    fps_map = load_metrics_for_fps(args.metrics_csv)
    print(f"[INFO] Loaded FPS data for {len(fps_map)} videos.")
    
    # Discover Videos
    vid_ids = []
    for f in os.listdir(args.transcripts_root):
        vid = None
        if f.endswith('_timestamped.txt'):
            vid = f.replace('_timestamped.txt', '')
        elif f.endswith('.json') and not f.endswith('topics.json'):
            vid = f.replace('.json', '')
        if vid:
            # Check if caption file exists to confirm it's a valid pair
            if os.path.exists(os.path.join(args.captions_root, vid, f"{vid}_captions_integrated.json")):
                vid_ids.append(vid)
                
    vid_ids = sorted(list(set(vid_ids)))
    print(f"[INFO] Found {len(vid_ids)} videos to process.")
    
    # Sections output subfolder
    sections_dir = os.path.join(args.output_dir, "sections")
    if not os.path.exists(sections_dir):
        os.makedirs(sections_dir)
        
    for vid in vid_ids:
        # Check if output already exists
        json_out = os.path.join(sections_dir, f"{vid}_sections.json")
        if os.path.exists(json_out):
            print(f"[INFO] Skipping {vid}, output already exists.")
            continue

        print(f"\nProcessing {vid}...")
        
        # Load Data
        t_path = os.path.join(args.transcripts_root, f"{vid}_timestamped.txt")
        if not os.path.exists(t_path): t_path = os.path.join(args.transcripts_root, f"{vid}.json")
        c_path = os.path.join(args.captions_root, vid, f"{vid}_captions_integrated.json")
        
        fps = fps_map.get(vid, 1.0)
        transcripts = load_transcript(t_path)
        captions = load_captions(c_path, fps)
        
        max_t = 0
        if transcripts: max_t = max(max_t, transcripts[-1][1])
        if captions: max_t = max(max_t, captions[-1][1])
        if max_t == 0: max_t = 60
        
        # Fusion & Filtering
        # Iteratively increase filtering step if text is too long
        cap_step = 1
        stream = []
        stream_text = ""
        
        while True:
            stream = create_fusion_stream(transcripts, captions, caption_step=cap_step)
            stream_text = format_stream_for_llm(stream)
            
            if len(stream_text) <= SAFE_CHAR_LIMIT:
                break
            
            # If we are filtering significantly, check if we should stop
            # Let's limit filtering to say 1/20 (~30s timestamps), otherwise we just fallback to chunking
            if cap_step >= 20: 
                print(f"  [WARN] Stream still exceeds {SAFE_CHAR_LIMIT} chars even at 1/{cap_step} captions. Proceeding with chunking.")
                break
                
            cap_step += 1
            
        if cap_step > 1:
            print(f"  > Filtered captions: taking 1 out of every {cap_step} (Reduced size to {len(stream_text)})")
                
        print(f"  > Segmenting ({len(stream)} events)...")
        sections = segment_video(stream_text, max_t, args.model)
        
        # 1. Save JSON
        json_out = os.path.join(sections_dir, f"{vid}_sections.json")
        result = {
            "video_id": vid,
            "duration": max_t,
            "sections": [s.to_dict() for s in sections]
        }
        with open(json_out, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2)
        
        # 2. Generate DOCX
        # Find frames dir for this video.
        # User passes frames_root which might contain <vid>/raw_frames or just <vid>
        # Let's try standard patterns
        v_frames = os.path.join(args.frames_root, vid, "raw_frames")
        if not os.path.isdir(v_frames):
             v_frames = os.path.join(args.frames_root, vid) # Direct
        
        docx_out = os.path.join(sections_dir, f"{vid}_sections.docx")
        print(f"  > Generating Report with frames from {v_frames}...")
        generate_docx_report(vid, sections, transcripts, v_frames, fps, docx_out)
        
        print(f"  > Saved JSON and DOCX to {sections_dir}")

if __name__ == "__main__":
    main()
