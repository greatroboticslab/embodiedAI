#!/usr/bin/env python3
"""
Collect per-video metrics for a Caption Analysis study (Embodied vs Conventional),
save a JSON summary, and generate a DOCX table report.

Inputs:
  --results_root              Path to the folder holding correlation results (*.results.jsonl)
                              Structure can be flat or recursively mirrored.
  --integrated_captions_root  Path to the folder holding integrated captions (*_captions_integrated.json).
  --videos_root               (Optional) Path to raw videos for duration calculation.
  --out_dir                   Output directory to write results JSON and DOCX.

Example:
  python collect_results.py \
      --results_root ../results \
      --integrated_captions_root ../data/integrated_caption \
      --out_dir ../output/conventional \
      --label Conventional
"""

import csv
import argparse
import json
import os
import re
import subprocess
import glob
from dataclasses import dataclass, asdict
from statistics import mean, median
from typing import Dict, List, Optional, Tuple, Any, Set

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    raise SystemExit("python-docx is required. Install with: pip install python-docx")

# ---------------------- Utilities ----------------------

def _none(x):
    return "" if x is None else x

def list_video_ids_from_results(results_root: str) -> List[Tuple[str, str]]:
    """
    Scan results_root for *_correlation.results.jsonl.
    Returns list of (video_id, absolute_path_to_jsonl).
    """
    if not os.path.isdir(results_root):
        return []
    
    found = []
    # Recursive search for .results.jsonl
    for root, dirs, files in os.walk(results_root):
        for fn in files:
            if fn.endswith("_correlation.results.jsonl"):
                # filename format: <video_id>_correlation.results.jsonl
                # robust extraction:
                suffix = "_correlation.results.jsonl"
                vid = fn[:-len(suffix)]
                found.append((vid, os.path.join(root, fn)))
    
    # Sort by video_id
    found.sort(key=lambda x: x[0])
    return found

def find_integrated_caption_file(captions_root: str, vid: str) -> Optional[str]:
    """
    Search recursively for <vid>_captions_integrated.json in captions_root.
    """
    target = f"{vid}_captions_integrated.json"
    for root, dirs, files in os.walk(captions_root):
        if target in files:
            return os.path.join(root, target)
    return None

def get_caption_metrics(json_path: str) -> Tuple[int, int]:
    """
    Returns (frames_count, caption_word_count).
    """
    if not json_path or not os.path.isfile(json_path):
        return 0, 0
    
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return 0, 0

    frames = data.get("frames", [])
    if not isinstance(frames, list):
        return 0, 0

    f_count = len(frames)
    word_count = 0
    for item in frames:
        if isinstance(item, dict):
            cap = item.get("caption", "")
            if isinstance(cap, str):
                # Simple word count
                word_count += len(re.findall(r"\b\w+\b", cap))
    
    return f_count, word_count

def parse_correlation_jsonl(path: str) -> Tuple[int, List[float]]:
    """
    Returns (comments_count, scores[]).
    Each line in JSONL is a frame record with "candidates" list.
    We verify unique comments by "comment" text.
    """
    if not path or not os.path.isfile(path):
        return 0, []

    comments_seen: Set[str] = set()
    scores: List[float] = []

    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except Exception:
                    continue

                cands = obj.get("candidates")
                if isinstance(cands, list):
                    for c in cands:
                        if not isinstance(c, dict):
                            continue
                        
                        # Extract score
                        sc = c.get("score")
                        if isinstance(sc, (int, float)):
                            scores.append(float(sc))
                        
                        # Extract unique comment
                        # We use the text itself to identify unique comments
                        txt = c.get("comment", "")
                        if txt and isinstance(txt, str) and txt.strip():
                            comments_seen.add(txt.strip())

    except Exception:
        pass

    return len(comments_seen), scores

def get_duration_seconds(videos_root: Optional[str], vid: str) -> Optional[float]:
    if not videos_root or not os.path.isdir(videos_root):
        return None
    
    def _find_file_for_id(target_id):
        def _scan(d):
            try:
                for fn in os.listdir(d):
                    if fn.startswith(target_id):
                        if os.path.splitext(fn)[0] == target_id and os.path.splitext(fn)[1].lower() in {'.mp4', '.mkv', '.mov', '.avi', '.webm'}:
                            return os.path.join(d, fn)
            except OSError:
                pass
            return None

        # Check root
        cand = _scan(videos_root)
        if cand: return cand

        # Check subdirectories (1 level deep)
        for subd in os.listdir(videos_root):
            subp = os.path.join(videos_root, subd)
            if os.path.isdir(subp):
                c = _scan(subp)
                if c: return c
        return None

    cand = _find_file_for_id(vid)
    
    # Fallback: if ID is longer than 11 chars (e.g. has suffix), try first 11 chars
    if not cand and len(vid) > 11:
        cand = _find_file_for_id(vid[:11])

    if not cand:
        return None

    try:
        proc = subprocess.run([
            'ffprobe', '-v', 'error', '-select_streams', 'v:0',
            '-show_entries', 'format=duration', '-of', 'default=nw=1:nk=1', cand
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=10)
        out = proc.stdout.strip()
        if out:
            return float(out)
    except Exception:
        return None
    return None

# ---------------------- Data Classes ----------------------

@dataclass
class VideoMetrics:
    video_id: str
    frames_count: int
    caption_word_count: int
    comments_count: int
    corr_avg: Optional[float]
    corr_median: Optional[float]
    corr_n: int
    corr_n_ge_threshold: int
    duration_seconds: Optional[float]

@dataclass
class Aggregate:
    label: str
    num_videos: int
    avg_frames: Optional[float]
    avg_caption_words: Optional[float]
    avg_comments: Optional[float]
    avg_corr_avg: Optional[float]
    median_corr_avg: Optional[float]

# ---------------------- Core ----------------------

def build_csv_files(out_dir: str, label: str, metrics: List[VideoMetrics], agg: Aggregate, corr_threshold: float):
    # --- Per-video CSV ---
    headers = [
        "video_id", "frames", "caption_words", "comments",
        "corr_avg", "corr_median", "corr_n", f"corr_n_ge_{corr_threshold}",
        "duration_s"
    ]
    per_video_csv = os.path.join(out_dir, f"{label}_metrics.csv")
    with open(per_video_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(headers)
        for m in metrics:
            w.writerow([
                m.video_id,
                m.frames_count,
                m.caption_word_count,
                m.comments_count,
                _none(m.corr_avg),
                _none(m.corr_median),
                m.corr_n,
                m.corr_n_ge_threshold,
                _none(m.duration_seconds),
            ])

    # --- Aggregate CSV ---
    agg_headers = [
        "label", "num_videos", "avg_frames", "avg_caption_words",
        "avg_comments", "avg_corr_avg", "median_corr_avg",
        "corr_threshold"
    ]
    agg_csv = os.path.join(out_dir, f"{label}_aggregate.csv")
    with open(agg_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(agg_headers)
        w.writerow([
            agg.label,
            agg.num_videos,
            _none(agg.avg_frames),
            _none(agg.avg_caption_words),
            _none(agg.avg_comments),
            _none(agg.avg_corr_avg),
            _none(agg.median_corr_avg),
            corr_threshold,
        ])
    
    print(f"Wrote: {per_video_csv}")
    print(f"Wrote: {agg_csv}")

def compute_metrics_for_set(
    videos_root: Optional[str],
    results_root: str,
    integrated_captions_root: str,
    label: str,
    corr_threshold: float = 60.0,
) -> Tuple[List[VideoMetrics], Aggregate]:

    # 1. Discover videos from results
    video_list = list_video_ids_from_results(results_root)
    
    metrics: List[VideoMetrics] = []
    corr_avgs: List[float] = []

    for vid, res_path in video_list:
        # Metrics from captions
        cap_path = find_integrated_caption_file(integrated_captions_root, vid)
        frames_count, caption_words = get_caption_metrics(cap_path)

        # Metrics from results
        comments_count, scores = parse_correlation_jsonl(res_path)

        corr_avg = mean(scores) if scores else None
        corr_median = median(scores) if scores else None
        corr_n = len(scores)
        corr_n_ge = sum(1 for s in scores if s >= corr_threshold)
        
        dur_s = get_duration_seconds(videos_root, vid)

        vm = VideoMetrics(
            video_id=vid,
            frames_count=frames_count,
            caption_word_count=caption_words,
            comments_count=comments_count,
            corr_avg=corr_avg,
            corr_median=corr_median,
            corr_n=corr_n,
            corr_n_ge_threshold=corr_n_ge,
            duration_seconds=dur_s,
        )
        metrics.append(vm)
        if corr_avg is not None:
            corr_avgs.append(corr_avg)
            
    def _avg(vals: List[int]) -> Optional[float]:
        return mean(vals) if vals else None

    # Aggregate
    agg = Aggregate(
        label=label,
        num_videos=len(metrics),
        avg_frames=_avg([m.frames_count for m in metrics]),
        avg_caption_words=_avg([m.caption_word_count for m in metrics]),
        avg_comments=_avg([m.comments_count for m in metrics]),
        avg_corr_avg=mean(corr_avgs) if corr_avgs else None,
        median_corr_avg=median(corr_avgs) if corr_avgs else None,
    )

    return metrics, agg

# ---------------------- DOCX ----------------------

def build_docx_table(out_docx_path: str, metrics: List[VideoMetrics], agg: Aggregate, corr_threshold: float):
    doc = Document()

    # Title
    title = doc.add_paragraph(f"{agg.label} Caption Metrics")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(16)
    run.bold = True

    # Summary
    p = doc.add_paragraph()
    p.add_run("Summary\n").bold = True
    summary_lines = [
        f"Videos: {agg.num_videos}",
        f"Avg frames: {round(agg.avg_frames, 2) if agg.avg_frames is not None else '—'}",
        f"Avg caption words: {round(agg.avg_caption_words, 2) if agg.avg_caption_words is not None else '—'}",
        f"Avg distinct comments: {round(agg.avg_comments, 2) if agg.avg_comments is not None else '—'}",
        f"Avg of per-video correlation averages: {round(agg.avg_corr_avg, 3) if agg.avg_corr_avg is not None else '—'}",
        f"Median of per-video correlation averages: {round(agg.median_corr_avg, 3) if agg.median_corr_avg is not None else '—'}",
        f"Threshold for counted correlations: >= {corr_threshold}",
    ]
    p.add_run("\n".join(summary_lines))

    # Table
    headers = [
        "video_id", "frames", "caption_words", "comments",
        "corr_avg", "corr_median", "corr_n", f"corr_n≥{corr_threshold}",
        "duration_s"
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for m in metrics:
        row = table.add_row().cells
        row[0].text = m.video_id
        row[1].text = str(m.frames_count)
        row[2].text = str(m.caption_word_count)
        row[3].text = str(m.comments_count)
        row[4].text = "—" if m.corr_avg is None else f"{m.corr_avg:.3f}"
        row[5].text = "—" if m.corr_median is None else f"{m.corr_median:.3f}"
        row[6].text = str(m.corr_n)
        row[7].text = str(m.corr_n_ge_threshold)
        row[8].text = "—" if m.duration_seconds is None else f"{m.duration_seconds:.2f}"

    doc.save(out_docx_path)

# ---------------------- Main ----------------------

def main():
    ap = argparse.ArgumentParser(description="Collect metrics for Caption Analysis results (Embodied vs Conventional)")
    ap.add_argument('--results_root', type=str, default="../results/frames_embodied", help='Root directory containing *_correlation.results.jsonl files')
    ap.add_argument('--integrated_captions_root', type=str, default="../data/integrated_caption/frames_embodied_captions_integrated", help='Root directory containing *_captions_integrated.json files')
    ap.add_argument('--videos_root', type=str, default="../../VideoAnalysis/rawvideos/embodied_videos", help='Raw videos folder (optional; used for duration via ffprobe)')
    ap.add_argument('--out_dir', type=str, default="../output", help='Output directory for JSON + DOCX')
    ap.add_argument('--label', type=str, default="Embodied", help='Label to tag this run (e.g., Conventional or Embodied)')
    ap.add_argument('--corr_threshold', type=float, default=60, help='Threshold for counting a correlation as present')

    args = ap.parse_args()

    os.makedirs(args.out_dir, exist_ok=True)

    metrics, agg = compute_metrics_for_set(
        videos_root=args.videos_root,
        results_root=args.results_root,
        integrated_captions_root=args.integrated_captions_root,
        label=args.label,
        corr_threshold=args.corr_threshold,
    )

    # Save JSON
    out_json = os.path.join(args.out_dir, f"{args.label}_metrics.json")
    blob = {
        'label': agg.label,
        'aggregate': asdict(agg),
        'corr_threshold': args.corr_threshold,
        'videos': [asdict(m) for m in metrics],
    }
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(blob, f, ensure_ascii=False, indent=2)

    # Save DOCX
    out_docx = os.path.join(args.out_dir, f"{args.label}_metrics.docx")
    build_docx_table(out_docx, metrics, agg, args.corr_threshold)

    print(f"Wrote: {out_json}")
    print(f"Wrote: {out_docx}")

    # Save CSVs
    build_csv_files(args.out_dir, args.label, metrics, agg, args.corr_threshold)

if __name__ == '__main__':
    main()
