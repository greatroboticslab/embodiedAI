import os
from config import Config
from utils.llm_utils import generate_image_desc, stream_parser, generate_response
from docx import Document

def extract_captioned_frames_from_docx(docx_path):
    captioned = set()
    try:
        doc = Document(docx_path)
        paragraphs = doc.paragraphs
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i]
            if para.style.name.startswith("Heading 1"):
                heading = para.text.strip()
                if heading.lower().startswith("frame:"):
                    frame = heading.split("Frame:", 1)[-1].strip()
                else:
                    frame = heading
                frame = frame.lower()

                # Look ahead for Description and Caption
                found_desc = False
                found_caption = False
                for j in range(i+1, min(i+6, len(paragraphs))):
                    text = paragraphs[j].text.strip()
                    if text.startswith("Description:"):
                        found_desc = True
                    if text.startswith("Caption:"):
                        found_caption = True
                    if found_desc and found_caption:
                        captioned.add(frame)
                        break
            i += 1
    except Exception as e:
        print(f"⚠️ Failed to read existing docx: {docx_path} — {e}")
    return captioned

def caption_images_in_folder(folder_path, llava_model, llama_model, docx_output_path):
    image_files = sorted([
        f for f in os.listdir(folder_path)
        if f.lower().endswith(('.jpg', '.jpeg', '.png'))
    ])
    if not image_files:
        return None

    folder_name = os.path.basename(folder_path.rstrip('/'))
    captioned = set()
    if os.path.exists(docx_output_path):
        print(f"🔁 Resuming from existing file: {docx_output_path}")
        doc = Document(docx_output_path)
        captioned = extract_captioned_frames_from_docx(docx_output_path)
    else:
        doc = Document()
        doc.add_heading(f"Image Caption Report: {folder_name}", 0)

    processed_count = 0
    for image_file in image_files:
        if image_file.lower() in captioned:
            print(f"⏭️ Skipping {image_file} (already captioned)")
            continue

        image_path = os.path.join(folder_path, image_file)
        print(f"🖼️ Processing {image_file}")

        try:
            stream = generate_image_desc(image_path, llava_model, "What do you see in this image?")
            description = ''.join(stream_parser(stream))
        except Exception as e:
            description = f"[Error: {str(e)}]"
            print(f"    ❌ Description failed: {description}")

        if description.startswith("[Error:") or description.startswith("[No content received"):
            caption = "[No caption generated due to description error]"
        else:
            try:
                prompt = f"Please generate a caption for this image using this description: {description}"
                stream = generate_response(llama_model, prompt)
                caption = ''.join(stream_parser(stream))
            except Exception as e:
                caption = f"[Error generating caption: {str(e)}]"

        print(f"    ✅ Description: {description}")
        print(f"    ✅ Caption: {caption}")

        doc.add_heading(f"Frame: {image_file}", level=1)
        doc.add_paragraph(f"Description: {description}")
        doc.add_paragraph(f"Caption: {caption}")
        doc.add_paragraph('')

        try:
            doc.save(docx_output_path)
        except Exception as e:
            print(f"⚠️ Failed to save after {image_file}: {e}")
            break

        processed_count += 1

    print(f"✅ Finished {processed_count} new frames. Saved to: {docx_output_path}")
    return doc

def generate_combined_and_individual_docs(root_folder: str, combined_filename: str):
    llava_model = Config.ollama_models[1]
    llama_model = Config.ollama_models[0]
    combined_doc = Document()
    combined_doc.add_heading('Combined Image Caption Report', 0)

    for subdir, _, files in os.walk(root_folder):
        image_files = sorted([f for f in files if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
        if not image_files:
            continue

        folder_name = os.path.relpath(subdir, root_folder)
        print(f"\n📂 Processing folder: {folder_name}")
        output_docx_name = f"{folder_name.replace(os.sep, '_')}_captions_llava.docx"
        output_docx_path = os.path.join(subdir, output_docx_name)

        existing_captions = extract_captioned_frames_from_docx(output_docx_path) if os.path.exists(output_docx_path) else set()
        image_set = {img.lower().strip() for img in image_files}
        missing_images = image_set - existing_captions

        print("🧾 Images in folder:", len(image_set))
        print("📝 Captions in docx:", len(existing_captions))
        print("🔍 Missing frames:", missing_images)

        if missing_images:
            individual_doc = caption_images_in_folder(subdir, llava_model, llama_model, output_docx_path)
        else:
            print(f"✅ All frames already captioned in {output_docx_name}")
            individual_doc = Document(output_docx_path)

        if individual_doc:
            combined_doc.add_page_break()
            combined_doc.add_heading(f"Folder: {folder_name}", level=1)
            for element in individual_doc.element.body:
                combined_doc.element.body.append(element)

    combined_doc_path = os.path.join(root_folder, f"{combined_filename}.docx")
    combined_doc.save(combined_doc_path)
    print(f"\n📄 Combined document saved as {combined_doc_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Generate captions per folder and a combined docx.")
    parser.add_argument("root_folder", help="Root folder containing subfolders of frames")
    parser.add_argument("combined_docx_filename", help="Filename for the combined .docx (no extension)")
    args = parser.parse_args()
    generate_combined_and_individual_docs(args.root_folder, args.combined_docx_filename)
