import os
import json
import argparse
from config import Config
from utils.llm_utils import generate_image_desc, stream_parser
from docx import Document
from docx.shared import Inches

VALID_EXTS = ('.jpg', '.jpeg', '.png')


def extract_captioned_frames_from_json(json_path):
    """Extract set of frame_ids that have already been captioned."""
    captioned = set()
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            for frame in data.get('frames', []):
                captioned.add(frame['frame_id'].lower())
    except Exception as e:
        print(f"⚠️ Failed to read existing json: {json_path} — {e}")
    return captioned


def infer_video_id_from_subdir(subdir):
    """
    If frames are under .../<video_id>/raw_frames or similar,
    use the parent folder as video_id. Otherwise use the basename.
    """
    base = os.path.basename(subdir.rstrip(os.sep))
    if base.lower() in {'raw_frames', 'frames', 'images'}:
        return os.path.basename(os.path.dirname(subdir.rstrip(os.sep)))
    return base


def caption_folder_to_docx_and_json(
        frames_folder,
        llava_model,
        out_docx_path,
        out_json_path,
        image_width_inches=5.0
):
    """
    Create a DOCX with image previews and captions for all frames in frames_folder,
    and a JSON file with (video_id, frame_id, caption) for downstream processing.
    Supports resume functionality by checking existing JSON file.
    """
    image_files = sorted(
        f for f in os.listdir(frames_folder)
        if f.lower().endswith(VALID_EXTS)
    )
    if not image_files:
        return 0

    # Check for existing captions to resume
    captioned = set()
    existing_data = {"video_id": infer_video_id_from_subdir(frames_folder), "frames": []}

    if os.path.exists(out_json_path):
        print(f"🔁 Resuming from existing file: {out_json_path}")
        captioned = extract_captioned_frames_from_json(out_json_path)
        try:
            with open(out_json_path, 'r', encoding='utf-8') as f:
                existing_data = json.load(f)
        except Exception as e:
            print(f"⚠️ Could not load existing JSON, starting fresh: {e}")
            existing_data = {"video_id": infer_video_id_from_subdir(frames_folder), "frames": []}

    # Prepare DOCX
    if os.path.exists(out_docx_path):
        try:
            doc = Document(out_docx_path)
            print(f"📄 Loaded existing DOCX: {out_docx_path}")
        except Exception as e:
            print(f"⚠️ Existing DOCX is corrupted, creating new one: {e}")
            doc = Document()
            folder_name = os.path.basename(frames_folder.rstrip(os.sep))
            doc.add_heading(f"Captions for {folder_name}", level=0)
    else:
        doc = Document()
        folder_name = os.path.basename(frames_folder.rstrip(os.sep))
        doc.add_heading(f"Captions for {folder_name}", level=0)

    # Prepare JSON data
    video_id = existing_data["video_id"]
    json_payload = {
        "video_id": video_id,
        "frames": existing_data["frames"].copy()
    }

    processed_count = 0
    for filename in image_files:
        frame_id = os.path.splitext(filename)[0]

        # Skip if already captioned
        if frame_id.lower() in captioned:
            print(f"⏭️ Skipping {filename} (already captioned)")
            continue

        image_path = os.path.join(frames_folder, filename)
        print(f"🖼️ Processing {filename}")

        # Frame header
        doc.add_heading(filename, level=1)

        # Insert the actual frame image
        try:
            doc.add_picture(image_path, width=Inches(image_width_inches))
        except Exception as e:
            doc.add_paragraph(f"[Error inserting image: {e}]")
            caption = f"[Error inserting image: {e}]"
        else:
            # Generate caption using LLaVa
            try:
                prompt = (
                    "Describe what is visible in this frame in detail. "
                    "Focus on the visual content and teaching context — such as the person, devices, slides, diagrams, or text shown. "
                    "Mention if the instructor is writing, pointing, holding a model, or speaking from a slide. "
                    "Include any readable text, equations, or graph types visible on the screen. "
                    "Do not infer topics or intentions beyond what can be seen."
                )

                stream = generate_image_desc(image_path, llava_model, prompt)
                caption = ''.join(stream_parser(stream))

                # Check if caption is an error message
                if caption.startswith("[Error:") or caption.startswith("[No content received"):
                    print(f"    ❌ Caption generation issue: {caption}")
            except Exception as e:
                caption = f"[Error generating caption: {str(e)}]"
                print(f"    ❌ Caption failed: {caption}")

        # Add caption to DOCX
        doc.add_paragraph(caption)
        print(f"    ✅ Caption: {caption}")

        # Append to JSON list
        json_payload["frames"].append({
            "frame_id": frame_id,
            "caption": caption
        })

        # Save after each image (for resume capability)
        try:
            os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)
            doc.save(out_docx_path)
            with open(out_json_path, "w", encoding="utf-8") as f:
                json.dump(json_payload, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"⚠️ Failed to save after {filename}: {e}")
            break

        processed_count += 1

    print(f"  ✅ Saved caption .docx: {out_docx_path}")
    print(f"  ✅ Saved caption .json: {out_json_path}")

    return processed_count


def caption_all_folders(src_dir, out_dir):
    """
    Walk src_dir for subfolders containing images.
    For each such subfolder, produce:
      DOCX: <out_dir>/<basename(src_dir)>_captions_llava/<video_id>/<video_id>_captions_llava.docx
      JSON: <out_dir>/<basename(src_dir)>_captions_llava/<video_id>/<video_id>_captions_llava.json
    """
    llava_model = Config.ollama_models[1]

    src_dir = os.path.abspath(src_dir)
    out_dir = os.path.abspath(out_dir)

    parent_out = os.path.join(
        out_dir,
        f"{os.path.basename(src_dir.rstrip(os.sep))}_captions_llava"
    )
    os.makedirs(parent_out, exist_ok=True)

    total_videos = 0
    total_images = 0

    for subdir, _, files in os.walk(src_dir):
        # Check if this subdir contains images
        image_files = [f for f in files if f.lower().endswith(VALID_EXTS)]
        if not image_files:
            continue

        # Determine video_id
        video_id = infer_video_id_from_subdir(subdir)

        # Output locations
        video_out_dir = os.path.join(parent_out, video_id)
        out_docx_name = f"{video_id}_captions_llava.docx"
        out_json_name = f"{video_id}_captions_llava.json"
        out_docx_path = os.path.join(video_out_dir, out_docx_name)
        out_json_path = os.path.join(video_out_dir, out_json_name)

        rel_path = os.path.relpath(subdir, src_dir)
        print(f"\n📂 Processing folder: {rel_path}  -> video_id: {video_id}")

        count = caption_folder_to_docx_and_json(
            subdir,
            llava_model,
            out_docx_path=out_docx_path,
            out_json_path=out_json_path,
        )
        if count > 0:
            total_videos += 1
            total_images += count

    if total_videos == 0:
        print("\n⚠️ No images were found. No documents were created.")
    else:
        print(f"\n✅ Done. Videos processed: {total_videos}, frames captioned: {total_images}")
        print(f"📁 Output root: {parent_out}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Batch caption frames using LLaVa; save per-video DOCX (with images) and JSON."
    )
    parser.add_argument(
        '--src_dir',
        default="../../../VideoAnalysis/data/frames/frames_conventional",
        help='Root folder containing subfolders of images (frames_root)'
    )
    parser.add_argument(
        '--out_dir',
        default="../../data/minicpm_caption",
        help='Output directory to store caption documents'
    )
    args = parser.parse_args()

    caption_all_folders(args.src_dir, args.out_dir)