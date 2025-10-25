#!/usr/bin/env python3
"""
Create a DOCX per video containing:
  1) Topics (id, title, description, evidence)
  2) Full transcript (plain text, no timestamps)
  3) Timestamped transcript (if available)

Transcript sources in transcripts/:
- <video_id>.txt (plain, preferred)
- <video_id>_timestamped.txt (if present, included as a separate section)

Topics source in topics/:
- <video_id>.topics.json

Usage:
    python build_topic_transcript_docx.py --transcripts_root transcripts --topics_root topics --out_dir docs_conventional

Each video gets one DOCX: <out_dir>/<video_id>_topics_and_transcript.docx
"""

import os
import json
import argparse
from os.path import join, exists, isdir
from typing import Any, Dict, List, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# -------------------- Helpers --------------------

def safe_get(d: Dict, key: str, default=None):
    v = d.get(key, default)
    return v if v is not None else default

def load_json(path: str) -> Optional[Dict[str, Any]]:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"[WARN] Failed to read JSON: {path}: {e}")
        return None

def load_txt_transcript(path: str) -> Optional[Dict[str, Any]]:
    """Plain or timestamped TXT with 3-line header then body."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        if not lines:
            return None
        video_id = lines[0].strip() if len(lines) >= 1 else None
        url      = lines[1].strip() if len(lines) >= 2 else None
        category = lines[2].strip() if len(lines) >= 3 else None

        body_start = 3
        while body_start < len(lines) and lines[body_start].strip() == "":
            body_start += 1

        # 👇 keep original line breaks for timestamped files
        if path.endswith("_timestamped.txt"):
            body = "\n".join(lines[body_start:]).strip()
        else:
            body = "\n".join(lines[body_start:]).strip()  # keep \n; plain will be paragraphized later

        return {"video_id": video_id, "url": url, "category": category, "text": body}
    except Exception as e:
        print(f"[WARN] Failed to parse TXT transcript: {path}: {e}")
        return None

def add_transcript_section_timestamped(doc: Document, text: str):
    doc.add_page_break()
    doc.add_heading('Transcript (timestamped)', level=1)
    if not text:
        doc.add_paragraph("(No transcript text found.)")
        return
    # one segment per non-empty line
    for line in (l.strip() for l in text.splitlines() if l.strip()):
        doc.add_paragraph(line)


# -------------------- DOCX Builders --------------------

def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def add_title_block(doc: Document, video_meta: Dict[str, Any], topic_meta: Optional[Dict[str, Any]]):
    title = safe_get(video_meta, "title") or safe_get(topic_meta or {}, "title") or safe_get(video_meta, "video_id") or "Video"
    url = safe_get(video_meta, "url") or safe_get(topic_meta or {}, "url")

    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if url:
        p = doc.add_paragraph()
        add_hyperlink(p, url, url)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = []
    if video_meta.get("video_id"):
        meta.append(f"Video ID: {video_meta['video_id']}")
    if safe_get(video_meta, "category"):
        meta.append(f"Category: {safe_get(video_meta, 'category')}")
    if safe_get(topic_meta or {}, "model"):
        meta.append(f"Topics model: {topic_meta['model']}")
    if safe_get(topic_meta or {}, "created_at"):
        meta.append(f"Topics created: {topic_meta['created_at']}")
    if meta:
        doc.add_paragraph(" | ".join(meta))

def add_topics_section(doc: Document, topics: List[Dict[str, Any]]):
    doc.add_heading('Topics', level=1)
    if not topics:
        doc.add_paragraph("(No topics found.)")
        return

    for t in topics:
        tid = safe_get(t, 'id', '—')
        ttitle = safe_get(t, 'title', '(untitled)')
        doc.add_heading(f"{tid}: {ttitle}", level=2)

        desc = safe_get(t, 'description')
        if desc:
            doc.add_paragraph(desc)

        ev = safe_get(t, 'evidence') or []
        if isinstance(ev, list) and ev:
            doc.add_paragraph("Evidence:")
            for e in ev:
                etxt = safe_get(e, 'text')
                if etxt:
                    doc.add_paragraph(etxt, style='List Bullet')

import re

def add_transcript_section(doc: Document, text: str, heading: str = 'Transcript'):
    doc.add_page_break()
    doc.add_heading(heading, level=1)
    if not text:
        doc.add_paragraph("(No transcript text found.)")
        return

    # Normalize newlines
    text = text.replace('\r\n', '\n').strip()

    # Treat 2+ newlines as a paragraph break; single newlines become spaces
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    for p in paragraphs:
        p = re.sub(r'[ \t]*\n[ \t]*', ' ', p)   # collapse single newlines inside a paragraph
        p = re.sub(r'[ \t]{2,}', ' ', p)       # collapse multiple spaces
        doc.add_paragraph(p)


# -------------------- Build Logic --------------------

def build_doc_for_video(video_id: str, transcripts_root: str, topics_root: str, out_dir: str, overwrite: bool = True) -> Optional[str]:
    topics_path = join(topics_root, f"{video_id}.topics.json")
    plain_path = join(transcripts_root, f"{video_id}.txt")
    stamped_path = join(transcripts_root, f"{video_id}_timestamped.txt")

    tx_plain = load_txt_transcript(plain_path) if exists(plain_path) else None
    tx_stamped = load_txt_transcript(stamped_path) if exists(stamped_path) else None
    topics = load_json(topics_path) if exists(topics_path) else None

    if not tx_plain and not tx_stamped and not topics:
        print(f"[WARN] Skipping {video_id}: no transcript (.txt/_timestamped.txt) or topics found.")
        return None

    os.makedirs(out_dir, exist_ok=True)
    out_path = join(out_dir, f"{video_id}_topics_and_transcript.docx")
    if exists(out_path) and not overwrite:
        print(f"[SKIP] Exists: {out_path}")
        return out_path

    video_meta: Dict[str, Any] = {"video_id": video_id}
    meta_src = tx_plain or tx_stamped
    if meta_src:
        video_meta.update({k: v for k, v in meta_src.items() if k in ("video_id", "url", "category") and v})

    doc = Document()
    add_title_block(doc, video_meta, topics)
    add_topics_section(doc, (topics or {}).get('topics', []))

    # Transcript sections — plain first, then timestamped if present
    if tx_plain and tx_plain.get('text'):
        add_transcript_section(doc, tx_plain['text'], heading='Transcript (plain)')

    if tx_stamped and tx_stamped.get('text'):
        add_transcript_section_timestamped(doc, tx_stamped['text'])

    doc.save(out_path)
    print(f"[OK] Wrote {out_path}")
    return out_path

# -------------------- CLI --------------------

def find_video_ids(transcripts_root: str, topics_root: str, explicit_ids: List[str]) -> List[str]:
    if explicit_ids:
        return explicit_ids

    vids = set()

    # From transcripts: collect base ids from .txt and _timestamped.txt
    if isdir(transcripts_root):
        for name in os.listdir(transcripts_root):
            if name.endswith('_timestamped.txt'):
                vids.add(name[:-len('_timestamped.txt')])
            elif name.endswith('.txt'):
                vids.add(name[:-4])

    # From topics
    if isdir(topics_root):
        for name in os.listdir(topics_root):
            if name.endswith('.topics.json'):
                vids.add(name[:-len('.topics.json')])

    return sorted(vids)

def main():
    ap = argparse.ArgumentParser(description="Create per-video DOCX files (topics + plain transcript).")
    ap.add_argument('--transcripts_root', default='../data/transcripts', help='Directory containing <video_id>.txt transcripts')
    ap.add_argument('--topics_root', default='../data/topics', help='Directory containing <video_id>.topics.json files')
    ap.add_argument('--out_dir', default='../results/docs', help='Output directory for generated DOCX files')
    ap.add_argument('-v', '--video_id', action='append', default=[], help='Specific video_id(s) to process; can repeat')
    ap.add_argument('--no-overwrite', dest='overwrite', action='store_false', help='Do not overwrite existing files')
    args = ap.parse_args()

    video_ids = find_video_ids(args.transcripts_root, args.topics_root, args.video_id)
    if not video_ids:
        print('[ERROR] No videos found.')
        return

    for vid in video_ids:
        build_doc_for_video(vid, args.transcripts_root, args.topics_root, args.out_dir, overwrite=args.overwrite)

if __name__ == '__main__':
    main()
