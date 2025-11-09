import os
import json
import torch
from PIL import Image
from transformers import AutoModel, AutoTokenizer
from docx import Document
from docx.shared import Inches
import traceback
import argparse


VALID_EXTS = ('.jpg', '.jpeg', '.png')


def generate_caption(image_path, model, tokenizer, device):
    try:
        image = Image.open(image_path).convert('RGB')
    except Exception as e:
        return f"[Error loading image: {e}]"

    question = 'Please provide a detailed caption for this image.'
    msgs = [{'role': 'user', 'content': question}]
    context = "This is an image captioning task."

    try:
        res = model.chat(
            image=image,
            context=context,
            msgs=msgs,
            tokenizer=tokenizer,
            sampling=True,
            temperature=0.7
        )
        return res[0]
    except Exception as e:
        traceback.print_exc()
        return f"[Error generating caption: {e}]"


def infer_video_id_from_subdir(subdir):
    """
    If frames are under .../<video_id>/raw_frames or similar,
    use the parent folder as video_id. Otherwise use the basename.
    """
    base = os.path.basename(subdir.rstrip(os.sep))
    if base.lower() in {'raw_frames', 'frames', 'images'}:
        return os.path.basename(os.path.dirname(subdir.rstrip(os.sep)))
    return base


def caption_folder_to_docx_and_json(
    frames_folder,
    model,
    tokenizer,
    device,
    out_docx_path,
    out_json_path,
    image_width_inches=5.0
):
    """
    Create a DOCX with image previews and captions for all frames in frames_folder,
    and a JSON file with (video_id, frame_id, caption) for downstream processing.
    """
    image_files = sorted(
        f for f in os.listdir(frames_folder)
        if f.lower().endswith(VALID_EXTS)
    )
    if not image_files:
        return 0

    # Prepare DOCX
    doc = Document()
    folder_name = os.path.basename(frames_folder.rstrip(os.sep))
    doc.add_heading(f"Captions for {folder_name}", level=0)

    # Prepare JSON data
    video_id = infer_video_id_from_subdir(frames_folder)
    json_payload = {
        "video_id": video_id,
        "frames": []  # list of {"frame_id": str, "caption": str}
    }

    count = 0
    for filename in image_files:
        image_path = os.path.join(frames_folder, filename)
        frame_id = os.path.splitext(filename)[0]

        print(f"  Processing {filename}")

        # Frame header
        doc.add_heading(filename, level=1)

        # Insert the actual frame image
        try:
            doc.add_picture(image_path, width=Inches(image_width_inches))
        except Exception as e:
            doc.add_paragraph(f"[Error inserting image: {e}]")

        # Generate and insert caption
        caption = generate_caption(image_path, model, tokenizer, device)
        doc.add_paragraph(caption)

        # Append to JSON list
        json_payload["frames"].append({
            "frame_id": frame_id,
            "caption": caption
        })

        count += 1

    # Ensure output directory exists then save
    os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)

    doc.save(out_docx_path)
    with open(out_json_path, "w", encoding="utf-8") as f:
        json.dump(json_payload, f, ensure_ascii=False, indent=2)

    print(f"  Saved caption .docx: {out_docx_path}")
    print(f"  Saved caption .json: {out_json_path}")

    return count


def caption_all_folders(root_folder, output_dir):
    """
    Walk root_folder for subfolders containing images.
    For each such subfolder, produce:
      DOCX: <output_dir>/<basename(root)>_captions_MiniCPM/<video_id>/<video_id>_captions_MiniCPM.docx
      JSON: <output_dir>/<basename(root)>_captions_MiniCPM/<video_id>/<video_id>_captions_MiniCPM.json
    """
    print("Loading model and tokenizer...")
    # model_id = 'openbmb/MiniCPM-V'
    # model = AutoModel.from_pretrained(model_id, trust_remote_code=True, torch_dtype=torch.float32)
    # device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    # model = model.to(device)
    # model.eval()
    # tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True)
    model_id = "openbmb/MiniCPM-V"

    # Tokenizer (slow version avoids some platform/Rust issues)
    tokenizer = AutoTokenizer.from_pretrained(model_id, trust_remote_code=True, use_fast=False)

    # Model
    use_cuda = torch.cuda.is_available()
    device = torch.device("cuda" if use_cuda else "cpu")
    dtype = torch.float16 if use_cuda else torch.float32

    if use_cuda:
        print(f"Using GPU: {torch.cuda.get_device_name(0)}")
    else:
        print("WARNING: CUDA not available. Using CPU (will be slow)")

    model = AutoModel.from_pretrained(
        model_id,
        trust_remote_code=True,
        torch_dtype=dtype,
    ).to(device)

    # Force all parameters and buffers to the correct dtype
    if use_cuda:
        model = model.half()  # Convert everything to float16
    else:
        model = model.float()  # Convert everything to float32

    model.eval()
    print("Model and tokenizer loaded.")

    root_folder = os.path.abspath(root_folder)
    output_dir = os.path.abspath(output_dir)

    parent_out = os.path.join(
        output_dir,
        f"{os.path.basename(root_folder.rstrip(os.sep))}_captions_MiniCPM"
    )
    os.makedirs(parent_out, exist_ok=True)

    total_videos = 0
    total_images = 0

    for subdir, _, files in os.walk(root_folder):
        # Check if this subdir contains images
        image_files = [f for f in files if f.lower().endswith(VALID_EXTS)]
        if not image_files:
            continue

        # Determine video_id
        video_id = infer_video_id_from_subdir(subdir)

        # Output locations
        video_out_dir = os.path.join(parent_out, video_id)
        out_docx_name = f"{video_id}_captions_MiniCPM.docx"
        out_json_name = f"{video_id}_captions_MiniCPM.json"
        out_docx_path = os.path.join(video_out_dir, out_docx_name)
        out_json_path = os.path.join(video_out_dir, out_json_name)

        rel_path = os.path.relpath(subdir, root_folder)
        print(f"\nProcessing folder: {rel_path}  -> video_id: {video_id}")

        count = caption_folder_to_docx_and_json(
            subdir,
            model,
            tokenizer,
            device,
            out_docx_path=out_docx_path,
            out_json_path=out_json_path,
        )
        if count > 0:
            total_videos += 1
            total_images += count

    if total_videos == 0:
        print("\nNo images were found. No documents were created.")
    else:
        print(f"\nDone. Videos processed: {total_videos}, frames captioned: {total_images}")
        print(f"Output root: {parent_out}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Batch caption frames using MiniCPM-V; save per-video DOCX (with images) and JSON.")
    parser.add_argument('--src_dir', default="../../../VideoAnalysis/data/frames/frames_conventional", help='Root folder containing subfolders of images (frames_root)')
    parser.add_argument('--out_dir', default="../../data/minicpm_caption", help='Output directory to store caption documents')
    args = parser.parse_args()

    caption_all_folders(args.src_dir, args.out_dir)
