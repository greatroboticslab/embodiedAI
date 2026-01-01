#!/usr/bin/env python3
"""
Correlate viewer comments (from JSON files) to *topics* (title + description),
then generate a rebuildable DOCX report with coverage and top matches.

Inputs
  1) frames_path: Either a project root containing <video_id>/raw_frames
     subfolders, or a single .../raw_frames folder.
  2) comments_dir: A directory containing JSON/JSONL comment files named <video_id>.json.
  3) topics_root (arg): A directory holding JSON topic files, default "topics".

Outputs (per video)
  - JSONL cache (resume-safe):
      <video_id>_topiccorr_n{n}_p{pen}_m{min}_<embed>_k{retrieval}.results.jsonl
  - DOCX report: <video_id>_correlation_topics_n{n}_p{pen}_m{min}.docx
"""

import os
import re
import json
import argparse
from urllib.parse import urlparse, parse_qs
import sys
from collections import defaultdict

# local imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from config import Config
from utils.llm_utils import generate_response, stream_parser
from docx import Document
from docx.shared import Inches

YES = {"yes", "y", "true", "t", "1"}


# -------------------- URL parsing --------------------

def get_video_id_from_url(text: str):
    text = text.strip()
    if not text or "youtu" not in text:
        return None
    m = re.search(r"https?://[^\s)\]>]+", text)
    if not m:
        return None
    url = re.sub(r'[:.,;!?"\)\]]+$', '', m.group(0))
    try:
        u = urlparse(url)
        host = (u.netloc or "").lower()
        if "youtu.be" in host:
            vid = u.path.strip("/").split("/")[0]
            return clean_video_id(vid)
        qs = parse_qs(u.query or "")
        if "v" in qs and qs["v"]:
            return clean_video_id(qs["v"][0])
        parts = [p for p in u.path.split("/") if p]
        if "embed" in parts:
            idx = parts.index("embed")
            if idx + 1 < len(parts):
                return clean_video_id(parts[idx + 1])
    except Exception:
        pass
    return None


def clean_video_id(vid: str):
    m = re.match(r"([A-Za-z0-9_-]{6,})", vid)
    return m.group(1) if m else vid


# -------------------- JSON Comment Parsers --------------------

def parse_comments_dir(comments_dir: str):
    """
    Returns dict: {video_id: [{"url": url, "comment": text}, ...]}
    Scans comments_dir recursively for .json files.
    Assumes filename is <video_id>.json.
    """
    by_video = defaultdict(list)
    
    if not os.path.isdir(comments_dir):
        print(f"[WARN] Comments directory not found: {comments_dir}")
        return dict(by_video)

    print(f"[INFO] Scanning comments in: {comments_dir}")
    
    for root, dirs, files in os.walk(comments_dir):
        for f in files:
            if not f.endswith(".json"):
                continue
                
            # Assume filename is video_id.json
            video_id = os.path.splitext(f)[0]
            # Verify it looks like a video ID
            if not re.match(r"^[A-Za-z0-9_-]{11}$", video_id):
                 # Relaxed check for 11 chars or likely video ID
                 if len(video_id) < 5: 
                     continue

            full_path = os.path.join(root, f)
            url = f"https://www.youtube.com/watch?v={video_id}"
            
            try:
                with open(full_path, "r", encoding="utf-8") as f_in:
                    # Try reading as JSONL (one JSON per line) first
                    # because the sample showed JSONL format
                    lines = f_in.readlines()
                    
                parsed_comments = []
                # Check if it's a single JSON object or JSONL
                first_char = lines[0].strip()[0] if lines and lines[0].strip() else ""
                
                if len(lines) > 0:
                    # Attempt JSONL
                    try:
                        for line in lines:
                            line = line.strip()
                            if not line: continue
                            data = json.loads(line)
                            if "text" in data:
                                parsed_comments.append(data["text"])
                    except json.JSONDecodeError:
                        # Fallback: maybe it's a single big JSON list?
                        try:
                            full_text = "".join(lines)
                            data = json.loads(full_text)
                            if isinstance(data, list):
                                for item in data:
                                    if isinstance(item, dict) and "text" in item:
                                        parsed_comments.append(item["text"])
                            elif isinstance(data, dict) and "text" in data:
                                parsed_comments.append(data["text"])
                        except:
                            print(f"[WARN] Could not parse {f} as JSONL or JSON.")
                
                for text in parsed_comments:
                    if text and isinstance(text, str):
                         by_video[video_id].append({"url": url, "comment": text})
                         
            except Exception as e:
                print(f"[ERROR] processing {f}: {e}")

    return dict(by_video)


# -------------------- Topic loading --------------------

def _find_topics_json(raw_frames_dir: str, topics_root: str):
    """Look for <topics_root>/<video_id>.topics.json, else scan raw_frames for *topics.json"""
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))
    
    # Potential search roots
    roots_to_check = []
    if topics_root:
        roots_to_check.append(topics_root)
    
    # Add smart defaults relative to this script: ../data/topics_embodied, ../data/topics_conventional
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) # TranscriptAnalysis/
    data_dir = os.path.join(base_dir, "data")
    roots_to_check.append(os.path.join(data_dir, "topics_embodied"))
    roots_to_check.append(os.path.join(data_dir, "topics_conventional"))

    # 1) Check preferred roots
    for root in roots_to_check:
        if not root or not os.path.exists(root):
            continue
        p = os.path.join(root, f"{video_id}.topics.json")
        if os.path.isfile(p):
            return p
            
    # 2) sibling/within raw_frames folder (legacy)
    for f in os.listdir(raw_frames_dir):
        if f.endswith(".topics.json") and video_id in f:
            return os.path.join(raw_frames_dir, f)
            
    # 3) one level up under a 'topics' folder (legacy)
    candidate = os.path.join(os.path.dirname(os.path.dirname(raw_frames_dir)), "topics", f"{video_id}.topics.json")
    if os.path.isfile(candidate):
        return candidate
        
    return None


def parse_topics(topics_json_path: str):
    """
    Returns list of topics with keys: id, title, description, start_s, end_s.
    Missing fields are tolerated.
    """
    with open(topics_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    topics = []
    for t in data.get("topics", []):
        topics.append({
            "id": t.get("id") or t.get("topic_id") or f"T{len(topics) + 1}",
            "title": t.get("title", "").strip(),
            "description": t.get("description", "").strip(),
            "start_s": t.get("start_s"),
            "end_s": t.get("end_s"),
        })
    return topics


# -------------------- Embeddings + helpers --------------------

def _embed_texts(texts, model_name="all-mpnet-base-v2"):
    """Returns L2-normalized embeddings. If sentence-transformers is missing, raises ImportError."""
    try:
        from sentence_transformers import SentenceTransformer
        try:
            import numpy as np
        except ImportError:
            pass
    except Exception as e:
        raise ImportError("sentence-transformers not installed") from e

    model = SentenceTransformer(model_name)
    X = model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
    return X


def pick_representative_frames(frame_names, k=3):
    """Evenly spaced picks from the list, preserving order."""
    if not frame_names:
        return []
    if k >= len(frame_names):
        return frame_names
    idxs = [int(round(i * (len(frame_names) - 1) / (k - 1))) for i in range(k)]
    seen = set()
    picks = []
    for idx in idxs:
        if idx not in seen:
            picks.append(frame_names[idx])
            seen.add(idx)
    return picks


def discover_raw_frames(frames_input: str):
    frames_input = os.path.abspath(frames_input)
    targets = []
    if os.path.basename(frames_input) == "raw_frames":
        targets.append(frames_input)
        return targets
    if not os.path.exists(frames_input):
         return []
    for entry in sorted(os.listdir(frames_input)):
        d = os.path.join(frames_input, entry, "raw_frames")
        if os.path.isdir(d):
            targets.append(d)
    return targets


# -------------------- LLM scoring --------------------

def ask_model_topic(model_name: str, topic_title: str, topic_desc: str, comment: str):
    """Return {"correlated": bool, "score": int, "reason": str} for TOPIC vs COMMENT."""
    topic_text = (topic_title or "").strip()
    if topic_desc:
        topic_text += "\n" + topic_desc.strip()

    prompt = (
        "You are checking whether a TOPIC (title + description) and a VIEWER COMMENT are strongly related.\n"
        "Strong = the comment directly discusses the same steps, objects, settings, parameters, claims, or goals described by the topic; not generic praise.\n\n"
        "Give a correlation confidence 0-100, and a brief reason citing specific overlaps.\n\n"
        f"TOPIC TITLE:\n{topic_title}\n\nTOPIC DESCRIPTION:\n{topic_desc}\n\nCOMMENT:\n{comment}\n\n"
        "Respond ONLY as compact JSON: {\"correlated\": true|false, \"score\": 0-100, \"reason\": \"brief explanation\"}"
    )
    stream = generate_response(model_name, prompt)
    text = "".join(stream_parser(stream)).strip()
    m = re.search(r"\{.*?\}", text, flags=re.S)
    jtxt = m.group(0) if m else text
    try:
        obj = json.loads(jtxt)
        corr = str(obj.get("correlated", False)).lower() in YES
        score = int(obj.get("score", 0))
        score = max(0, min(100, score))
        reason = str(obj.get("reason", "")).strip()
        return {"correlated": corr, "score": score, "reason": reason}
    except Exception:
        corr = "yes" in text.lower()
        m2 = re.search(r"(\d{1,3})", text)
        score = int(m2.group(1)) if m2 else 0
        score = max(0, min(100, score))
        return {"correlated": corr, "score": score, "reason": text[:400]}


# -------------------- IO helpers (resume-safe) --------------------

def results_path(output_dir: str, video_id: str, seg_tag: str, embed_model: str, retrieval_topk: int):
    safe_model = re.sub(r"[^A-Za-z0-9_-]+", "", embed_model or "mpnet")
    return os.path.join(
        output_dir,
        f"{video_id}_topiccorr_{seg_tag}_{safe_model}_k{retrieval_topk}.results.jsonl"
    )


def load_existing_results(results_file: str):
    done = {}
    if os.path.isfile(results_file):
        with open(results_file, "r", encoding="utf-8") as f:
            for line in f:
                try:
                    rec = json.loads(line)
                    key = (rec.get("topic_id"),)
                    done[key] = rec
                except Exception:
                    continue
    return done


def append_result(results_file: str, record: dict):
    with open(results_file, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


# -------------------- DOCX helpers --------------------

def _add_thumbnails_row(doc, base_dir, frame_list, thumb_width_in=1.3):
    tbl = doc.add_table(rows=1, cols=max(1, len(frame_list)))
    tbl.autofit = True
    for col, relpath in enumerate(frame_list):
        cell = tbl.rows[0].cells[col]
        img_path = os.path.join(base_dir, relpath)
        if os.path.exists(img_path):
            try:
                run = cell.paragraphs[0].add_run()
                run.add_picture(img_path, width=Inches(thumb_width_in))
            except Exception:
                cell.text = relpath
        else:
            cell.text = relpath
    return tbl


def build_docx_from_results(out_docx: str, topics: list, results_map: dict,
                            min_score: int, top_k: int, model_name: str,
                            frame_dir: str,
                            total_comments_available: int | None = None):
    doc = Document()
    doc.add_heading("Comment–Topic Correlation Report", 0)
    doc.add_paragraph(f"Model: {model_name}")

    # Aggregate
    # Aggregate stats
    total_pairs = sum(rec.get("checked_pairs", 0) for rec in results_map.values())
    base_dir = frame_dir

    # Topic metadata map
    topic_map = {str(t.get("id")): t for t in topics}

    # Calculate summary metrics and per-topic correlation scores
    corr_true_ids = set()       # correlated=True
    high_score_ids = set()      # correlated=True AND score >= min
    
    topic_sort_data = [] # (avg_score, topic_obj)

    for t in topics:
        tid = t["id"]
        rec = results_map.get((tid,), {})
        candidates = rec.get("candidates", [])
        
        valid_scores = []
        for c in candidates:
            cid = (c.get("url", ""), c.get("comment", ""))
            s = int(c.get("score", 0))
            is_corr = c.get("correlated")
            
            if is_corr:
                corr_true_ids.add(cid)
                if s >= min_score:
                    high_score_ids.add(cid)
                    valid_scores.append(s)
        
        # Metric for sorting topics: Avg score of VALID matches (>= min_score)
        # If no valid matches, score is 0.
        avg_score = sum(valid_scores) / len(valid_scores) if valid_scores else 0.0
        topic_sort_data.append((avg_score, t))

    # Sort topics: descending by avg score, then ID
    topic_sort_data.sort(key=lambda x: (x[0], str(x[1]["id"])), reverse=True)
    sorted_topics = [x[1] for x in topic_sort_data]

    # Summary Section
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Topics processed: {len(topics)}")
    if total_comments_available is not None:
        doc.add_paragraph(f"Comments available for this video: {total_comments_available}")
    
    # Coverage stats
    pct_any = 0.0
    pct_high = 0.0
    if total_comments_available and total_comments_available > 0:
        pct_any = 100.0 * len(corr_true_ids) / total_comments_available
        pct_high = 100.0 * len(high_score_ids) / total_comments_available

    doc.add_paragraph(f"Comments with correlated=True: {len(corr_true_ids)} ({pct_any:.1f}%)")
    doc.add_paragraph(f"Comments with score >= {min_score}: {len(high_score_ids)} ({pct_high:.1f}%)")
    doc.add_paragraph(f"Total candidate pairs checked: {total_pairs}")
    
    doc.add_paragraph("-" * 40)

    # Per-Topic Details (Sorted)
    for t in sorted_topics:
        tid = t["id"]
        title = t.get("title", "").strip()
        desc = t.get("description", "").strip()
        rep = t.get("rep_frames", [])
        
        # Header
        doc.add_heading(f"Topic {tid}: {title}", level=2)
        
        # Thumbnails
        if rep:
            _add_thumbnails_row(doc, base_dir, rep, thumb_width_in=1.5)
            
        # Description
        if desc:
            p = doc.add_paragraph()
            run = p.add_run(f"Description: {desc}")
            run.italic = True
            
        # Matched Comments
        rec = results_map.get((tid,), {})
        candidates = rec.get("candidates", [])
        # Filter for display: score >= min_score AND correlated
        kept = [c for c in candidates if c.get("score", 0) >= min_score and c.get("correlated", False)]
        kept.sort(key=lambda x: x.get("score", 0), reverse=True)
        kept = kept[:top_k]

        if kept:
            doc.add_paragraph(f"Top {len(kept)} Correlated Comments:", style="Heading 3")
            
            # Create a table for cleaner layout
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Score"
            hdr_cells[1].text = "Comment"
            hdr_cells[2].text = "Reason"
            # Set widths roughly (optional, python-docx autofit handles mostly)
            
            for r in kept:
                row_cells = table.add_row().cells
                row_cells[0].text = str(r.get("score", 0))
                
                # Comment + URL
                c_text = r.get("comment", "")
                if r.get("url"):
                    c_text += f"\n\n[Source]({r['url']})"
                row_cells[1].text = c_text
                
                row_cells[2].text = r.get("reason", "")
                
            doc.add_paragraph("") # Spacing after table
        else:
            doc.add_paragraph("No strongly correlated comments found.", style="Body Text")
            doc.add_paragraph("") # Spacing

    doc.save(out_docx)


# -------------------- Frame discovery for thumbnails --------------------

def list_all_frames(raw_frames_dir: str):
    files = [f for f in os.listdir(raw_frames_dir) if f.lower().endswith((".jpg", ".jpeg", ".png"))]
    files.sort()
    return files


# -------------------- Core processing --------------------

def process_one_video(raw_frames_dir: str, comments_by_video: dict, topics_root: str, results_root: str, model: str,
                      min_score: int, top_k_scan: int,
                      seg_n_sections: int | None, seg_penalty: int | None,
                      seg_min_size: int, embed_model: str,
                      retrieval_topk: int, rep_k: int):
    """Segment knobs (n/penalty/min) are kept for tag consistency only; topics are pre-defined."""
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))

    # Prepare output dir
    video_out_dir = os.path.join(results_root, video_id)
    os.makedirs(video_out_dir, exist_ok=True)

    topics_json = _find_topics_json(raw_frames_dir, topics_root)
    if not topics_json:
        print(f"[SKIP] {video_id}: topics JSON not found (looked under {topics_root} and nearby).")
        return

    topics = parse_topics(topics_json)
    if not topics:
        print(f"[SKIP] {video_id}: no topics parsed.")
        return

    # Representative frames per topic (best effort): evenly from whole folder
    # If you later add per-topic frame ranges, swap to filtered subset here.
    all_frames = list_all_frames(raw_frames_dir)
    for t in topics:
        t["rep_frames"] = pick_representative_frames(all_frames, k=rep_k)

    comments_for_video = comments_by_video.get(video_id, [])
    if not comments_for_video:
        print(f"[WARN] {video_id}: no comments found for this video.")
        return
    else:
        print(f"[INFO] {video_id}: {len(comments_for_video)} comments found for this video.")

    # Pre-embed comments for retrieval (if possible)
    comment_texts = [c["comment"] for c in comments_for_video]
    C = None
    retrieval_ok = False
    try:
        C = _embed_texts(comment_texts, model_name=embed_model)
        retrieval_ok = True
    except ImportError:
        print(
            f"[WARN] {video_id}: sentence-transformers missing; skipping fast retrieval (LLM will score all comments).")

    # Resume state
    seg_tag = f"n{seg_n_sections or 0}_p{seg_penalty or 0}_m{seg_min_size}"
    res_file = results_path(video_out_dir, video_id, seg_tag, embed_model, retrieval_topk)
    existing = load_existing_results(res_file)
    print(f"[RESUME] {video_id}: {len(existing)}/{len(topics)} topics already processed.")

    try:
        import numpy as np
    except ImportError:
        np = None

    # Process topics
    for t in topics:
        tid = t["id"]
        key = (tid,)
        if key in existing:
            continue

        title = t.get("title", "")
        desc = t.get("description", "")
        topic_text = (title or "").strip()
        if desc:
            topic_text += "\n" + desc.strip()

        checked_pairs = 0
        candidates = []

        # choose candidate comments via retrieval
        candidate_indices = range(len(comment_texts))
        if retrieval_ok and retrieval_topk > 0 and topic_text and np:
            t_vec = _embed_texts([topic_text], model_name=embed_model)[0]
            sims = np.dot(C, t_vec)
            idx = np.argsort(-sims)[:min(retrieval_topk, len(comment_texts))]
            candidate_indices = idx

        for i in candidate_indices:
            item = comments_for_video[i]
            res = ask_model_topic(model, title, desc, item["comment"])
            checked_pairs += 1
            candidates.append({
                "correlated": bool(res["correlated"]),
                "score": int(res["score"]),
                "reason": res.get("reason", ""),
                "comment": item["comment"],
                "url": item["url"]
            })

        record = {
            "video_id": video_id,
            "topic_id": tid,
            "title": title,
            "checked_pairs": checked_pairs,
            "candidates": candidates,
        }
        append_result(res_file, record)
        existing[key] = record
        print(f"[OK] {video_id}: topic {tid} processed ({checked_pairs} pairs).")

    # Build DOCX
    out_docx = os.path.join(video_out_dir, f"{video_id}_correlation_topics_{seg_tag}.docx")
    build_docx_from_results(
        out_docx, topics, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        frame_dir=raw_frames_dir,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[DONE] {video_id}: report saved -> {out_docx}")


def rebuild_docx_only(raw_frames_dir: str, comments_by_video: dict, topics_root: str, results_root: str, model: str,
                      min_score: int, top_k_scan: int,
                      seg_n_sections: int | None, seg_penalty: int | None,
                      seg_min_size: int, embed_model: str,
                      retrieval_topk: int, rep_k: int):
    video_id = os.path.basename(os.path.dirname(raw_frames_dir))

    # Prepare output dir
    video_out_dir = os.path.join(results_root, video_id)
    os.makedirs(video_out_dir, exist_ok=True)

    topics_json = _find_topics_json(raw_frames_dir, topics_root)
    if not topics_json:
        print(f"[SKIP] {video_id}: topics JSON not found for rebuild.")
        return
    topics = parse_topics(topics_json)
    if not topics:
        print(f"[SKIP] {video_id}: no topics parsed for rebuild.")
        return

    # thumbnails best-effort
    all_frames = list_all_frames(raw_frames_dir)
    for t in topics:
        t["rep_frames"] = pick_representative_frames(all_frames, k=rep_k)

    seg_tag = f"n{seg_n_sections or 0}_p{seg_penalty or 0}_m{seg_min_size}"
    res_file = results_path(video_out_dir, video_id, seg_tag, embed_model, retrieval_topk)
    if not os.path.isfile(res_file):
        print(f"[SKIP] {video_id}: no results file found to rebuild.")
        return

    existing = load_existing_results(res_file)
    comments_for_video = comments_by_video.get(video_id, [])

    out_docx = os.path.join(video_out_dir, f"{video_id}_correlation_topics_{seg_tag}.docx")
    build_docx_from_results(
        out_docx, topics, existing,
        min_score=min_score, top_k=top_k_scan,
        model_name=model,
        frame_dir=raw_frames_dir,
        total_comments_available=len(comments_for_video) if comments_for_video else None
    )
    print(f"[REBUILT] {video_id}: report saved -> {out_docx}")


def determine_subfolder(raw_frames_dir: str):
    """
    Heuristic: if the path contains 'frames_embodied', return 'correlation_embodied'.
    If 'frames_conventional', return 'correlation_conventional'.
    """
    parts = os.path.normpath(raw_frames_dir).split(os.sep)
    for p in parts:
        if p == "frames_embodied":
            return "correlation_embodied"
        if p == "frames_conventional":
            return "correlation_conventional"
    # Fallback: check direct parent of video_id if it's named frames_something
    # raw_frames_dir is .../<video_id>/raw_frames
    # parent is <video_id>
    # grantparent might be frames_conventional
    try:
        grandparent = os.path.basename(os.path.dirname(os.path.dirname(raw_frames_dir)))
        if grandparent.startswith("frames_"):
            return grandparent.replace("frames_", "correlation_")
    except:
        pass
    return None


# -------------------- CLI --------------------

def main():
    ap = argparse.ArgumentParser(
        description="Correlate viewer comments (from JSON) to topics (resume-safe) and build a DOCX report.")
    ap.add_argument("--frames_path", default="../../VideoAnalysis/data/frames/frames_conventional",
                    help="Path to frames root (containing <video_id>/raw_frames) OR a single raw_frames folder.")
    ap.add_argument("--results_root", default="../results_youtube",
                    help="Root directory for results. Defaults to ../results_youtube relative to this script.")
    ap.add_argument("--comments_dir",
                    default="../../VideoAnalysis/data/comments/conventional",
                    help="Directory containing <video_id>.json comment files.")

    ap.add_argument("--topics_root", default="../data/topics_conventional", help="Directory containing <video_id>.topics.json files.")
    ap.add_argument("--model", default=None, help="LLM name for scoring (defaults to Config.ollama_models[0]).")
    ap.add_argument("--min_score", type=int, default=60,
                    help="Minimum score to include a matched comment in the final report.")
    ap.add_argument("--top_k", type=int, default=5, help="Top-K comments per topic shown in the final report.")

    # Tags only (for naming parity with your section pipeline)
    ap.add_argument("--seg_n_sections", type=int, default=None,
                    help="Unused for topics; kept for tag naming consistency.")
    ap.add_argument("--seg_penalty", type=int, default=15, help="Unused for topics; kept for tag naming consistency.")
    ap.add_argument("--seg_min_size", type=int, default=6, help="Unused for topics; kept for tag naming consistency.")

    ap.add_argument("--embed_model", default="all-mpnet-base-v2", help="SentenceTransformer model for embeddings.")
    ap.add_argument("--retrieval_topk", type=int, default=200,
                    help="How many candidate comments to retrieve per topic before LLM re-score (0 = score all).")
    ap.add_argument("--rep_k", type=int, default=3, help="Representative thumbnails per topic in the report.")

    ap.add_argument("--rebuild", action="store_true",
                    help="Rebuild DOCX from existing results without running the model.")

    args = ap.parse_args()
    model = args.model or Config.ollama_models[0]

    # Resolve results root
    if args.results_root:
        results_root = os.path.abspath(args.results_root)
    else:
        # Default: ../results_youtube relative to this script
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        results_root = os.path.join(base, "results_youtube")

    print(f"[INFO] Results will be saved to: {results_root}")
    
    # Check absolute path for comments dir if needed
    if not os.path.isabs(args.comments_dir):
         # Try relative to this script first, or just keep as is if it might be relative to cwd
         # But the default is relative to CWD usually or script dir? 
         # The default provided "../../VideoAnalysis/data/comments" is likely relative to the script location if running from there.
         # But usually argparse paths are relative to CWD.
         # For safety let's use abspath if it exists relative to script, else CWD
         script_dir = os.path.dirname(os.path.abspath(__file__))
         candidate = os.path.join(script_dir, args.comments_dir)
         if os.path.exists(candidate):
             args.comments_dir = candidate
         else:
             args.comments_dir = os.path.abspath(args.comments_dir)
    else:
        args.comments_dir = os.path.abspath(args.comments_dir)

    comments_by_video = parse_comments_dir(args.comments_dir)
    if not comments_by_video:
        print("[FATAL] No comments parsed from comments directory.")
        return

    targets = discover_raw_frames(args.frames_path)
    if not targets:
        print("[FATAL] No raw_frames folders found.")
        return

    print(f"[INFO] Found {len(targets)} video(s) to process.")

    for raw_frames in targets:
        print(f"\n=== Processing: {os.path.basename(os.path.dirname(raw_frames))} ===")
        
        # Adjust results root for this specific video if subfolder detected
        current_results_root = results_root
        sub = determine_subfolder(raw_frames)
        if sub:
            current_results_root = os.path.join(results_root, sub)
            
        try:
            if args.rebuild:
                rebuild_docx_only(
                    raw_frames, comments_by_video, args.topics_root, current_results_root, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    seg_n_sections=args.seg_n_sections, seg_penalty=args.seg_penalty,
                    seg_min_size=args.seg_min_size, embed_model=args.embed_model,
                    retrieval_topk=args.retrieval_topk, rep_k=args.rep_k
                )
            else:
                process_one_video(
                    raw_frames, comments_by_video, args.topics_root, current_results_root, model,
                    min_score=args.min_score, top_k_scan=args.top_k,
                    seg_n_sections=args.seg_n_sections, seg_penalty=args.seg_penalty,
                    seg_min_size=args.seg_min_size, embed_model=args.embed_model,
                    retrieval_topk=args.retrieval_topk, rep_k=args.rep_k
                )
        except KeyboardInterrupt:
            print("\n[INTERRUPTED] Stopping cleanly. You can rerun to resume.")
            break
        except Exception as e:
            print(f"[ERROR] {raw_frames}: {e}")


if __name__ == "__main__":
    main()
