#!/usr/bin/env python3
"""
divide_sections.py

Manually segments videos into fixed 10-second intervals.
Outputs JSON and DOCX reports with frames and transcripts.
"""

import os
import sys
import json
import csv
import re
import argparse
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, asdict

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("[ERR] python-docx not installed. Install with 'pip install python-docx' to generate reports.")
    sys.exit(1)

# -------------------- Data Structures --------------------

@dataclass
class SegmentSection:
    segment_index: int
    start_time: float
    end_time: float
    frames: List[str]  # List of filenames or paths
    transcript: str

    def to_dict(self):
        return asdict(self)

# -------------------- Data Loading --------------------

def load_metrics_for_fps(metrics_csv_paths: List[str]) -> Dict[str, float]:
    """Load duration and frame counts to calculate FPS."""
    vid_meta = {} 
    for path in metrics_csv_paths:
        if not os.path.exists(path): continue
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                vid = row.get('video_id') or row.get('VideoID')
                if not vid: continue
                frames = row.get('frames_count') or row.get('frames') or row.get('total_frames')
                dur = row.get('duration_seconds') or row.get('duration') or row.get('DurationSec') or row.get('duration_s')
                
                if frames and dur:
                    try:
                        f_val = float(frames)
                        d_val = float(dur)
                        if vid not in vid_meta:
                            vid_meta[vid] = {'frames': f_val, 'duration': d_val}
                        else:
                            if f_val > 0: vid_meta[vid]['frames'] = f_val
                            if d_val > 0: vid_meta[vid]['duration'] = d_val
                    except ValueError: continue

    fps_map = {}
    for vid, data in vid_meta.items():
        if data['duration'] > 0 and data['frames'] > 0:
            fps_map[vid] = data['frames'] / data['duration']
        else:
            fps_map[vid] = 30.0 # Default if unknown
    return fps_map

def load_transcript(path: str) -> List[Tuple[float, float, str]]:
    """Parses _timestamped.txt or .json. Returns [(start, end, text), ...]"""
    segments = []
    if not os.path.exists(path): return segments
    
    if path.endswith('.json'):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                segs = data.get('segments', [])
                for s in segs:
                    segments.append((float(s.get('start', 0)), float(s.get('end', 0)), s.get('text', '')))
        except Exception as e:
            print(f"[ERR] Failed to load transcript JSON {path}: {e}")
            
    elif path.endswith('.txt'):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                for line in f:
                    # Regex for [HH:MM:SS.mmm → HH:MM:SS.mmm]
                    m = re.search(r'\[(\d{2}):(\d{2}):(\d{2}\.\d+)\s+(?:→|->)\s+(\d{2}):(\d{2}):(\d{2}\.\d+)\]\s+(.*)', line)
                    if m:
                        h1, m1, s1 = int(m.group(1)), int(m.group(2)), float(m.group(3))
                        start_sec = h1*3600 + m1*60 + s1
                        h2, m2, s2 = int(m.group(4)), int(m.group(5)), float(m.group(6))
                        end_sec = h2*3600 + m2*60 + s2
                        content = m.group(7).strip()
                        segments.append((start_sec, end_sec, content))
        except Exception as e:
            print(f"[WARN] Text transcript parse error {path}: {e}")
    return segments

def get_frames_in_range(video_frames_dir: str, start_time: float, end_time: float, fps: float) -> List[str]:
    """Find all frames in the directory that fall within [start_time, end_time)."""
    if not os.path.exists(video_frames_dir):
        return []
    
    frames_in_segment = []
    files = sorted([f for f in os.listdir(video_frames_dir) if f.endswith('.jpg') or f.endswith('.png')])
    
    for f in files:
        # Assuming format frame_{index}.jpg
        nums = re.findall(r'\d+', f)
        if nums:
            f_idx = int(nums[-1])
            # Time of this frame
            t = f_idx / fps
            if start_time <= t < end_time:
                frames_in_segment.append(f)
    
    # Sort by index
    frames_in_segment.sort(key=lambda x: int(re.findall(r'\d+', x)[-1]))
    return frames_in_segment

# -------------------- Report Generation --------------------

def generate_docx_report(video_id: str, segments: List[SegmentSection], video_frames_dir: str, out_path: str):
    doc = Document()
    doc.add_heading(f"Manual Segmentation Analysis: {video_id}", 0)
    
    for sec in segments:
        heading = f"Segment {sec.segment_index}: {int(sec.start_time)}s - {int(sec.end_time)}s"
        doc.add_heading(heading, level=1)
        
        # Transcript
        if sec.transcript:
            doc.add_paragraph(sec.transcript)
        else:
            doc.add_paragraph("(No speech)")
            
        # Frames
        if sec.frames:
            doc.add_heading("Frames", level=2)
            # Create table for frames (up to 5 per row)
            # We want to show ALL frames found in this segment?
            # Or just a few? The user said "insert the actual frames". 
            # If 1 frame per 3.33s, a 10s segment has ~3 frames. Showing all is fine.
            
            # Batch into rows of 3
            cols_per_row = 3
            num_rows = (len(sec.frames) + cols_per_row - 1) // cols_per_row
            table = doc.add_table(rows=num_rows, cols=cols_per_row)
            table.autofit = True
            
            for i, fname in enumerate(sec.frames):
                row_idx = i // cols_per_row
                col_idx = i % cols_per_row
                
                cell = table.rows[row_idx].cells[col_idx]
                p = cell.paragraphs[0]
                p.add_run(fname).bold = True
                p.add_run("\n")
                
                img_path = os.path.join(video_frames_dir, fname)
                try:
                    p.add_run().add_picture(img_path, width=Inches(1.5))
                except Exception:
                    p.add_run("[Image Error]")
                    
    doc.save(out_path)

# -------------------- Main --------------------

def main():
    parser = argparse.ArgumentParser(description="Manual Video Segmentation (10s intervals)")
    parser.add_argument("--transcripts_root", default="../../TranscriptAnalysis/data/transcripts_embodied", help="Directory containing transcripts")
    parser.add_argument("--captions_root", default="../../CaptionAnalysis/data/integrated_caption/frames_embodied_captions_integrated", help="Used only for video discovery if needed, or ignored")
    parser.add_argument("--frames_root", default="../../VideoAnalysis/data/frames/frames_embodied", help="Root dir containing <video_id>/raw_frames")
    parser.add_argument("--metrics_csv", default=["../../CaptionAnalysis/output/Embodied_metrics.csv"], nargs='+', help="CSVs for FPS")
    parser.add_argument("--output_dir", default="../data/manual_segmentation")
    parser.add_argument("--interval", type=int, default=10, help="Segment length in seconds")
    
    args = parser.parse_args()
    
    fps_map = load_metrics_for_fps(args.metrics_csv)
    print(f"[INFO] Loaded FPS data for {len(fps_map)} videos.")
    
    # Discover Videos (using transcripts as source of truth)
    vid_ids = []
    if os.path.exists(args.transcripts_root):
        for f in os.listdir(args.transcripts_root):
            vid = None
            if f.endswith('_timestamped.txt'):
                vid = f.replace('_timestamped.txt', '')
            elif f.endswith('.json') and not f.endswith('topics.json'):
                vid = f.replace('.json', '')
            if vid:
                vid_ids.append(vid)
    
    vid_ids = sorted(list(set(vid_ids)))
    print(f"[INFO] Found {len(vid_ids)} videos to process.")
    
    if not os.path.exists(args.output_dir):
        os.makedirs(args.output_dir)
        
    for vid in vid_ids:
        print(f"\nProcessing {vid}...")
        
        # Load Transcript
        t_path = os.path.join(args.transcripts_root, f"{vid}_timestamped.txt")
        if not os.path.exists(t_path): 
            t_path = os.path.join(args.transcripts_root, f"{vid}.json")
            
        transcripts = load_transcript(t_path)
        
        # FPS
        fps = fps_map.get(vid, 30.0)
        
        # Determine Duration
        max_t = 0
        if transcripts: max_t = max(max_t, transcripts[-1][1])
        # If no transcript, try to guess from frame count in dir?
        # Let's rely on metric csv if available, else max_t from transcript
        
        # Check frames dir
        v_frames_dir = os.path.join(args.frames_root, vid, "raw_frames")
        if not os.path.isdir(v_frames_dir):
             v_frames_dir = os.path.join(args.frames_root, vid)
             
        # Create Segments
        if max_t == 0: max_t = 60 # Default fallback
        
        segments = []
        num_segments = int(math.ceil(max_t / args.interval))
        if num_segments == 0: num_segments = 1
        
        for i in range(num_segments):
            start = i * args.interval
            end = min((i + 1) * args.interval, max_t)
            
            # Filter Transcript
            # Include if midpoint is in range, or any overlap?
            # User wants "transcript in the segment". Let's loosely check overlap.
            seg_text_parts = []
            for t_s, t_e, txt in transcripts:
                # Overlap check: start < t_e AND end > t_s
                if start < t_e and end > t_s:
                    seg_text_parts.append(f"[{int(t_s//60)}:{int(t_s%60):02d}] {txt}")
            
            seg_transcript = "\n".join(seg_text_parts)
            
            # Find Frames
            # Using FPS and frame index logic
            relevant_frames = get_frames_in_range(v_frames_dir, start, end, fps)
            
            segments.append(SegmentSection(
                segment_index=i+1,
                start_time=start,
                end_time=end,
                frames=relevant_frames,
                transcript=seg_transcript
            ))
            
        # 1. Save JSON
        json_out = os.path.join(args.output_dir, f"{vid}_manual_sections.json")
        result = {
            "video_id": vid,
            "duration": max_t,
            "segments": [s.to_dict() for s in segments]
        }
        with open(json_out, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2)
            
        # 2. Save DOCX
        docx_out = os.path.join(args.output_dir, f"{vid}_manual_sections.docx")
        generate_docx_report(vid, segments, v_frames_dir, docx_out)
        print(f"  > Saved {json_out} and {docx_out}")

import math

if __name__ == "__main__":
    main()
